"""
Tests for the Markdown -> HTML -> PDF/DOCX report renderer.

The defect these guard against is specific and was shipped: report Markdown was pushed
into a single text run, so readers received PDFs containing `## 1. Summary`,
`| Spec | Verdict |` and `**Confirmed malicious**` as literal text. Most assertions here
therefore check that markup became *structure* — a real `<h2>`, a real `<table>` — rather
than that a file was produced. A renderer that emits a valid but syntax-transcribing PDF
passes a "did it render" test and fails the users.

Run on the host with:
    python3 -m pytest --noconftest tests/test_report_rendering.py
(the repo conftest imports src/api/main.py, which needs loguru).
"""

import re
import sys
import time
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from src.api.utils import zda_report  # noqa: E402
from src.reporting import (  # noqa: E402
    DocumentMeta,
    MarkdownRenderError,
    markdown_to_docx,
    markdown_to_html,
    markdown_to_pdf,
    render_file,
)
from src.reporting.md_to_pdf import _offline_url_fetcher, split_front_matter  # noqa: E402
from src.reporting.pdf_generator import ReportGenerator, build_scan_markdown  # noqa: E402
from src.reports.generator import _safe_name  # noqa: E402

SAMPLE = """\
# Zero Day Analysis Report

## 1. Summary

A **confirmed malicious** version was found in `keyv@4.5.4`.

## 2. Verdicts

| Spec | Verdict | Basis |
|---|---|---|
| `keyv@4.5.4` | Confirmed | registry |
| `cacheable-lookup` | ~~Rejected~~ | contradicted |

> Important: absence of a hit is not absence of exposure.

```json
{"spec": "keyv@4.5.4"}
```
"""


# --------------------------------------------------------------------------- #
# Markdown -> HTML: the step that makes formatting real
# --------------------------------------------------------------------------- #

def test_markdown_becomes_structure_not_literal_syntax():
    html = markdown_to_html(SAMPLE, DocumentMeta(title="Zero Day Analysis Report"))

    assert "<h2" in html and ">1. Summary<" in html
    assert "## 1. Summary" not in html
    assert "<strong>confirmed malicious</strong>" in html
    assert "**confirmed malicious**" not in html
    assert "<table" in html and "<th>Spec</th>" in html
    assert "| Spec | Verdict | Basis |" not in html
    assert "<s>Rejected</s>" in html          # strikethrough is enabled
    assert "<blockquote>" in html
    assert "<code" in html


def test_raw_html_in_the_body_is_escaped_not_executed():
    """Report bodies are LLM-generated; markup passthrough is an injection channel."""
    html = markdown_to_html(
        '<img src="https://attacker.example/pixel.gif">\n\n<script>alert(1)</script>',
        DocumentMeta(title="T", cover=False, toc=False),
    )
    assert "<script>" not in html
    assert "&lt;script&gt;" in html
    assert "attacker.example" in html  # visible as text, not fetched


def test_headings_get_anchors_that_the_toc_targets():
    html = markdown_to_html(SAMPLE, DocumentMeta(title="Zero Day Analysis Report"))
    anchors = set(re.findall(r'<h[1-6] id="([^"]+)"', html))
    targets = set(re.findall(r'<a href="#([^"]+)"', html))
    assert targets, "no table of contents was generated"
    # An unresolvable target renders as a blank page number, silently.
    assert targets <= anchors


def test_duplicate_headings_get_distinct_anchors():
    html = markdown_to_html(
        "## Coverage\n\ntext\n\n## Coverage\n\nmore",
        DocumentMeta(title="T", cover=False),
    )
    anchors = re.findall(r'<h2 id="([^"]+)"', html)
    assert len(anchors) == len(set(anchors)) == 2


def test_leading_h1_matching_the_cover_title_is_dropped_but_others_are_kept():
    meta = DocumentMeta(title="Zero Day Analysis Report", toc=False)
    kept = markdown_to_html("# Something Else\n\nbody", meta)
    assert ">Something Else<" in kept

    dropped = markdown_to_html("# Zero Day Analysis Report\n\nbody", meta)
    body = dropped.split('<main class="body">', 1)[1]
    assert "Zero Day Analysis Report" not in body


def test_wide_tables_are_tagged_for_condensed_layout():
    five = "| a | b | c | d | e |\n" + "|---" * 5 + "|\n" + "| 1 | 2 | 3 | 4 | 5 |\n"
    four = "| a | b | c | d |\n" + "|---" * 4 + "|\n" + "| 1 | 2 | 3 | 4 |\n"
    assert 'class="wide"' in markdown_to_html(five, DocumentMeta(title="T", cover=False))
    assert 'class="wide"' not in markdown_to_html(four, DocumentMeta(title="T", cover=False))


def test_front_matter_is_split_and_never_printed_as_body():
    data, body = split_front_matter("---\ntitle: Hunt\nsubtitle: npm\n---\n# Body\n")
    assert data == {"title": "Hunt", "subtitle": "npm"}
    assert body.startswith("# Body")

    html = markdown_to_html("---\ntitle: Hunt\nsubtitle: npm\n---\n\n## Section\n")
    assert "<title>Hunt</title>" in html
    assert "subtitle: npm" not in html


def test_malformed_front_matter_falls_back_to_body_text():
    """A broken header must not cost the reader the document."""
    text = "---\ntitle: [unclosed\n---\n\n## Section\n"
    data, body = split_front_matter(text)
    assert data == {}
    assert body == text


# --------------------------------------------------------------------------- #
# Security posture
# --------------------------------------------------------------------------- #

@pytest.mark.parametrize("url", [
    "https://attacker.example/pixel.gif",
    "http://169.254.169.254/latest/meta-data/",
    "ftp://internal.example/x",
])
def test_remote_assets_are_refused_during_render(url):
    """A PDF build that reaches out turns every export into a callback."""
    with pytest.raises(ValueError, match="Refusing to fetch"):
        _offline_url_fetcher(url)


# --------------------------------------------------------------------------- #
# PDF output
# --------------------------------------------------------------------------- #

def test_pdf_is_produced_with_cover_toc_and_body():
    pdf = markdown_to_pdf(SAMPLE, DocumentMeta(
        title="Zero Day Analysis Report",
        subtitle="npm supply chain",
        fields=[("Query", "keyv")],
        generated="2026-08-07 10:00 CDT",
    ))
    assert pdf.startswith(b"%PDF")
    # Cover + contents + body: a single page means the cover swallowed the document.
    assert pdf.count(b"/Type /Page\n") > 1 or b"/Count 3" in pdf or len(pdf) > 5000


def test_pdf_render_is_reproducible_when_the_timestamp_is_fixed():
    """
    Evidence documents must re-render identically, or a reader cannot distinguish an
    edited finding from a re-exported one. The only intentional variable is the
    generated timestamp, so pinning it must pin the output.
    """
    meta = lambda: DocumentMeta(title="T", generated="2026-08-07 10:00 CDT")  # noqa: E731
    first = markdown_to_pdf(SAMPLE, meta())

    # Render something else in between, and specifically something that pulls in a font
    # the sample does not use. `@font-face` registrations accumulate in whatever font
    # configuration the render is given, so on a shared one this call used to shift the
    # next document by a byte — meaning an export's output depended on which reports the
    # worker had served before it. Comparing lengths hid that; compare the bytes.
    markdown_to_pdf("Confirmed ✅\n", DocumentMeta(title="Other", cover=False, toc=False))

    # Cross a second boundary, which is the other thing that used to move the bytes:
    # fontTools stamps the wall clock into every subset font's `head` table, so this
    # assertion held only for two renders that happened to land in the same second. It
    # failed roughly one run in five, and looked like flakiness rather than the missing
    # SOURCE_DATE_EPOCH it was. The sleep is the whole point of the test; do not remove it.
    time.sleep(1.2)

    assert markdown_to_pdf(SAMPLE, meta()) == first


# --------------------------------------------------------------------------- #
# The stylesheet, which fails silently
# --------------------------------------------------------------------------- #

CSS_EXERCISE = """\
# Section

Body with **bold**, `code` and https://example.invalid/a/very/long/advisory/path/here.

## Subsection

| ID | Finding | Detail |
|---|---|---|
| E1 | Thing | Detail text |

```json
{"spec": "keyv@4.5.4"}
```

> A quote.
"""


def test_the_stylesheet_has_no_declarations_weasyprint_discards():
    """
    An unsupported declaration is a warning on stderr, never an error, so the stylesheet
    can rot one property at a time while every render still "succeeds". Two had already
    rotted when this was written, and neither showed up on the host: WeasyPrint 68 took
    both values and 69 rejects them, so the container quietly rendered a different
    document from the one that had been reviewed.

    Asserting on the renderer's own log catches the *next* one too, on either version.
    """
    from tests.pdf_probes import discarded_declarations

    meta = DocumentMeta(title="T", subtitle="s", fields=[("Query", "keyv")])
    ignored = discarded_declarations(lambda: markdown_to_pdf(CSS_EXERCISE, meta))
    assert not ignored, "WeasyPrint discarded stylesheet declarations:\n" + "\n".join(ignored)


def test_bold_text_resolves_to_a_bold_face():
    """
    `strong` carried `font-weight: 650`, which WeasyPrint 69 drops as an invalid value,
    so emphasis rendered at the inherited weight. Nothing failed and nothing looked
    obviously broken; the words a finding leans on just stopped being emphasized.

    Deliberately no heading in the source: an `<h1>` is bold on its own and would satisfy
    this assertion no matter what `strong` resolved to.
    """
    from tests.pdf_probes import embedded_fonts

    pdf = markdown_to_pdf("Plain text with **emphasis** in it.\n",
                          DocumentMeta(title="T", cover=False, toc=False))
    fonts = embedded_fonts(pdf)
    assert any("Bold" in f for f in fonts), f"no bold face was embedded: {fonts}"


def test_digits_do_not_get_captured_by_the_emoji_font():
    """
    Color emoji fonts carry the ASCII digits, because those are the bases of the keycap
    sequences. Pango gives them every digit in the document the moment such a family
    appears anywhere in the stack — ahead of DejaVu Sans, not after it — and the CBDT
    bitmaps draw as nothing. The digits keep their place in the text layer and their
    advance width, so extraction and layout both look fine while the reader sees
    "Part " for "Part 1" and a blank where every count and page number should be.

    Nothing here is visible in the byte count, the page count, or `pdftotext`. The only
    signal is which font the digits were assigned to.
    """
    from tests.pdf_probes import embedded_fonts

    pdf = markdown_to_pdf("Digits 0123456789 and letters ABC.\n",
                          DocumentMeta(title="T", cover=False, toc=False))
    fonts = embedded_fonts(pdf)
    assert not any("Emoji" in f for f in fonts), (
        f"an emoji font was reached by a document with no emoji in it: {fonts}")


# Every mark the report templates and the hunt corpus actually put on a page. A glyph is
# either here and proven, or it is not used — the fences below are checked against this
# list, not against a general claim about emoji support.
REPORT_MARKS = "✅❌⚠🔍🔴🟠🟡🔵🤖🔐🔑👥🚩🛡⏰🚀✓"


def test_emoji_are_not_drawn_from_a_bitmap_only_font():
    """
    The other half of the fence, and the assertion the first version of this test was too
    weak to make. It checked only that *an* emoji font had been embedded — which was true
    the whole time the marks were rendering as blanks, because Noto Color Emoji embeds
    perfectly well and simply draws nothing from its CBDT bitmaps under WeasyPrint.

    What has to hold is that whatever font each mark landed in can actually put ink on
    the page.
    """
    from tests.pdf_probes import bitmap_only_fonts, embedded_fonts

    pdf = markdown_to_pdf(f"Marks {REPORT_MARKS} and digits 0123456789.\n",
                          DocumentMeta(title="T", cover=False, toc=False))
    assert not bitmap_only_fonts(pdf), (
        "these marks were embedded as color bitmaps and will render as gaps: "
        f"{bitmap_only_fonts(pdf)}")
    fonts = embedded_fonts(pdf)
    assert len(fonts) > 1, f"the marks did not reach a symbol face at all: {fonts}"


def _stylesheet_emoji_fences():
    """The `unicode-range` fences in report.css, as (font, set-of-codepoints) pairs."""
    css = (Path(__file__).resolve().parents[1]
           / "src/reporting/assets/report.css").read_text(encoding="utf-8")
    fences = []
    for block in re.findall(r"@font-face\s*\{(.*?)\}", css, re.DOTALL):
        if "Report Emoji" not in block:
            continue
        ranges = re.search(r"unicode-range:\s*([^;]+);", block)
        local = re.findall(r'local\("([^"]+)"\)', block)
        if not ranges:
            continue
        points = set()
        for item in ranges.group(1).split(","):
            item = item.strip().removeprefix("U+")
            low, _, high = item.partition("-")
            points.update(range(int(low, 16), int(high or low, 16) + 1))
        fences.append((local[-1], points))
    return fences


def test_the_emoji_fences_are_disjoint():
    """
    Two faces share the family "Report Emoji" because neither covers the mark set alone.
    Overlapping ranges would leave the choice to whatever order fontconfig happens to
    return, which is not a property to build a document on.
    """
    fences = _stylesheet_emoji_fences()
    assert len(fences) == 2, f"expected two fenced emoji faces, found {len(fences)}"
    overlap = fences[0][1] & fences[1][1]
    assert not overlap, (
        "the emoji fences overlap at " + ", ".join(f"U+{c:04X}" for c in sorted(overlap)[:8]))


@pytest.mark.parametrize("mark", list(REPORT_MARKS))
def test_every_mark_the_reports_use_is_inside_a_fence_that_covers_it(mark):
    """
    The fences are ranges; the fonts behind them are files. A range can name a codepoint
    its font does not have, and the result is a blank on the page with nothing logged.
    Symbola predates Unicode 12 and has no 🟠🟡🟢🟣 at all, which is how the second face
    came to exist — so the ranges are checked against the actual cmaps.

    Skipped where the Linux fallback is not installed: on a macOS host the color font
    earlier in each `src` list serves these, and it is the container that has to be right.
    """
    import glob

    from fontTools.ttLib import TTFont

    fences = _stylesheet_emoji_fences()
    owning = [(font, points) for font, points in fences if ord(mark) in points]
    assert owning, f"{mark} (U+{ord(mark):04X}) is in no fence, so it falls to the system"
    font_name, _ = owning[0]

    matches = glob.glob(f"/usr/share/fonts/**/{font_name.replace(' ', '')}*.ttf",
                        recursive=True)
    if not matches:
        pytest.skip(f"{font_name} is not installed on this host")

    covered = set()
    for table in TTFont(matches[0], fontNumber=0, lazy=True)["cmap"].tables:
        covered |= set(table.cmap)
    assert ord(mark) in covered, (
        f"{mark} (U+{ord(mark):04X}) is fenced to {font_name}, which has no glyph for it")


def test_marks_are_asked_for_in_their_monochrome_form_but_only_in_prose():
    """
    Pango routes an emoji-presentation run to a color font before the CSS stack is read,
    so the selector is what makes the fences reachable at all. It must not reach code:
    a reader copies a command out of a fenced block, and an invisible U+FE0E riding along
    turns a working command into one that fails somewhere else entirely.
    """
    html = markdown_to_html(
        "Confirmed ✅ and ⚠️ warned.\n\n```\ngrep ✅ log.txt\n```\n\nAlso `echo ✅`.\n",
        DocumentMeta(title="T", cover=False, toc=False),
        emoji_text_presentation=True,
    )
    # From <main> onwards: the <head> carries the stylesheet, whose own comments talk
    # about these marks and would otherwise answer for the body.
    _, _, body = html.partition('<main class="body">')
    prose, _, code = body.partition("<pre>")
    assert "✅︎" in prose and "⚠︎" in prose
    assert "️" not in prose, "an emoji-presentation selector survived in prose"
    assert "︎" not in code, "the text-presentation selector leaked into a code block"


@pytest.mark.parametrize("source,expected", [
    ("⚠️", "⚠︎"),              # U+FE0F is replaced, not doubled
    ("✅", "✅︎"),
    ("0123 ABC", "0123 ABC"),                      # nothing outside the fences is touched
    ("→", "→"),                          # an arrow is text already; leave it be
    ("\U0001F468‍\U0001F4BB",                 # a ZWJ sequence would break apart
     "\U0001F468‍\U0001F4BB"),
    ("\U0001F1FA\U0001F1F8",                       # regional indicators pair into a flag
     "\U0001F1FA\U0001F1F8"),
    ("1️⃣", "1︎⃣"),            # a keycap keeps its base and combiner
])
def test_the_text_presentation_rewrite_leaves_sequences_intact(source, expected):
    from src.reporting.md_to_pdf import _force_text_presentation

    assert _force_text_presentation(source) == expected


def test_text_renders_as_glyphs_rather_than_as_empty_boxes():
    """
    python:3.11-slim ships no fonts, and WeasyPrint on a fontless image produces a PDF of
    empty boxes rather than an error — a passing render proves nothing on its own. Run in
    the container this asserts the image's fonts survived, so a Dockerfile that drops the
    font packages fails here instead of shipping unreadable evidence documents.
    """
    from tests.pdf_probes import embedded_fonts

    pdf = markdown_to_pdf("# Heading\n\nBody text.\n",
                          DocumentMeta(title="T", cover=False, toc=False))
    fonts = embedded_fonts(pdf)
    assert fonts, "no font was embedded: this image has no fonts installed"
    assert not all("LastResort" in f for f in fonts), fonts


def test_render_file_writes_each_supported_format(tmp_path):
    source = tmp_path / "doc.md"
    source.write_text(SAMPLE, encoding="utf-8")

    pdf = render_file(source, tmp_path / "doc.pdf")
    assert pdf.read_bytes().startswith(b"%PDF")

    html = render_file(source, tmp_path / "doc.html")
    assert "<table" in html.read_text(encoding="utf-8")

    docx = render_file(source, tmp_path / "doc.docx")
    assert docx.read_bytes().startswith(b"PK")


def test_unsupported_output_format_is_reported(tmp_path):
    source = tmp_path / "doc.md"
    source.write_text("# x\n", encoding="utf-8")
    with pytest.raises(MarkdownRenderError):
        render_file(source, tmp_path / "doc.rtf")


# --------------------------------------------------------------------------- #
# DOCX output — same token stream, so it must agree with the PDF
# --------------------------------------------------------------------------- #

def _docx(payload_bytes):
    import io

    from docx import Document

    return Document(io.BytesIO(payload_bytes))


def test_docx_carries_headings_lists_and_tables_not_markdown_source():
    doc = _docx(markdown_to_docx(
        SAMPLE + "\n- first\n- second\n",
        DocumentMeta(title="Zero Day Analysis Report"),
    ))
    styles = [p.style.name for p in doc.paragraphs]
    text = "\n".join(p.text for p in doc.paragraphs)

    assert "Heading 2" in styles
    assert "List Bullet" in styles
    assert "Intense Quote" in styles
    assert "## 1. Summary" not in text
    assert "**confirmed malicious**" not in text

    assert doc.tables, "the verdict table did not survive"
    verdict = next(t for t in doc.tables if t.rows[0].cells[0].text == "Spec")
    assert [c.text for c in verdict.rows[0].cells] == ["Spec", "Verdict", "Basis"]
    assert verdict.rows[1].cells[0].text == "keyv@4.5.4"


def test_docx_emphasis_becomes_run_formatting():
    doc = _docx(markdown_to_docx(
        "Plain **bold** and *italic* text.",
        DocumentMeta(title="T", cover=False),
    ))
    runs = {r.text: r for p in doc.paragraphs for r in p.runs}
    assert runs["bold"].bold is True
    assert runs["italic"].italic is True
    assert runs["Plain "].bold is False


# --------------------------------------------------------------------------- #
# Zero-day export payload -> Markdown
# --------------------------------------------------------------------------- #

def test_analysis_markdown_passes_the_model_text_through_verbatim():
    """
    Escaping the analysis is the original bug. The analysis IS Markdown; only the
    generated scaffolding around it is escaped.
    """
    payload = {"analysis": "## Findings\n\n| a | b |\n|---|---|\n| 1 | 2 |\n"}
    markdown = zda_report.analysis_markdown(payload)
    assert "## Findings" in markdown
    assert "\\#" not in markdown
    assert "| 1 | 2 |" in markdown


def test_generated_cells_are_escaped_so_a_pipe_cannot_break_the_table():
    payload = {"affected_repositories": [
        {"repository": "org/re|po", "last_updated": "2026-08-01T00:00:00Z", "source": "sbom"},
    ]}
    markdown = zda_report.analysis_markdown(payload)
    assert r"re\|po" in markdown
    assert "2026-08-01" in markdown

    # The real check: the pipe must land inside a cell, not create an extra one. Scoped
    # to the table that carries the row -- the document has several, and counting <th>
    # across all of them would compare this row to a different table's header.
    html = markdown_to_html(markdown, DocumentMeta(title="T", cover=False, toc=False))
    table = next(t for t in re.findall(r"<table.*?</table>", html, re.S)
                 if "org/re|po" in t)
    row = next(r for r in re.findall(r"<tr>(.*?)</tr>", table, re.S) if "org/re|po" in r)
    assert row.count("<td>") == table.count("<th>")


def test_missing_analysis_says_so_rather_than_rendering_none():
    markdown = zda_report.analysis_markdown({"analysis": None})
    assert "None" not in markdown
    assert "No analysis text" in markdown


def test_bullet_paragraphs_become_a_list_and_do_not_swallow_the_next_paragraph():
    sections = [{"heading": "Coverage", "paragraphs": ["• one", "• two", "Trailing note."]}]
    lines = zda_report.sections_to_markdown(sections)
    assert "- one" in lines and "- two" in lines
    # The blank line is what stops "Trailing note." becoming part of the last bullet.
    assert lines[lines.index("- two") + 1] == ""
    html = markdown_to_html("\n".join(lines), DocumentMeta(title="T", cover=False, toc=False))
    assert html.count("<li>") == 2
    assert "<p>Trailing note.</p>" in html


def test_coverage_section_travels_with_the_repository_list():
    """A short list is exactly where the reader needs the blind-spot section."""
    markdown = zda_report.repo_list_markdown({
        "repositories": [],
        "total_repositories": 0,
        "hunt_enabled": False,
        "coverage_notes": ["Endpoint telemetry not queried"],
    })
    assert "No affected repositories found." in markdown
    assert "Coverage and Blind Spots" in markdown
    assert "Endpoint telemetry not queried" in markdown


def test_repo_list_rows_are_numbered_and_defaulted():
    markdown = zda_report.repo_list_markdown({
        "repositories": [{"repository": "org/a"}, {"repository": "org/b", "source": "sbom"}],
        "total_repositories": 2,
    })
    assert "| 1 | org/a | Context match | - |" in markdown
    assert "| 2 | org/b | Context match | sbom |" in markdown


def test_zero_day_payload_renders_end_to_end():
    payload = {
        "query": "keyv",
        "scope": ["all"],
        "analysis": SAMPLE,
        "affected_repositories": [{"repository": "org/a", "source": "sbom"}],
        "hunt_enabled": True,
        "coverage_notes": ["SBOM coverage 22.6%"],
    }
    markdown = zda_report.analysis_markdown(payload)
    meta = DocumentMeta(title="Zero Day Analysis Report",
                        fields=zda_report.document_fields(payload),
                        generated="2026-08-07 10:00 CDT")
    assert markdown_to_pdf(markdown, meta).startswith(b"%PDF")
    assert markdown_to_docx(markdown, meta).startswith(b"PK")


def test_document_fields_include_scope_and_organizations():
    fields = dict(zda_report.document_fields({
        "query": "keyv", "scope": ["all"], "organizations_in_scope": ["Org1", "Org2"],
    }))
    assert fields["Query"] == "keyv"
    assert fields["Scope"] == "all"
    assert fields["Organizations"] == "Org1, Org2"


# --------------------------------------------------------------------------- #
# Scan reports
# --------------------------------------------------------------------------- #

def test_scan_markdown_orders_by_severity_and_keeps_unclassified_findings():
    findings = [
        {"title": "Low thing", "severity": "low"},
        {"title": "Critical thing", "severity": "critical"},
        {"title": "Strange thing", "severity": "not-a-severity"},
    ]
    # Scoped to the evidence section. The briefed document opens with a plan ordered by
    # risk *and* blast radius, where an unclassified finding is weighted as medium and so
    # sits above a low -- a different and deliberate order that this test is not about.
    markdown = build_scan_markdown({"repo_name": "org/a"}, findings, briefed=False)
    order = [markdown.index(t) for t in ("Critical thing", "Low thing", "Strange thing")]
    assert order == sorted(order)
    # A finding with a bad severity must be counted somewhere, not dropped.
    assert "| Unclassified | 1 |" in markdown


def test_scan_report_writes_a_real_pdf(tmp_path):
    output = tmp_path / "scan.pdf"
    assert ReportGenerator().generate_scan_report(
        {"repo_name": "org/a", "scan_id": "s-1"},
        [{"title": "Hardcoded key", "severity": "critical", "file_path": "src/x.py",
          "line_number": 12, "description": "Committed credential."}],
        str(output),
    ) is True
    assert output.read_bytes().startswith(b"%PDF")


def test_scan_report_reports_failure_instead_of_writing_elsewhere(tmp_path):
    """
    The old code silently wrote Markdown to a different path and returned success.
    """
    output = tmp_path / "scan.rtf"
    assert ReportGenerator().generate_scan_report({"repo_name": "org/a"}, [], str(output)) is False
    assert not output.exists()
    assert list(tmp_path.iterdir()) == []


@pytest.mark.parametrize("name,expected", [
    ("SleepNumberInc/api", "SleepNumberInc_api"),
    ("owner/repo.name", "owner_repo.name"),
    ("", "report"),
    ("///", "report"),
])
def test_report_filenames_cannot_contain_a_path_separator(name, expected):
    assert _safe_name(name) == expected
