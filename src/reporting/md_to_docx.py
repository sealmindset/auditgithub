"""
Markdown -> DOCX, from the same token stream the PDF path uses.

Why this exists
---------------
The Word export had the identical defect as the PDF one: the analysis Markdown was
handed to `add_paragraph()` as a single string, so a reader opened a document containing
`## 1. Summary` and `| Spec | Verdict |` as literal text. Fixing only the PDF would have
left the two exports of the *same* report disagreeing about what it says.

Rather than a second parser, this walks the markdown-it token stream produced by
:mod:`.md_to_pdf` and emits python-docx elements. One parse, one interpretation of the
source, two output formats.

Fidelity is deliberately partial. Word has no CSS Paged Media, so there is no cover page
furniture beyond a title block, no table of contents with resolved page numbers, and no
syntax highlighting. Headings, emphasis, links, lists, tables, code blocks and
blockquotes — the structures that carry meaning in these reports — all survive.
"""

from __future__ import annotations

import logging
from typing import Any, List, Optional, Sequence

from .md_to_pdf import DocumentMeta, MarkdownRenderError, _build_parser, split_front_matter

logger = logging.getLogger(__name__)

# Word's built-in heading styles stop at 9; Markdown stops at 6, so no clamping is
# needed beyond guarding against a level of 0.
_MAX_HEADING = 6

_MONO = "Consolas"


def _add_runs(paragraph, tokens: Sequence[Any]) -> None:
    """
    Emit an inline token stream as runs on an existing paragraph.

    Emphasis nests, so bold/italic/strike/code are tracked as counters rather than flags:
    `**bold with *italic* inside**` closes the italic without clearing the bold.
    """
    bold = italic = strike = mono = 0
    link_url: Optional[str] = None

    def run(text: str):
        if not text:
            return
        r = paragraph.add_run(text)
        r.bold = bold > 0
        r.italic = italic > 0
        if strike:
            r.font.strike = True
        if mono:
            r.font.name = _MONO
        if link_url:
            # python-docx has no hyperlink API. Underlined accent text plus the URL
            # printed at the close is honest about the target without hand-writing
            # w:hyperlink relationship XML.
            r.underline = True
        return r

    for token in tokens:
        kind = token.type
        if kind == "text":
            run(token.content)
        elif kind == "code_inline":
            mono += 1
            run(token.content)
            mono -= 1
        elif kind == "strong_open":
            bold += 1
        elif kind == "strong_close":
            bold = max(0, bold - 1)
        elif kind == "em_open":
            italic += 1
        elif kind == "em_close":
            italic = max(0, italic - 1)
        elif kind == "s_open":
            strike += 1
        elif kind == "s_close":
            strike = max(0, strike - 1)
        elif kind == "link_open":
            link_url = token.attrGet("href")
        elif kind == "link_close":
            url = link_url
            link_url = None
            if url:
                italic += 1
                run(f" <{url}>")
                italic -= 1
        elif kind in ("softbreak", "hardbreak"):
            run("\n" if kind == "hardbreak" else " ")
        elif kind == "image":
            # Remote images are refused in the PDF path for exfil reasons; the same
            # applies here, so the alt text stands in for the image.
            alt = token.content or token.attrGet("alt") or "image"
            italic += 1
            run(f"[image: {alt}]")
            italic -= 1


def _paragraph(doc, style: Optional[str]):
    """
    Add a paragraph, degrading to the default style if Word's template lacks the named one.

    A missing style raises KeyError deep inside python-docx. Losing the indent on a
    third-level bullet is a cosmetic loss; losing the export is not.
    """
    if not style:
        return doc.add_paragraph()
    try:
        return doc.add_paragraph(style=style)
    except KeyError:
        logger.debug("DOCX template has no style %r; falling back to default", style)
        return doc.add_paragraph()


def _render_tokens(doc, tokens: Sequence[Any]) -> None:
    """Walk a block-level token stream and append to a python-docx Document."""
    list_stack: List[str] = []
    quote_depth = 0
    index = 0
    total = len(tokens)

    while index < total:
        token = tokens[index]
        kind = token.type

        if kind == "heading_open":
            level = min(_MAX_HEADING, max(1, int(token.tag[1:])))
            paragraph = doc.add_heading("", level=level)
            _add_runs(paragraph, (tokens[index + 1].children or [])
                      if index + 1 < total and tokens[index + 1].type == "inline" else [])
            index += 3
            continue

        if kind == "paragraph_open":
            inline = tokens[index + 1] if index + 1 < total else None
            if list_stack:
                base = "List Bullet" if list_stack[-1] == "ul" else "List Number"
                # Word names nested variants "List Bullet 2"/"List Bullet 3"; past that
                # depth it has none, so the deepest named style is reused.
                depth = min(len(list_stack), 3)
                style = base if depth == 1 else f"{base} {depth}"
            elif quote_depth:
                style = "Intense Quote"
            else:
                style = None
            paragraph = _paragraph(doc, style)
            if inline is not None and inline.type == "inline":
                _add_runs(paragraph, inline.children or [])
            index += 3
            continue

        if kind in ("bullet_list_open", "ordered_list_open"):
            list_stack.append("ul" if kind == "bullet_list_open" else "ol")
            index += 1
            continue

        if kind in ("bullet_list_close", "ordered_list_close"):
            if list_stack:
                list_stack.pop()
            index += 1
            continue

        if kind == "blockquote_open":
            quote_depth += 1
            index += 1
            continue

        if kind == "blockquote_close":
            quote_depth = max(0, quote_depth - 1)
            index += 1
            continue

        if kind in ("fence", "code_block"):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(token.content.rstrip("\n"))
            run.font.name = _MONO
            index += 1
            continue

        if kind == "hr":
            doc.add_paragraph("_" * 60)
            index += 1
            continue

        if kind == "table_open":
            index = _render_table(doc, tokens, index)
            continue

        if kind == "html_block":
            # html: False means this only appears when a caller bypasses the parser
            # config; render it as text rather than dropping it silently.
            doc.add_paragraph(token.content.strip())
            index += 1
            continue

        index += 1


def _render_table(doc, tokens: Sequence[Any], start: int) -> int:
    """Render one table starting at `tokens[start]`; return the index after it."""
    rows: List[List[Any]] = []
    header_count = 0
    index = start + 1
    current: Optional[List[Any]] = None
    in_head = False

    while index < len(tokens) and tokens[index].type != "table_close":
        kind = tokens[index].type
        if kind == "thead_open":
            in_head = True
        elif kind == "thead_close":
            in_head = False
        elif kind == "tr_open":
            current = []
        elif kind == "tr_close":
            if current is not None:
                rows.append(current)
                if in_head:
                    header_count = len(current)
            current = None
        elif kind in ("th_open", "td_open"):
            inline = tokens[index + 1] if index + 1 < len(tokens) else None
            if current is not None:
                current.append(inline.children or [] if inline is not None
                               and inline.type == "inline" else [])
        index += 1

    if not rows:
        return index + 1

    columns = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=columns)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for column in range(columns):
            cell = cells[column]
            paragraph = cell.paragraphs[0]
            if column < len(row):
                _add_runs(paragraph, row[column])
            if row_index == 0 and header_count:
                for run in paragraph.runs:
                    run.bold = True
    return index + 1


def _add_title_block(doc, meta: DocumentMeta) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    mark = doc.add_paragraph()
    mark_run = mark.add_run(meta.classification.upper())
    mark_run.bold = True
    mark.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title = doc.add_heading(meta.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if meta.subtitle:
        subtitle = doc.add_paragraph()
        subtitle.add_run(meta.subtitle).italic = True
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fields = list(meta.fields) + [("Generated", meta.resolved_generated())]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for key, value in fields:
        if not str(value).strip():
            continue
        cells = table.add_row().cells
        cells[0].paragraphs[0].add_run(str(key)).bold = True
        cells[1].text = str(value)
    doc.add_paragraph()


def markdown_to_docx(markdown_text: str, meta: Optional[DocumentMeta] = None) -> bytes:
    """Render Markdown to DOCX bytes."""
    try:
        from docx import Document
    except ImportError as exc:
        raise MarkdownRenderError(
            "DOCX generation requires python-docx. Install with: pip install python-docx"
        ) from exc

    try:
        parser = _build_parser()
    except ImportError as exc:  # pragma: no cover - dependency is declared
        raise MarkdownRenderError(
            "Markdown rendering requires markdown-it-py. Install with: "
            "pip install markdown-it-py"
        ) from exc

    from .md_to_pdf import _strip_leading_h1, meta_from_front_matter

    front_matter, body = split_front_matter(markdown_text or "")
    if meta is None:
        meta = meta_from_front_matter(front_matter, fallback_title="Report")

    doc = Document()
    if meta.cover:
        _add_title_block(doc, meta)

    try:
        tokens = parser.parse(body, {})
        if meta.cover:
            _strip_leading_h1(tokens, meta.title)
        _render_tokens(doc, tokens)
    except Exception as exc:
        logger.exception("DOCX rendering failed")
        raise MarkdownRenderError(f"DOCX layout failed: {exc}") from exc

    if meta.footer_note:
        note = doc.add_paragraph()
        note.add_run(meta.footer_note).italic = True

    import io

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
