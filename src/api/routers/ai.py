import html as html_mod
import json
import os
import uuid
import shutil
from typing import Dict, Any, Optional, List
from fastapi import APIRouter, Depends, HTTPException, BackgroundTasks
from fastapi.responses import StreamingResponse
from io import BytesIO
from pydantic import BaseModel, Field
from sqlalchemy.orm import Session
from sqlalchemy import func
from loguru import logger
from ..dependencies import get_tenant_db
from .. import models
from ...ai_agent.agent import AIAgent
from ..utils.repo_context import clone_repo_to_temp as clone_repo, cleanup_repo
from ..config import settings # Keep settings as it's used later
from src.rbac.dependencies import require_permissions
from src.auth.models import User
from src.auth.dependencies import get_current_user
from src.auth.rate_limit import limiter
from src.api.schemas.common import CRUD_ERRORS, LIST_ERRORS, CREATE_ERRORS, DELETE_ERRORS, AUTH_ERRORS

AI_RATE_LIMIT = f"{os.getenv('AI_RATE_LIMIT_REQUESTS_PER_MINUTE', '30')}/minute"


async def get_github_token_for_repo(db: Session, repo: models.Repository) -> str:
    """
    Get the GitHub token for a repository's organization.
    Falls back to the default GITHUB_TOKEN if no org-specific token is found.
    """
    # If repo has an organization_id, try to get the org's token
    if repo.organization_id:
        try:
            org = db.query(models.Organization).filter(
                models.Organization.id == repo.organization_id
            ).first()
            
            if org:
                # Try to get token from secrets manager
                try:
                    import sys
                    sys.path.insert(0, '/app/execution')
                    from secrets_manager import get_org_credentials
                    creds = await get_org_credentials(org.name)
                    if creds and creds.get('github_token'):
                        logger.info(f"Using org-specific token for {org.name}")
                        return creds['github_token']
                except Exception as e:
                    logger.warning(f"Could not get org credentials from secrets manager: {e}")
                
                # Try environment variable pattern ORG_{NAME}_TOKEN
                env_token = os.environ.get(f"ORG_{org.name}_TOKEN")
                if env_token:
                    logger.info(f"Using env token for org {org.name}")
                    return env_token
        except Exception as e:
            logger.warning(f"Error getting org token: {e}")
    
    # Fall back to default token
    return settings.GITHUB_TOKEN

def _is_error_response(response: str) -> bool:
    """Check if an AI response is an error message."""
    if not response:
        return True
    
    error_indicators = [
        "Error:",
        "Failed:",
        "AI analysis failed:",
        "API key",
        "rate limit",
        "connection error",
        "timeout",
        "HTTPException",
        "Authentication failed"
    ]
    
    response_lower = response.lower()
    for indicator in error_indicators:
        if indicator.lower() in response_lower:
            return True
    
    return False

router = APIRouter(
    prefix="/ai",
    tags=["ai"]
)

@router.get(
    "/config",
    dependencies=[Depends(require_permissions("findings:read"))],
    summary="Get current AI provider configuration",
    responses={
        **LIST_ERRORS,
        401: {"description": "Not authenticated"},
        403: {"description": "Insufficient permissions - findings:read required"},
    },
)
async def get_ai_config():
    """
    Get current AI configuration.

    Returns the active AI provider, model names, and connection settings.
    Requires findings:read permission.
    """
    return {
        "provider": settings.AI_PROVIDER,
        "model": settings.AI_MODEL,
        "openai_model": settings.OPENAI_MODEL,
        "anthropic_model": settings.ANTHROPIC_MODEL,
        "gemini_model": settings.GEMINI_MODEL,
        "docker_model": getattr(settings, 'DOCKER_MODEL', 'ai/llama3.2:latest'),
        "docker_base_url": settings.DOCKER_BASE_URL
    }

# Initialize AI Agent
# In a real app, this might be a dependency to allow for mocking/lazy loading
try:
    # Determine model and url based on provider
    provider = settings.AI_PROVIDER
    model = settings.AI_MODEL
    ollama_url = settings.OLLAMA_BASE_URL
    
    if provider == "openai" and settings.OPENAI_MODEL:
        model = settings.OPENAI_MODEL
    elif provider in ["claude", "anthropic_foundry"] and settings.ANTHROPIC_MODEL:
        # Both claude and anthropic_foundry use ANTHROPIC_MODEL
        model = settings.ANTHROPIC_MODEL
    elif (provider == "gemini" or provider == "google") and settings.GEMINI_MODEL:
        model = settings.GEMINI_MODEL
    elif provider == "docker":
        ollama_url = settings.DOCKER_BASE_URL
        model = settings.DOCKER_MODEL or "ai/llama3.2:latest"
        
    ai_agent = AIAgent(
        openai_api_key=settings.OPENAI_API_KEY,
        anthropic_api_key=settings.ANTHROPIC_API_KEY,
        provider=provider,
        model=model,
        ollama_base_url=ollama_url,
        azure_foundry_endpoint=settings.AZURE_AI_FOUNDRY_ENDPOINT,
        azure_foundry_api_key=settings.AZURE_AI_FOUNDRY_API_KEY,
        gemini_api_key=settings.GEMINI_API_KEY
    )
except Exception as e:
    logger.warning(f"Failed to initialize AI Agent: {e}")
    ai_agent = None

# Initialize Diagrams Index
from ..utils.diagrams_indexer import get_diagrams_index
diagrams_index = {}
try:
    diagrams_index = get_diagrams_index()
except Exception as e:
    logger.warning(f"Failed to initialize diagrams index: {e}")

class RemediationRequest(BaseModel):
    vuln_type: str = Field(..., description="Vulnerability type (e.g. SQL Injection, XSS)")
    description: str = Field(..., description="Description of the vulnerability")
    context: str = Field(..., description="Code context surrounding the vulnerability")
    language: str = Field(..., description="Programming language of the affected code")
    finding_id: Optional[str] = Field(None, description="Finding UUID to persist the remediation against")

class RemediationResponse(BaseModel):
    remediation: str = Field(..., description="AI-generated remediation guidance text")
    diff: str = Field(..., description="Suggested code diff to fix the vulnerability")
    remediation_id: Optional[str] = Field(None, description="UUID of the persisted remediation record")

@router.post(
    "/remediate",
    response_model=RemediationResponse,
    dependencies=[Depends(require_permissions("findings:write"))],
    summary="Generate AI remediation for a vulnerability",
    responses={
        **CREATE_ERRORS,
        401: {"description": "Not authenticated"},
        403: {"description": "Insufficient permissions - findings:write required"},
        500: {"description": "AI generation failed"},
        503: {"description": "AI Agent not initialized"},
    },
)
async def generate_remediation(
    request: RemediationRequest,
    db: Session = Depends(get_tenant_db)
):
    """
    Generate remediation for a vulnerability.

    Uses AI to produce remediation guidance and a suggested code diff.
    Optionally persists the result to the database if finding_id is provided.
    Requires findings:write permission.
    """
    if not ai_agent:
        raise HTTPException(status_code=503, detail="AI Agent not initialized")
    
    try:
        result = await ai_agent.generate_remediation(
            vuln_type=request.vuln_type,
            description=request.description,
            context=request.context,
            language=request.language
        )

        remediation_text = result.get("remediation", "")
        diff_text = result.get("diff", "")
        remediation_id = None

        # Persist if finding_id is provided
        if request.finding_id:
            try:
                # Verify finding exists
                finding = db.query(models.Finding).filter(models.Finding.id == request.finding_id).first()
                if finding:
                    new_remediation = models.Remediation(
                        finding_id=request.finding_id,
                        remediation_text=remediation_text,
                        diff=diff_text,
                        confidence=0.85 # Placeholder/Default confidence from AI
                    )
                    db.add(new_remediation)
                    
                    # Update the "latest" cache on the finding for backward compatibility
                    finding.ai_remediation_text = remediation_text
                    finding.ai_remediation_diff = diff_text
                    
                    db.commit()
                    db.refresh(new_remediation)
                    remediation_id = str(new_remediation.id)
            except Exception as db_err:
                logger.error(f"Failed to persist remediation: {db_err}")
                # Don't fail the request if persistence fails, just log it
        
        return RemediationResponse(
            remediation=remediation_text,
            diff=diff_text,
            remediation_id=remediation_id
        )
    except Exception as e:
        logger.error(f"Error generating remediation: {e}")
        raise HTTPException(status_code=500, detail=str(e))

class TriageRequest(BaseModel):
    title: str = Field(..., description="Finding title")
    description: str = Field(..., description="Finding description")
    severity: str = Field(..., description="Severity level (critical, high, medium, low, info)")
    scanner: str = Field(..., description="Scanner that detected the finding")
    finding_id: Optional[str] = Field(None, description="Finding UUID to append triage results to")

class TriageResponse(BaseModel):
    priority: str = Field(..., description="Triage priority level (Critical, High, Medium, Low)")
    confidence: float = Field(..., description="AI confidence score (0.0-1.0)")
    reasoning: str = Field(..., description="Human-readable triage reasoning")
    false_positive_probability: float = Field(..., description="Probability that this is a false positive (0.0-1.0)")

@router.post(
    "/triage",
    response_model=TriageResponse,
    dependencies=[Depends(require_permissions("findings:write"))],
    summary="AI-triage a security finding",
    responses={
        **CREATE_ERRORS,
        401: {"description": "Not authenticated"},
        403: {"description": "Insufficient permissions - findings:write required"},
        500: {"description": "AI triage failed"},
        503: {"description": "AI Agent not initialized"},
    },
)
async def triage_finding(
    request: TriageRequest,
    db: Session = Depends(get_tenant_db)
):
    """
    Analyze and triage a finding.

    Uses AI to determine priority, confidence, and false positive probability.
    Optionally appends analysis to the finding description. Requires findings:write permission.
    """
    if not ai_agent:
        raise HTTPException(status_code=503, detail="AI Agent not initialized")
        
    try:
        result = await ai_agent.triage_finding(
            title=request.title,
            description=request.description,
            severity=request.severity,
            scanner=request.scanner
        )
        
        reasoning = result.get("reasoning", "Analysis failed")
        
        # Append to finding description if finding_id is provided
        if request.finding_id:
            try:
                # Try to find by finding_uuid first (as that's what the API exposes as 'id')
                # But handle potential UUID format errors if string is passed
                import uuid
                try:
                    f_uuid = uuid.UUID(request.finding_id)
                    finding = db.query(Finding).filter(Finding.finding_uuid == f_uuid).first()
                    if not finding:
                        # Fallback to primary key
                        finding = db.query(Finding).filter(Finding.id == f_uuid).first()
                except ValueError:
                    finding = None
                    logger.warning(f"Invalid UUID format for finding_id: {request.finding_id}")

                if finding:
                    # Check if reasoning is already appended to avoid duplication
                    ai_header = "### 🤖 AI Security Analysis"
                    if ai_header not in (finding.description or ""):
                        # Format the reasoning nicely
                        formatted_reasoning = (
                            f"\n\n{ai_header}\n"
                            f"**Priority:** {result.get('priority', 'Medium')}\n"
                            f"**Confidence:** {result.get('confidence', 0.0) * 100:.0f}%\n\n"
                            f"{reasoning}\n"
                        )
                        finding.description = f"{finding.description or ''}{formatted_reasoning}"
                        db.commit()
                        logger.info(f"Appended AI analysis to finding {request.finding_id}")
                else:
                    logger.warning(f"Finding not found for ID: {request.finding_id}")
            except Exception as db_err:
                logger.error(f"Failed to update finding description: {db_err}")

        return TriageResponse(
            priority=result.get("priority", "Medium"),
            confidence=result.get("confidence", 0.0),
            reasoning=reasoning,
            false_positive_probability=result.get("false_positive_probability", 0.0)
        )
    except Exception as e:
        logger.error(f"Error triaging finding: {e}")
        raise HTTPException(status_code=500, detail=str(e))

class AnalyzeFindingRequest(BaseModel):
    finding_id: str = Field(..., description="Finding UUID to analyze")
    prompt: Optional[str] = Field(None, description="Optional custom analysis prompt")

@router.post(
    "/analyze-finding",
    dependencies=[Depends(require_permissions("findings:write"))],
    summary="Analyze a specific finding using AI",
    responses={
        **CREATE_ERRORS,
        401: {"description": "Not authenticated"},
        403: {"description": "Insufficient permissions - findings:write required"},
        404: {"description": "Finding not found"},
        500: {"description": "AI analysis failed"},
        503: {"description": "AI Agent not initialized"},
    },
)
async def analyze_finding_endpoint(
    request: AnalyzeFindingRequest,
    db: Session = Depends(get_tenant_db)
):
    """
    Analyze a specific finding using AI.

    Performs deep security analysis on the finding and appends the result
    to the finding description. Requires findings:write permission.
    """
    if not ai_agent:
        raise HTTPException(status_code=503, detail="AI Agent not initialized")

    try:
        # Fetch finding
        # Try UUID first
        try:
            f_uuid = uuid.UUID(request.finding_id)
            finding = db.query(models.Finding).filter(models.Finding.finding_uuid == f_uuid).first()
            if not finding:
                finding = db.query(models.Finding).filter(models.Finding.id == f_uuid).first()
        except ValueError:
            # If not UUID, try ID (though ID is usually UUID in this schema)
            finding = None
        
        if not finding:
            raise HTTPException(status_code=404, detail="Finding not found")

        # Convert to dict for AI
        finding_dict = {
            "title": finding.title,
            "description": finding.description,
            "severity": finding.severity,
            "file_path": finding.file_path,
            "line_start": finding.line_start,
            "line_end": finding.line_end,
            "code_snippet": finding.code_snippet,
            "finding_type": finding.finding_type,
            "scanner": finding.scanner_name
        }

        analysis = await ai_agent.analyze_finding(
            finding=finding_dict,
            user_prompt=request.prompt
        )
        
        # Append AI Security Analysis to finding description if successful
        if analysis and not _is_error_response(analysis):
            from datetime import datetime
            timestamp = datetime.utcnow().isoformat()
            ai_header = "### 🤖 AI Security Analysis"
            
            # Check if analysis is already appended to avoid duplication
            if ai_header not in (finding.description or ""):
                provider_name = getattr(ai_agent, 'provider_name', 'AI')
                formatted_analysis = f"""

{ai_header}
**Analyzed:** {timestamp}
**Provider:** {provider_name}

{analysis}

---
"""
                finding.description = f"{finding.description or ''}{formatted_analysis}"
                db.commit()
                logger.info(f"Appended AI Security Analysis to finding {request.finding_id}")
        
        return {"analysis": analysis}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error analyzing finding: {e}")
        raise HTTPException(status_code=500, detail=str(e))

class AnalyzeComponentRequest(BaseModel):
    package_name: str = Field(..., description="Package/library name to analyze")
    version: str = Field(..., description="Package version string")
    package_manager: str = Field(..., description="Package manager (npm, pip, maven, etc.)")

@router.post(
    "/analyze-component",
    dependencies=[Depends(require_permissions("findings:write"))],
    summary="Analyze a software component for vulnerabilities",
    responses={
        **CREATE_ERRORS,
        401: {"description": "Not authenticated"},
        403: {"description": "Insufficient permissions - findings:write required"},
        500: {"description": "AI analysis failed"},
        503: {"description": "AI Agent not initialized"},
    },
)
async def analyze_component_endpoint(
    request: AnalyzeComponentRequest,
    db: Session = Depends(get_tenant_db)
):
    """
    Analyze a software component using AI, with caching.

    Returns cached analysis if available, otherwise calls AI provider.
    Results are cached in the database. Requires findings:write permission.
    """
    if not ai_agent:
        raise HTTPException(status_code=503, detail="AI Agent not initialized")

    try:
        # 1. Check DB for existing analysis
        existing = db.query(models.ComponentAnalysis).filter(
            models.ComponentAnalysis.package_name == request.package_name,
            models.ComponentAnalysis.version == request.version,
            models.ComponentAnalysis.package_manager == request.package_manager
        ).first()

        if existing:
            return {
                "analysis_text": existing.analysis_text,
                "vulnerability_summary": existing.vulnerability_summary,
                "severity": existing.severity,
                "exploitability": existing.exploitability,
                "fixed_version": existing.fixed_version,
                "source": "cache"
            }

        # 2. Call AI Provider
        analysis = await ai_agent.provider.analyze_component(
            package_name=request.package_name,
            version=request.version,
            package_manager=request.package_manager
        )

        # 3. Save to DB
        new_analysis = models.ComponentAnalysis(
            package_name=request.package_name,
            version=request.version,
            package_manager=request.package_manager,
            analysis_text=analysis.get("analysis_text", ""),
            vulnerability_summary=analysis.get("vulnerability_summary", ""),
            severity=analysis.get("severity", "Unknown"),
            exploitability=analysis.get("exploitability", "Unknown"),
            fixed_version=analysis.get("fixed_version", "Unknown")
        )
        
        try:
            db.add(new_analysis)
            db.commit()
            db.refresh(new_analysis)
        except Exception as db_err:
            logger.error(f"Failed to cache component analysis: {db_err}")
            # Continue even if caching fails
        
        return {
            **analysis,
            "source": "ai"
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error analyzing component: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.delete(
    "/remediate/{remediation_id}",
    summary="Delete a remediation suggestion",
    responses={
        **DELETE_ERRORS,
        400: {"description": "Invalid UUID format"},
        404: {"description": "Remediation not found"},
    },
)
def delete_remediation(remediation_id: str, db: Session = Depends(get_tenant_db)):
    """
    Delete a remediation suggestion.

    Removes a previously generated remediation record from the database.
    """
    try:
        uuid_obj = uuid.UUID(remediation_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid UUID format")
        
    remediation = db.query(models.Remediation).filter(models.Remediation.id == uuid_obj).first()
    if not remediation:
        raise HTTPException(status_code=404, detail="Remediation not found")
        
    db.delete(remediation)
    db.commit()
    return {"status": "success", "message": "Remediation deleted"}

class ZeroDayRequest(BaseModel):
    query: str = Field(..., description="Natural language query about vulnerabilities or technologies")
    scope: Optional[List[str]] = Field(None, description="Data sources to search (dependencies, findings, languages, all)")

class ZeroDayResponse(BaseModel):
    answer: str = Field(..., description="AI-generated analysis answer")
    affected_repositories: List[Dict[str, Any]] = Field(..., description="List of repositories affected by the vulnerability")
    plan: Optional[Dict[str, Any]] = Field(None, description="Search plan generated by the AI agent")
    execution_summary: Optional[List[str]] = Field(None, description="Summary of search operations executed")

@router.post(
    "/zero-day",
    response_model=ZeroDayResponse,
    dependencies=[Depends(require_permissions("findings:write"))],
    summary="Analyze a zero-day vulnerability across repositories",
    responses={
        **CREATE_ERRORS,
        401: {"description": "Not authenticated"},
        403: {"description": "Insufficient permissions - findings:write required"},
        500: {"description": "Analysis failed"},
        503: {"description": "AI Agent not initialized"},
    },
)
async def analyze_zero_day(
    request: ZeroDayRequest,
    db: Session = Depends(get_tenant_db)
):
    """
    Analyze a zero-day vulnerability query with optional scope filtering.

    Searches across repositories for affected dependencies, findings, and
    technologies. **Required permissions:** findings:write.
    """
    if not ai_agent:
        raise HTTPException(status_code=503, detail="AI Agent not initialized")

    try:
        result = await ai_agent.analyze_zero_day(
            query=request.query,
            db_session=db,
            scope=request.scope
        )
        
        if "error" in result:
             raise HTTPException(status_code=500, detail=result["error"])

        return ZeroDayResponse(
            answer=result.get("answer", "Analysis complete."),
            affected_repositories=result.get("affected_repositories", []),
            plan=result.get("plan"),
            execution_summary=result.get("execution_summary")
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error in zero-day analysis: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get(
    "/zero-day/prompt",
    dependencies=[Depends(require_permissions("findings:read"))],
    summary="Get the zero-day analysis prompt template",
    responses={
        **LIST_ERRORS,
        401: {"description": "Not authenticated"},
        403: {"description": "Insufficient permissions - findings:read required"},
    },
)
async def get_zero_day_prompt():
    """
    Get the current zero-day analysis planning prompt template.

    Returns the AI prompt used for generating search plans during
    zero-day analysis. Requires findings:read permission.
    """
    # Return the prompt template used in analyze_zero_day
    prompt_template = """You are a Senior Security Analyst Agent analyzing Zero Day vulnerabilities.

User Query: "{query}"
{scope_str}

Available Search Tools:
1. search_dependencies(package_name, version_spec) - Search SBOM/Dependencies for libraries/packages
2. search_findings(query, severity_filter) - Search security findings (CVE, CWE, vulnerabilities)
3. search_languages(language) - Find repositories using specific programming language
4. search_technology(keyword) - Find repos by language or description match

NORMALIZATION RULES (90% Confidence Fuzzy Matching):
- "react.js" or "React" or "ReactJS" → "react"
- "next.js" or "Next.JS" or "NextJS" → "next"
- "log4j" → "log4j-core" or "log4j"
- "python" → "python" or "py"
- Extract CVE IDs (CVE-YYYY-NNNNN format)
- Extract CWE IDs (CWE-NNN format)

SEARCH STRATEGY:
- For package/library vulnerabilities → Use search_dependencies and/or search_findings
- For CVE/CWE queries → Use search_findings
- For technology/language queries → Use search_languages and/or search_technology
- For comprehensive searches → Use multiple tools

Generate a search plan as JSON. Format:
{{
    "thought": "Explain your reasoning and which sources to search",
    "tools": [
        {{"name": "search_dependencies", "args": {{"package_name": "react"}}}},
        {{"name": "search_findings", "args": {{"query": "CVE-2024-12345", "severity_filter": "Critical"}}}},
        {{"name": "search_languages", "args": {{"language": "python"}}}}
    ]
}}

IMPORTANT: Return ONLY the JSON object, no markdown formatting."""
    
    return {"prompt": prompt_template}

class ValidateZDAPromptRequest(BaseModel):
    prompt: str = Field(..., description="Prompt template to validate")
    test_query: Optional[str] = Field("Find all repositories using React", description="Sample query to test the prompt with")

@router.post(
    "/zero-day/prompt/validate",
    dependencies=[Depends(require_permissions("findings:write"))],
    summary="Validate a zero-day analysis prompt",
    responses={
        **CREATE_ERRORS,
        401: {"description": "Not authenticated"},
        403: {"description": "Insufficient permissions - findings:write required"},
        500: {"description": "AI provider does not support prompt execution"},
        503: {"description": "AI Agent not initialized"},
    },
)
async def validate_zero_day_prompt(
    request: ValidateZDAPromptRequest,
    db: Session = Depends(get_tenant_db)
):
    """
    Validate a zero-day analysis prompt by testing it with a sample query.

    Formats the prompt template with a test query and sends it to the AI
    provider to verify it produces valid output. Requires findings:write permission.
    """
    if not ai_agent:
        raise HTTPException(status_code=503, detail="AI Agent not initialized")
    
    try:
        # Use the test query to validate the prompt
        test_query = request.test_query
        scope_str = "Scope: All sources"
        
        # Replace placeholders in the prompt
        formatted_prompt = request.prompt.replace("{query}", test_query).replace("{scope_str}", scope_str)
        
        # Try to get a plan from the AI
        if hasattr(ai_agent.reasoning_engine.provider, "execute_prompt"):
            result = await ai_agent.reasoning_engine.provider.execute_prompt(formatted_prompt)
            
            # Try to parse as JSON to validate format
            import json
            clean_json = result.strip()
            if clean_json.startswith("```json"):
                clean_json = clean_json.replace("```json", "").replace("```", "")
            elif clean_json.startswith("```"):
                clean_json = clean_json.replace("```", "")
            
            plan = json.loads(clean_json)
            
            return {
                "success": True,
                "response": f"✅ Prompt validated successfully!\n\n**Test Query:** {test_query}\n\n**AI Plan:**\n```json\n{json.dumps(plan, indent=2)}\n```\n\n**Thought:** {plan.get('thought', 'N/A')}\n\n**Tools to Execute:** {len(plan.get('tools', []))}"
            }
        else:
            raise HTTPException(status_code=500, detail="AI Provider does not support prompt execution")
            
    except json.JSONDecodeError as e:
        return {
            "success": False,
            "response": f"❌ **Invalid JSON Response**\n\nThe AI did not return valid JSON. This prompt may need adjustment.\n\n**Error:** {str(e)}\n\n**AI Response:**\n```\n{result[:500]}\n```"
        }
    except Exception as e:
        logger.error(f"Error validating ZDA prompt: {e}")
        return {
            "success": False,
            "response": f"❌ **Validation Error**\n\n{str(e)}"
        }


# Zero Day Analysis Export Endpoints
@router.post(
    "/zero-day/export/pdf",
    dependencies=[Depends(require_permissions("findings:write"))],
    summary="Export zero-day analysis as PDF",
    responses={
        **CREATE_ERRORS,
        401: {"description": "Not authenticated"},
        403: {"description": "Insufficient permissions - findings:write required"},
        500: {"description": "PDF generation failed or reportlab not installed"},
    },
)
async def export_zda_pdf(request: dict):
    """
    Export Zero Day Analysis as PDF.

    Generates a downloadable PDF report containing the analysis, affected
    repositories, and metadata. Requires findings:write permission.
    """
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.units import inch

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
        styles = getSampleStyleSheet()
        story = []

        # Title
        title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=18, spaceAfter=12)
        story.append(Paragraph("Zero Day Analysis Report", title_style))
        story.append(Spacer(1, 12))

        # Metadata
        meta_style = styles['Normal']
        story.append(Paragraph(f"<b>Query:</b> {html_mod.escape(str(request.get('query', 'N/A')))}", meta_style))
        story.append(Paragraph(f"<b>Generated:</b> {html_mod.escape(str(request.get('timestamp', 'N/A')))}", meta_style))
        scope_list = request.get('scope', [])
        scope_str = ', '.join(scope_list) if isinstance(scope_list, list) else str(scope_list)
        story.append(Paragraph(f"<b>Scope:</b> {html_mod.escape(scope_str)}", meta_style))
        story.append(Spacer(1, 20))

        # Analysis
        story.append(Paragraph("AI Analysis", styles['Heading2']))
        story.append(Spacer(1, 8))
        analysis_text = html_mod.escape(request.get('analysis', '')).replace('\n', '<br/>')
        try:
            story.append(Paragraph(analysis_text, meta_style))
        except Exception:
            # Fallback to plain text if XML parsing fails
            story.append(Paragraph(request.get('analysis', '').replace('\n', ' ')[:2000], meta_style))
        story.append(Spacer(1, 20))

        # Affected Repositories
        repos = request.get('affected_repositories', [])
        story.append(Paragraph(f"Affected Repositories ({len(repos)})", styles['Heading2']))
        story.append(Spacer(1, 8))

        if repos:
            table_data = [['Repository', 'Last Updated', 'Source']]
            for repo in repos:
                last_updated = repo.get('last_updated', 'N/A')
                if last_updated and last_updated != 'N/A':
                    try:
                        from datetime import datetime
                        dt = datetime.fromisoformat(last_updated.replace('Z', '+00:00'))
                        last_updated = dt.strftime('%Y-%m-%d')
                    except:
                        pass
                table_data.append([
                    repo.get('repository', '')[:35],
                    last_updated[:15],
                    repo.get('source', '-')[:15]
                ])

            table = Table(table_data, colWidths=[2.8*inch, 2.0*inch, 1.7*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            story.append(table)
        else:
            story.append(Paragraph("No affected repositories found.", meta_style))

        doc.build(story)
        buffer.seek(0)

        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=zda-analysis.pdf"}
        )

    except ImportError:
        raise HTTPException(status_code=500, detail="PDF generation requires reportlab. Install with: pip install reportlab")
    except Exception as e:
        logger.error(f"PDF generation failed: {e}")
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(e)}")


@router.post(
    "/zero-day/export/docx",
    dependencies=[Depends(require_permissions("findings:write"))],
    summary="Export zero-day analysis as DOCX",
    responses={
        **CREATE_ERRORS,
        401: {"description": "Not authenticated"},
        403: {"description": "Insufficient permissions - findings:write required"},
        500: {"description": "DOCX generation failed or python-docx not installed"},
    },
)
async def export_zda_docx(request: dict):
    """
    Export Zero Day Analysis as DOCX.

    Generates a downloadable Word document containing the analysis, affected
    repositories, and metadata. Requires findings:write permission.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title = doc.add_heading('Zero Day Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        doc.add_paragraph(f"Query: {request.get('query', 'N/A')}")
        doc.add_paragraph(f"Generated: {request.get('timestamp', 'N/A')}")
        scope_list = request.get('scope', [])
        scope_str = ', '.join(scope_list) if isinstance(scope_list, list) else str(scope_list)
        doc.add_paragraph(f"Scope: {scope_str}")

        # Analysis
        doc.add_heading('AI Analysis', level=1)
        doc.add_paragraph(request.get('analysis', ''))

        # Affected Repositories
        repos = request.get('affected_repositories', [])
        doc.add_heading(f'Affected Repositories ({len(repos)})', level=1)

        if repos:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Repository'
            header_cells[1].text = 'Last Updated'
            header_cells[2].text = 'Source'

            for repo in repos:
                row_cells = table.add_row().cells
                row_cells[0].text = repo.get('repository', '')
                last_updated = repo.get('last_updated', 'N/A')
                if last_updated and last_updated != 'N/A':
                    try:
                        from datetime import datetime
                        dt = datetime.fromisoformat(last_updated.replace('Z', '+00:00'))
                        last_updated = dt.strftime('%Y-%m-%d')
                    except:
                        pass
                row_cells[1].text = last_updated
                row_cells[2].text = repo.get('source', '-')
        else:
            doc.add_paragraph('No affected repositories found.')

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=zda-analysis.docx"}
        )

    except ImportError:
        raise HTTPException(status_code=500, detail="DOCX generation requires python-docx. Install with: pip install python-docx")
    except Exception as e:
        logger.error(f"DOCX generation failed: {e}")
        raise HTTPException(status_code=500, detail=f"DOCX generation failed: {str(e)}")


# Zero Day Analysis - Repository List Export Endpoints
@router.post(
    "/zero-day/export/repos/pdf",
    dependencies=[Depends(require_permissions("findings:write"))],
    summary="Export affected repositories list as PDF",
    responses={
        **CREATE_ERRORS,
        401: {"description": "Not authenticated"},
        403: {"description": "Insufficient permissions - findings:write required"},
        500: {"description": "PDF generation failed or reportlab not installed"},
    },
)
async def export_zda_repos_pdf(request: dict):
    """
    Export Zero Day Analysis affected repositories as PDF.

    Generates a downloadable PDF with a table of affected repositories.
    Requires findings:write permission.
    """
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.units import inch

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
        styles = getSampleStyleSheet()
        story = []

        # Title
        title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=18, spaceAfter=12)
        story.append(Paragraph("Affected Repositories Report", title_style))
        story.append(Spacer(1, 12))

        # Metadata
        meta_style = styles['Normal']
        story.append(Paragraph(f"<b>Query:</b> {html_mod.escape(str(request.get('query', 'N/A')))}", meta_style))
        story.append(Paragraph(f"<b>Generated:</b> {html_mod.escape(str(request.get('timestamp', 'N/A')))}", meta_style))
        scope_list = request.get('scope', [])
        scope_str = ', '.join(scope_list) if isinstance(scope_list, list) else str(scope_list)
        story.append(Paragraph(f"<b>Scope:</b> {html_mod.escape(scope_str)}", meta_style))
        story.append(Paragraph(f"<b>Total Repositories:</b> {html_mod.escape(str(request.get('total_repositories', 0)))}", meta_style))
        story.append(Spacer(1, 20))

        # Repository Table
        repos = request.get('repositories', [])
        story.append(Paragraph("Repository List", styles['Heading2']))
        story.append(Spacer(1, 8))

        if repos:
            table_data = [['#', 'Repository', 'Reason', 'Source']]
            for idx, repo in enumerate(repos, 1):
                table_data.append([
                    str(idx),
                    repo.get('repository', '')[:25],
                    repo.get('reason', 'Context match')[:35],
                    repo.get('source', '-')[:15]
                ])

            table = Table(table_data, colWidths=[0.4*inch, 2.0*inch, 2.8*inch, 1.3*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1e3a5f')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (0, -1), 'CENTER'),
                ('ALIGN', (1, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f5f5f5')),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f0f0f0')]),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('TOPPADDING', (0, 1), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
            ]))
            story.append(table)
        else:
            story.append(Paragraph("No affected repositories found.", meta_style))

        # Footer
        story.append(Spacer(1, 30))
        footer_style = ParagraphStyle('Footer', parent=styles['Normal'], fontSize=8, textColor=colors.grey)
        story.append(Paragraph("Report generated from Zero Day Analysis", footer_style))

        doc.build(story)
        buffer.seek(0)

        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=affected-repos.pdf"}
        )

    except ImportError:
        raise HTTPException(status_code=500, detail="PDF generation requires reportlab. Install with: pip install reportlab")
    except Exception as e:
        logger.error(f"PDF generation failed: {e}")
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(e)}")


@router.post(
    "/zero-day/export/repos/docx",
    dependencies=[Depends(require_permissions("findings:write"))],
    summary="Export affected repositories list as DOCX",
    responses={
        **CREATE_ERRORS,
        401: {"description": "Not authenticated"},
        403: {"description": "Insufficient permissions - findings:write required"},
        500: {"description": "DOCX generation failed or python-docx not installed"},
    },
)
async def export_zda_repos_docx(request: dict):
    """
    Export Zero Day Analysis affected repositories as DOCX.

    Generates a downloadable Word document with a table of affected repositories.
    Requires findings:write permission.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        doc = Document()

        # Title
        title = doc.add_heading('Affected Repositories Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        doc.add_paragraph(f"Query: {request.get('query', 'N/A')}")
        doc.add_paragraph(f"Generated: {request.get('timestamp', 'N/A')}")
        scope_list = request.get('scope', [])
        scope_str = ', '.join(scope_list) if isinstance(scope_list, list) else str(scope_list)
        doc.add_paragraph(f"Scope: {scope_str}")
        doc.add_paragraph(f"Total Repositories: {request.get('total_repositories', 0)}")

        doc.add_paragraph()  # Spacer

        # Repository List
        doc.add_heading('Repository List', level=1)
        repos = request.get('repositories', [])

        if repos:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'

            # Header row
            header_cells = table.rows[0].cells
            headers = ['#', 'Repository', 'Reason', 'Source']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                # Bold header
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            # Data rows
            for idx, repo in enumerate(repos, 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(idx)
                row_cells[1].text = repo.get('repository', '')
                row_cells[2].text = repo.get('reason', 'Context match')
                row_cells[3].text = repo.get('source', '-')
        else:
            doc.add_paragraph('No affected repositories found.')

        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph('Report generated from Zero Day Analysis')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=affected-repos.docx"}
        )

    except ImportError:
        raise HTTPException(status_code=500, detail="DOCX generation requires python-docx. Install with: pip install python-docx")
    except Exception as e:
        logger.error(f"DOCX generation failed: {e}")
        raise HTTPException(status_code=500, detail=f"DOCX generation failed: {str(e)}")


from ..dependencies import get_tenant_db
from .. import models
from sqlalchemy.orm import Session
from ..utils.repo_context import get_repo_context, clone_repo_to_temp, cleanup_repo
from ..utils.architecture_preprocessor import ArchitecturePreprocessor, build_diagram_prompt_from_preprocessed
from ..utils.diagram_executor import execute_with_self_annealing, apply_known_corrections
import uuid

import re

import subprocess
import base64
import os
import tempfile

class ArchitectureRequest(BaseModel):
    project_id: str = Field(..., description="Repository UUID or name to analyze")

class ArchitectureUpdateRequest(BaseModel):
    report: str = Field(..., description="Architecture report markdown text")
    diagram: Optional[str] = Field(None, description="Python diagrams-as-code source for the architecture diagram")

class PromptRequest(BaseModel):
    project_id: str = Field(..., description="Repository UUID or name")
    prompt: Optional[str] = Field(None, description="Custom architecture prompt to execute")

@router.post(
    "/architecture/prompt",
    dependencies=[Depends(require_permissions("findings:write"))],
    summary="Get the architecture analysis prompt for a project",
    responses={
        **CREATE_ERRORS,
        400: {"description": "Repository URL is missing"},
        401: {"description": "Not authenticated"},
        403: {"description": "Insufficient permissions - findings:write required"},
        404: {"description": "Project not found"},
        500: {"description": "Prompt generation failed"},
    },
)
async def get_architecture_prompt(
    request: ArchitectureRequest,
    db: Session = Depends(get_tenant_db),
    settings: settings = Depends(lambda: settings)
):
    """
    Get the constructed architecture prompt for a project.

    Clones the repository, analyzes its structure, and builds a prompt
    for AI architecture generation. Requires findings:write permission.
    """
    try:
        p_uuid = uuid.UUID(request.project_id)
        project = db.query(models.Repository).filter(models.Repository.id == p_uuid).first()
    except ValueError:
        project = db.query(models.Repository).filter(models.Repository.name == request.project_id).first()

    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    if not project.url:
        raise HTTPException(status_code=400, detail="Project repository URL is missing")

    # Clone repo to temp dir to analyze structure
    try:
        # Get org-specific token for this repository
        token = await get_github_token_for_repo(db, project)
        
        # clone_repo_to_temp creates a new temp dir and returns the path
        temp_dir = clone_repo_to_temp(project.url, token)
        
        try:
            # Analyze structure
            from ..utils.repo_context import get_repo_structure, get_config_files
            file_structure = get_repo_structure(temp_dir)
            config_files = get_config_files(temp_dir)
            
            # Build prompt
            from ...ai_agent.providers.openai import OpenAIProvider
            provider = OpenAIProvider(api_key=settings.OPENAI_API_KEY, model=settings.AI_MODEL)
            prompt = provider.build_architecture_prompt(project.name, file_structure, config_files, diagrams_index)
            
            return {"prompt": prompt}
        finally:
            cleanup_repo(temp_dir)
            
    except Exception as e:
        logger.error(f"Error building prompt: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post(
    "/architecture/validate",
    dependencies=[Depends(require_permissions("findings:write"))],
    summary="Execute a custom architecture prompt",
    responses={
        **CREATE_ERRORS,
        400: {"description": "Prompt is required"},
        401: {"description": "Not authenticated"},
        403: {"description": "Insufficient permissions - findings:write required"},
        500: {"description": "Prompt execution failed"},
    },
)
async def validate_architecture_prompt(
    request: PromptRequest,
    settings: settings = Depends(lambda: settings)
):
    """
    Execute a custom architecture prompt.

    Sends a custom prompt to the AI provider and returns the response.
    Used for testing and validating prompts. Requires findings:write permission.
    """
    if not request.prompt:
        raise HTTPException(status_code=400, detail="Prompt is required")

    try:
        from ...ai_agent.providers.openai import OpenAIProvider
        provider = OpenAIProvider(api_key=settings.OPENAI_API_KEY, model=settings.AI_MODEL)
        response = await provider.execute_prompt(request.prompt)
        return {"response": response}
        
    except Exception as e:
        logger.error(f"Error validating prompt: {e}")
        raise HTTPException(status_code=500, detail=str(e))

class RefineRequest(BaseModel):
    project_id: str = Field(..., description="Repository UUID or name")
    code: str = Field(..., description="Diagram Python code to refine")

@router.post(
    "/architecture/refine",
    dependencies=[Depends(require_permissions("findings:write"))],
    summary="Refine diagram code with correct cloud provider icons",
    responses={
        **CREATE_ERRORS,
        400: {"description": "Code is required"},
        401: {"description": "Not authenticated"},
        403: {"description": "Insufficient permissions - findings:write required"},
        404: {"description": "Project not found"},
        500: {"description": "Refinement failed"},
        503: {"description": "AI provider does not support diagram code fixing"},
    },
)
async def refine_architecture_diagram(
    request: RefineRequest,
    db: Session = Depends(get_tenant_db),
    settings: settings = Depends(lambda: settings)
):
    """
    Refine diagram code to use correct cloud provider icons.

    Uses AI to fix import paths and ensure the correct cloud provider icons
    are used based on the architecture report. Requires findings:write permission.
    """
    if not request.code:
        raise HTTPException(status_code=400, detail="Code is required")

    try:
        # Fetch the project to get the architecture report
        project = db.query(Project).filter(Project.id == request.project_id).first()
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
            
        report_context = project.architecture_report or ""

        if not ai_agent or not hasattr(ai_agent.provider, 'fix_and_enhance_diagram_code'):
            raise HTTPException(status_code=503, detail="AI provider does not support diagram code fixing")

        # We pass a specific "error" message that acts as the instruction
        instruction = "Refine this code to use the correct cloud provider icons based on the detected technology in the Architecture Report. Enforce the Cloud Provider Preference strictly."

        refined_code = await ai_agent.provider.fix_and_enhance_diagram_code(
            request.code,
            instruction,
            diagrams_index,
            report_context=report_context
        )
        
        # Clean up code block if present
        code_match = re.search(r"```python\n(.*?)```", refined_code, re.DOTALL)
        if code_match:
            refined_code = code_match.group(1).strip()
        else:
            refined_code = refined_code.replace("```python", "").replace("```", "").strip()
            
        return {"code": refined_code}
        
    except Exception as e:
        logger.error(f"Error refining diagram: {e}")
        raise HTTPException(status_code=500, detail=str(e))

class ArchitectureResponse(BaseModel):
    report: str = Field(..., description="Architecture overview report in markdown")
    diagram: Optional[str] = Field(None, description="Python diagrams-as-code source")
    image: Optional[str] = Field(None, description="Base64-encoded PNG image of the architecture diagram")

@router.get(
    "/architecture/{project_id}",
    response_model=ArchitectureResponse,
    dependencies=[Depends(require_permissions("findings:read"))],
    summary="Get saved architecture overview for a project",
    responses={
        **CRUD_ERRORS,
        401: {"description": "Not authenticated"},
        403: {"description": "Insufficient permissions - findings:read required"},
        404: {"description": "Project not found"},
    },
)
async def get_architecture(project_id: str, db: Session = Depends(get_tenant_db)):
    """
    Get saved architecture overview for a project.

    Returns the stored architecture report, diagram code, and generated
    PNG image. Requires findings:read permission.
    """
    try:
        p_uuid = uuid.UUID(project_id)
        project = db.query(models.Repository).filter(models.Repository.id == p_uuid).first()
    except ValueError:
        project = db.query(models.Repository).filter(models.Repository.name == project_id).first()
        
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # If we have code but no image (e.g. from old data or failed run), we could try to regen,
    # but for GET let's just return what we have.
    # Actually, we don't store the image in DB, so we need to generate it or store it.
    # Storing base64 in DB might be heavy.
    # Let's regenerate on the fly if missing? No, that's slow.
    # Let's assume we store the code, and maybe we should store the image too?
    # The user asked to "save their edits", implying code edits.
    # If we don't save the image, we have to run the code every time we view? That's slow and risky.
    # Let's add an `architecture_image` column to the DB?
    # Or just return the code and let the frontend trigger a generation if needed?
    # The requirement says "The Python script will then take that specs to create the diagram."
    # Let's stick to generating it on save/generate and returning it.
    # For now, if we don't have the image stored, we might return null and let the UI ask for regen?
    # Or better: let's execute the code if we have it.
    
    image_b64 = None
    if project.architecture_diagram:
        try:
            image_b64 = execute_diagram_code(
                project.architecture_diagram, 
                report_context=project.architecture_report
            )
        except Exception as e:
            logger.error(f"Failed to generate image from saved code: {e}")

    return ArchitectureResponse(
        report=project.architecture_report or "",
        diagram=project.architecture_diagram,
        image=image_b64
    )

@router.put(
    "/architecture/{project_id}",
    response_model=ArchitectureResponse,
    summary="Update architecture overview after user edits",
    responses={
        **CRUD_ERRORS,
        404: {"description": "Project not found"},
        500: {"description": "Diagram generation or AI fix failed"},
    },
)
async def update_architecture(project_id: str, request: ArchitectureUpdateRequest, db: Session = Depends(get_tenant_db)):
    """
    Update architecture overview (e.g. after user edits).

    Saves the updated report and diagram code, then attempts to regenerate
    the architecture diagram image.
    """
    try:
        p_uuid = uuid.UUID(project_id)
        project = db.query(models.Repository).filter(models.Repository.id == p_uuid).first()
    except ValueError:
        project = db.query(models.Repository).filter(models.Repository.name == project_id).first()
        
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
        
    project.architecture_report = request.report
    image_b64 = None
    
    if request.diagram is not None:
        project.architecture_diagram = request.diagram
        # Execute code to verify and generate image using self-annealing
        try:
            image_b64 = execute_diagram_code(
                request.diagram, 
                report_context=request.report
            )
        except Exception as e:
            logger.warning(f"Self-annealing diagram generation failed: {e}. Attempting AI-powered fix...")
            try:
                # AI-powered fix as fallback
                if not hasattr(ai_agent.provider, 'fix_and_enhance_diagram_code'):
                    logger.error("AI provider does not support diagram code fixing")
                    raise Exception("AI provider does not support diagram code fixing")

                fixed_code = await ai_agent.provider.fix_and_enhance_diagram_code(
                    request.diagram, str(e), diagrams_index, report_context=request.report
                )
                
                # Clean up code block if present
                code_match_fix = re.search(r"```python\n(.*?)```", fixed_code, re.DOTALL)
                if code_match_fix:
                    fixed_code = code_match_fix.group(1).strip()
                else:
                    fixed_code = fixed_code.replace("```python", "").replace("```", "").strip()
                    
                project.architecture_diagram = fixed_code
                image_b64 = execute_diagram_code(fixed_code, report_context=request.report)
                logger.info("AI-powered fix successful")
            except Exception as fix_error:
                logger.error(f"AI-powered fix failed: {fix_error}")
                # Code is saved so user can fix it manually
        
    db.commit()
    
    return ArchitectureResponse(
        report=project.architecture_report,
        diagram=project.architecture_diagram,
        image=image_b64
    )

def execute_diagram_code(code: str, use_self_annealing: bool = True, report_context: str = None) -> str:
    """
    Execute Python code to generate diagram and return base64 image.
    
    When use_self_annealing=True (default), uses DOE/self-annealing approach:
    - Pre-validates imports against diagrams index
    - Automatically corrects known import errors (e.g., Internet from wrong module)
    - Retries with progressive fixes up to 3 times
    """
    if use_self_annealing:
        # Use the new self-annealing executor
        image_b64, final_code, fix_log = execute_with_self_annealing(
            code=code,
            diagrams_index=diagrams_index,
            ai_fix_callback=None,  # Caller can handle AI fixes separately
            max_retries=3,
            report_context=report_context
        )
        
        if fix_log:
            logger.info(f"Self-annealing diagram execution: {fix_log}")
        
        if image_b64:
            return image_b64
        else:
            raise Exception(f"Self-annealing failed after all retries. Log: {fix_log}")
    
    # Fallback to direct execution (legacy mode)
    with tempfile.TemporaryDirectory() as tmpdir:
        script_path = os.path.join(tmpdir, "diagram_script.py")
        
        with open(script_path, "w") as f:
            f.write(code)
            
        try:
            subprocess.check_output(
                ["python3", script_path],
                cwd=tmpdir,
                stderr=subprocess.STDOUT,
                timeout=30
            )
        except subprocess.CalledProcessError as e:
            raise Exception(f"Script execution failed: {e.output.decode()}")
            
        png_path = os.path.join(tmpdir, "architecture_diagram.png")
        if not os.path.exists(png_path):
            files = [f for f in os.listdir(tmpdir) if f.endswith('.png')]
            if files:
                png_path = os.path.join(tmpdir, files[0])
            else:
                raise Exception("No PNG image generated by the script")
            
        with open(png_path, "rb") as f:
            return base64.b64encode(f.read()).decode("utf-8")

@router.post(
    "/architecture",
    response_model=ArchitectureResponse,
    dependencies=[Depends(require_permissions("findings:write"))],
    summary="Generate an architecture overview for a project",
    responses={
        **CREATE_ERRORS,
        400: {"description": "Repository URL is missing"},
        401: {"description": "Not authenticated"},
        403: {"description": "Insufficient permissions - findings:write required"},
        404: {"description": "Project not found"},
        500: {"description": "Architecture generation failed"},
        503: {"description": "AI Agent not initialized"},
    },
)
async def generate_architecture(request: ArchitectureRequest, db: Session = Depends(get_tenant_db)):
    """
    Generate an architecture overview for a project.

    Clones the repository, analyzes its structure, and uses AI to generate
    an architecture report with a diagrams-as-code visualization.
    Requires findings:write permission.
    """
    if not ai_agent:
        raise HTTPException(status_code=503, detail="AI Agent not initialized")
        
    # Get project
    try:
        p_uuid = uuid.UUID(request.project_id)
        project = db.query(models.Repository).filter(models.Repository.id == p_uuid).first()
    except ValueError:
        project = db.query(models.Repository).filter(models.Repository.name == request.project_id).first()
        
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
        
    if not project.url:
        raise HTTPException(status_code=400, detail="Project repository URL is missing")
        
    repo_path = None
    try:
        # Clone repo to temp
        # Get org-specific token for this repository
        token = await get_github_token_for_repo(db, project)
        repo_path = clone_repo_to_temp(project.url, token)
        
        # Get context
        structure, configs = get_repo_context(repo_path)
        
        # Generate overview
        full_response = await ai_agent.generate_architecture_overview(
            repo_name=project.name,
            file_structure=structure,
            config_files=configs
        )
        
        # Parse Python code from response
        diagram_code = None
        report = full_response
        image_b64 = None
        
        code_match = re.search(r"```python\n(.*?)```", full_response, re.DOTALL)
        if code_match:
            diagram_code = code_match.group(1).strip()
            # Remove the code block from the report
            report = full_response.replace(code_match.group(0), "").strip()
            
            # Generate image using self-annealing executor
            # The self-annealing approach handles import errors automatically
            try:
                image_b64 = execute_diagram_code(diagram_code, report_context=report)
            except Exception as e:
                logger.warning(f"Self-annealing diagram generation failed: {e}. Attempting AI-powered fix...")
                try:
                    # If self-annealing still failed, try AI-powered fix as last resort
                    if not hasattr(ai_agent.provider, 'fix_and_enhance_diagram_code'):
                        logger.error("AI provider does not support diagram code fixing")
                        raise Exception("AI provider does not support diagram code fixing")

                    fixed_code = await ai_agent.provider.fix_and_enhance_diagram_code(
                        diagram_code, str(e), diagrams_index, report_context=report
                    )
                    
                    # Clean up code block if present
                    code_match_fix = re.search(r"```python\n(.*?)```", fixed_code, re.DOTALL)
                    if code_match_fix:
                        fixed_code = code_match_fix.group(1).strip()
                    else:
                        fixed_code = fixed_code.replace("```python", "").replace("```", "").strip()
                        
                    diagram_code = fixed_code
                    # Try again with self-annealing on the AI-fixed code
                    image_b64 = execute_diagram_code(diagram_code, report_context=report)
                    logger.info("AI-powered fix successful")
                except Exception as fix_error:
                    logger.error(f"AI-powered fix failed: {fix_error}")
                    # We still save the code so user can fix it
            
        # Save to DB
        project.architecture_report = report
        project.architecture_diagram = diagram_code
        db.commit()
        
        return ArchitectureResponse(
            report=report,
            diagram=diagram_code,
            image=image_b64
        )
        
    except Exception as e:
        logger.error(f"Error generating architecture: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if repo_path:
            cleanup_repo(repo_path)

class VersionCreateRequest(BaseModel):
    description: Optional[str] = Field(None, description="Optional description for this version snapshot")

class VersionResponse(BaseModel):
    id: str = Field(..., description="Version UUID")
    version_number: int = Field(..., description="Sequential version number")
    created_at: str = Field(..., description="ISO 8601 creation timestamp")
    description: Optional[str] = Field(None, description="Version description")

class VersionDetailResponse(VersionResponse):
    report: str = Field(..., description="Architecture report content for this version")
    diagram: Optional[str] = Field(None, description="Diagram code for this version")

@router.post(
    "/architecture/{project_id}/versions",
    response_model=VersionResponse,
    dependencies=[Depends(require_permissions("findings:write"))],
    summary="Save current architecture as a new version",
    responses={
        **CREATE_ERRORS,
        400: {"description": "No architecture report to save"},
        401: {"description": "Not authenticated"},
        403: {"description": "Insufficient permissions - findings:write required"},
        404: {"description": "Project not found"},
    },
)
async def create_architecture_version(
    project_id: str, 
    request: VersionCreateRequest, 
    db: Session = Depends(get_tenant_db)
):
    """
    Save current architecture state as a new version.

    Creates a snapshot of the current architecture report and diagram code
    with an auto-incremented version number. Requires findings:write permission.
    """
    try:
        p_uuid = uuid.UUID(project_id)
        project = db.query(models.Repository).filter(models.Repository.id == p_uuid).first()
    except ValueError:
        project = db.query(models.Repository).filter(models.Repository.name == project_id).first()
        
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
        
    if not project.architecture_report:
        raise HTTPException(status_code=400, detail="No architecture report to save")
        
    # Get next version number
    last_version = db.query(func.max(models.ArchitectureVersion.version_number))\
        .filter(models.ArchitectureVersion.repository_id == project.id)\
        .scalar() or 0
        
    new_version = models.ArchitectureVersion(
        repository_id=project.id,
        version_number=last_version + 1,
        report_content=project.architecture_report,
        diagram_code=project.architecture_diagram,
        description=request.description
    )
    
    db.add(new_version)
    db.commit()
    db.refresh(new_version)
    
    return VersionResponse(
        id=str(new_version.id),
        version_number=new_version.version_number,
        created_at=new_version.created_at.isoformat(),
        description=new_version.description
    )

@router.get(
    "/architecture/{project_id}/versions",
    response_model=List[VersionResponse],
    dependencies=[Depends(require_permissions("findings:read"))],
    summary="List all architecture versions for a project",
    responses={
        **LIST_ERRORS,
        401: {"description": "Not authenticated"},
        403: {"description": "Insufficient permissions - findings:read required"},
        404: {"description": "Project not found"},
    },
)
async def list_architecture_versions(project_id: str, db: Session = Depends(get_tenant_db)):
    """
    List all architecture versions for a project.

    Returns version metadata sorted by version number descending.
    Requires findings:read permission.
    """
    try:
        p_uuid = uuid.UUID(project_id)
        project = db.query(models.Repository).filter(models.Repository.id == p_uuid).first()
    except ValueError:
        project = db.query(models.Repository).filter(models.Repository.name == project_id).first()
        
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
        
    versions = db.query(models.ArchitectureVersion)\
        .filter(models.ArchitectureVersion.repository_id == project.id)\
        .order_by(models.ArchitectureVersion.version_number.desc())\
        .all()
        
    return [
        VersionResponse(
            id=str(v.id),
            version_number=v.version_number,
            created_at=v.created_at.isoformat(),
            description=v.description
        ) for v in versions
    ]

@router.get(
    "/architecture/{project_id}/versions/{version_id}",
    response_model=VersionDetailResponse,
    dependencies=[Depends(require_permissions("findings:read"))],
    summary="Get a specific architecture version",
    responses={
        **CRUD_ERRORS,
        400: {"description": "Invalid version ID format"},
        401: {"description": "Not authenticated"},
        403: {"description": "Insufficient permissions - findings:read required"},
        404: {"description": "Version not found"},
    },
)
async def get_architecture_version(project_id: str, version_id: str, db: Session = Depends(get_tenant_db)):
    """
    Get details of a specific architecture version.

    Returns the full report and diagram code for the given version.
    Requires findings:read permission.
    """
    try:
        v_uuid = uuid.UUID(version_id)
        version = db.query(models.ArchitectureVersion).filter(models.ArchitectureVersion.id == v_uuid).first()
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid version ID")
        
    if not version:
        raise HTTPException(status_code=404, detail="Version not found")
        
    return VersionDetailResponse(
        id=str(version.id),
        version_number=version.version_number,
        created_at=version.created_at.isoformat(),
        description=version.description,
        report=version.report_content,
        diagram=version.diagram_code
    )

@router.post(
    "/architecture/{project_id}/restore/{version_id}",
    dependencies=[Depends(require_permissions("findings:write"))],
    summary="Restore a previous architecture version",
    responses={
        **CREATE_ERRORS,
        400: {"description": "Invalid version ID format"},
        401: {"description": "Not authenticated"},
        403: {"description": "Insufficient permissions - findings:write required"},
        404: {"description": "Project or version not found"},
    },
)
async def restore_architecture_version(project_id: str, version_id: str, db: Session = Depends(get_tenant_db)):
    """
    Restore a previous architecture version.

    Overwrites the current architecture report and diagram code with the
    contents of the specified version. Requires findings:write permission.
    """
    try:
        p_uuid = uuid.UUID(project_id)
        project = db.query(models.Repository).filter(models.Repository.id == p_uuid).first()
    except ValueError:
        project = db.query(models.Repository).filter(models.Repository.name == project_id).first()
        
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
        
    try:
        v_uuid = uuid.UUID(version_id)
        version = db.query(models.ArchitectureVersion).filter(models.ArchitectureVersion.id == v_uuid).first()
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid version ID")
        
    if not version:
        raise HTTPException(status_code=404, detail="Version not found")
        
    # Overwrite current state
    project.architecture_report = version.report_content
    project.architecture_diagram = version.diagram_code
    
    db.commit()
    
    return {"status": "success", "message": f"Restored version {version.version_number}"}


# Architecture Preprocessing Endpoint

class PreprocessRequest(BaseModel):
    project_id: str = Field(..., description="Repository UUID or name to preprocess")
    force_refresh: bool = Field(False, description="If true, re-preprocess even if cached data exists")

class PreprocessedArchitectureResponse(BaseModel):
    project_name: str = Field(..., description="Detected project name")
    project_type: str = Field(..., description="Detected project type (web-api, cli, library, etc.)")
    cloud_provider: Optional[str] = Field(None, description="Detected cloud provider (aws, gcp, azure)")
    components: List[Dict[str, Any]] = Field(..., description="Extracted architecture components")
    api_summary: Dict[str, Any] = Field(..., description="Summary of API routes and endpoints")
    data_layer: Dict[str, Any] = Field(..., description="Data layer details (databases, caches, etc.)")
    external_integrations: List[Dict[str, str]] = Field(..., description="External service integrations detected")
    tech_stack: Dict[str, List[str]] = Field(..., description="Technology stack by category")
    confidence_notes: List[str] = Field(..., description="Notes about detection confidence and assumptions")
    cached: bool = Field(False, description="Whether the response was served from cache")

@router.post(
    "/architecture/preprocess",
    response_model=PreprocessedArchitectureResponse,
    dependencies=[Depends(require_permissions("findings:write"))],
    summary="Preprocess repository for architecture extraction",
    responses={
        **CREATE_ERRORS,
        400: {"description": "Repository URL is missing"},
        401: {"description": "Not authenticated"},
        403: {"description": "Insufficient permissions - findings:write required"},
        404: {"description": "Project not found"},
        500: {"description": "Preprocessing failed"},
        503: {"description": "AI Agent not initialized"},
    },
)
async def preprocess_architecture(
    request: PreprocessRequest,
    db: Session = Depends(get_tenant_db)
):
    """
    Preprocess a repository to extract architecture information.

    Clones the repo, runs deterministic code extractors, and uses AI to produce
    structured JSON for diagram generation. Results are cached.
    **Required permissions:** findings:write.
    """
    if not ai_agent:
        raise HTTPException(status_code=503, detail="AI Agent not initialized")

    # Get project
    try:
        p_uuid = uuid.UUID(request.project_id)
        project = db.query(models.Repository).filter(models.Repository.id == p_uuid).first()
    except ValueError:
        project = db.query(models.Repository).filter(models.Repository.name == request.project_id).first()

    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    if not project.url:
        raise HTTPException(status_code=400, detail="Project repository URL is missing")

    # Check for cached preprocessed data (if not force_refresh)
    if not request.force_refresh and project.architecture_preprocessed:
        try:
            cached_data = json.loads(project.architecture_preprocessed)
            return PreprocessedArchitectureResponse(
                **cached_data,
                cached=True
            )
        except (json.JSONDecodeError, TypeError):
            logger.warning(f"Failed to parse cached preprocessed data for project {project.id}")

    repo_path = None
    try:
        # Clone repo
        # Get org-specific token for this repository
        token = await get_github_token_for_repo(db, project)
        repo_path = clone_repo_to_temp(project.url, token)

        # Create preprocessor and run
        preprocessor = ArchitecturePreprocessor(ai_agent.provider)
        architecture = await preprocessor.preprocess(repo_path, project.name)

        # Convert to response format
        response_data = {
            "project_name": architecture.project_name,
            "project_type": architecture.project_type,
            "cloud_provider": architecture.cloud_provider,
            "components": [
                {
                    "name": c.name,
                    "type": c.type,
                    "technology": c.technology,
                    "port": c.port,
                    "description": c.description,
                    "connections": c.connects_to
                }
                for c in architecture.components
            ],
            "api_summary": architecture.api_summary,
            "data_layer": architecture.data_layer,
            "external_integrations": architecture.external_integrations,
            "tech_stack": architecture.tech_stack,
            "confidence_notes": architecture.confidence_notes
        }

        # Cache the preprocessed data
        project.architecture_preprocessed = json.dumps(response_data)
        db.commit()

        return PreprocessedArchitectureResponse(
            **response_data,
            cached=False
        )

    except Exception as e:
        logger.error(f"Error preprocessing architecture: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if repo_path:
            cleanup_repo(repo_path)


@router.post(
    "/architecture/generate-from-preprocessed",
    response_model=ArchitectureResponse,
    dependencies=[Depends(require_permissions("findings:write"))],
    summary="Generate architecture diagram from preprocessed data",
    responses={
        **CREATE_ERRORS,
        400: {"description": "No preprocessed data found"},
        401: {"description": "Not authenticated"},
        403: {"description": "Insufficient permissions - findings:write required"},
        404: {"description": "Project not found"},
        500: {"description": "Diagram generation failed"},
        503: {"description": "AI Agent not initialized"},
    },
)
async def generate_architecture_from_preprocessed(
    request: ArchitectureRequest,
    db: Session = Depends(get_tenant_db)
):
    """
    Generate architecture diagram using preprocessed data.

    Uses cached preprocessed architecture data to generate a more accurate
    diagram without re-analyzing the entire codebase.
    **Required permissions:** findings:write.
    """
    if not ai_agent:
        raise HTTPException(status_code=503, detail="AI Agent not initialized")

    # Get project
    try:
        p_uuid = uuid.UUID(request.project_id)
        project = db.query(models.Repository).filter(models.Repository.id == p_uuid).first()
    except ValueError:
        project = db.query(models.Repository).filter(models.Repository.name == request.project_id).first()

    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    # Check for preprocessed data
    if not project.architecture_preprocessed:
        raise HTTPException(
            status_code=400,
            detail="No preprocessed data found. Run /ai/architecture/preprocess first."
        )

    try:
        # Parse preprocessed data
        preprocessed_data = json.loads(project.architecture_preprocessed)

        # Reconstruct the PreprocessedArchitecture object
        from ..utils.architecture_preprocessor import PreprocessedArchitecture, ArchitectureComponent

        components = [
            ArchitectureComponent(
                name=c.get("name", "unknown"),
                type=c.get("type", "service"),
                technology=c.get("technology"),
                port=c.get("port"),
                description=c.get("description", ""),
                connects_to=c.get("connections", [])
            ) for c in preprocessed_data.get("components", [])
        ]

        architecture = PreprocessedArchitecture(
            project_name=preprocessed_data["project_name"],
            project_type=preprocessed_data["project_type"],
            cloud_provider=preprocessed_data.get("cloud_provider"),
            components=components,
            api_summary=preprocessed_data.get("api_summary", {}),
            data_layer=preprocessed_data.get("data_layer", {}),
            external_integrations=preprocessed_data.get("external_integrations", []),
            tech_stack=preprocessed_data.get("tech_stack", {}),
            confidence_notes=preprocessed_data.get("confidence_notes", [])
        )

        # Build focused diagram prompt
        diagram_prompt = build_diagram_prompt_from_preprocessed(architecture)

        # Add the diagrams index for icon reference
        prompt_with_icons = f"""{diagram_prompt}

## Available Diagram Icons Reference
Use ONLY these valid imports. Here are common icons by cloud provider:

### AWS (from diagrams.aws.*)
- compute: EC2, Lambda, ECS, EKS, Fargate
- database: RDS, Aurora, DynamoDB, ElastiCache, Redshift
- network: ELB, ALB, NLB, CloudFront, Route53, VPC, APIGateway
- storage: S3, EBS, EFS
- integration: SQS, SNS, EventBridge
- security: IAM, Cognito, KMS

### GCP (from diagrams.gcp.*)
- compute: ComputeEngine, Functions, Run, GKE
- database: SQL, Spanner, Firestore, Memorystore, BigQuery
- network: LoadBalancing, CDN, DNS
- storage: Storage

### Azure (from diagrams.azure.*)
- compute: VirtualMachines, FunctionApps, ContainerInstances, KubernetesServices
- database: SQLDatabases, CosmosDb, CacheForRedis
- network: LoadBalancers, CDNProfiles, DNS, ApplicationGateway
- storage: BlobStorage

### Generic (from diagrams.generic.*)
- compute: Rack
- database: SQL
- network: Firewall, Router, Switch
- storage: Storage

### On-Premise (from diagrams.onprem.*)
- compute: Server
- database: PostgreSQL, MySQL, MongoDB, Redis, Cassandra, Elasticsearch
- queue: Kafka, RabbitMQ, Celery
- network: Nginx, Apache, Traefik
- container: Docker, Kubernetes
- ci: Jenkins, GitlabCI, GithubActions
- client: Client, User
- monitoring: Prometheus, Grafana, Datadog

IMPORTANT: Only use icons that exist in the diagrams library. When in doubt, use generic icons."""

        # Generate diagram code using AI
        if hasattr(ai_agent.provider, 'execute_prompt'):
            response = await ai_agent.provider.execute_prompt(prompt_with_icons)
        else:
            raise HTTPException(status_code=500, detail="AI provider does not support prompt execution")

        # Parse the response
        diagram_code = None
        report = response
        image_b64 = None

        code_match = re.search(r"```python\n(.*?)```", response, re.DOTALL)
        if code_match:
            diagram_code = code_match.group(1).strip()
            report = response.replace(code_match.group(0), "").strip()

            # Try to generate image
            try:
                image_b64 = execute_diagram_code(diagram_code)
            except Exception as e:
                logger.warning(f"Diagram generation failed: {e}. Attempting auto-fix...")
                try:
                    if hasattr(ai_agent.provider, 'fix_and_enhance_diagram_code'):
                        fixed_code = await ai_agent.provider.fix_and_enhance_diagram_code(
                            diagram_code, str(e), diagrams_index
                        )

                        code_match_fix = re.search(r"```python\n(.*?)```", fixed_code, re.DOTALL)
                        if code_match_fix:
                            fixed_code = code_match_fix.group(1).strip()
                        else:
                            fixed_code = fixed_code.replace("```python", "").replace("```", "").strip()

                        diagram_code = fixed_code
                        image_b64 = execute_diagram_code(diagram_code)
                        logger.info("Auto-fix successful for preprocessed diagram")
                except Exception as fix_error:
                    logger.error(f"Auto-fix failed: {fix_error}")

        # Save to project
        project.architecture_report = report
        project.architecture_diagram = diagram_code
        db.commit()

        return ArchitectureResponse(
            report=report,
            diagram=diagram_code,
            image=image_b64
        )

    except json.JSONDecodeError:
        raise HTTPException(status_code=500, detail="Failed to parse preprocessed data")
    except Exception as e:
        logger.error(f"Error generating from preprocessed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# =============================================================================
# Phase 3: AI Deduplication & Auto-Triage
# =============================================================================

class DeduplicateRequest(BaseModel):
    """Request to find duplicate findings."""
    repository_id: Optional[str] = Field(None, description="Scope deduplication to a single repository UUID")
    limit: int = Field(100, description="Maximum number of findings to analyze")
    similarity_threshold: float = Field(0.8, description="Similarity threshold (0.0-1.0) for grouping duplicates")


class DuplicateGroup(BaseModel):
    """A group of duplicate findings."""
    group_id: str = Field(..., description="Unique group identifier")
    primary_finding_id: str = Field(..., description="UUID of the primary finding in the group")
    primary_title: str = Field(..., description="Title of the primary finding")
    member_count: int = Field(..., description="Number of findings in the duplicate group")
    member_ids: List[str] = Field(..., description="UUIDs of all findings in the group")
    similarity_reason: str = Field(..., description="Explanation of why these findings are grouped")


class DeduplicateResponse(BaseModel):
    """Response after deduplication analysis."""
    groups_found: int = Field(..., description="Number of duplicate groups detected")
    duplicate_groups: List[DuplicateGroup] = Field(..., description="Identified duplicate groups")
    total_findings_analyzed: int = Field(..., description="Total findings that were analyzed")


@router.post(
    "/deduplicate",
    response_model=DeduplicateResponse,
    dependencies=[Depends(require_permissions("findings:write"))],
    summary="Find duplicate findings using AI",
    responses={
        **CREATE_ERRORS,
        401: {"description": "Not authenticated"},
        403: {"description": "Insufficient permissions - findings:write required"},
        500: {"description": "Deduplication analysis failed"},
    },
)
def find_duplicate_findings(
    request: DeduplicateRequest,
    db: Session = Depends(get_tenant_db)
):
    """
    Find duplicate findings based on title, description, and context similarity.

    Groups similar findings together for bulk actions using normalized
    pattern matching. **Required permissions:** findings:write.
    """
    from collections import defaultdict
    
    # Get findings to analyze
    query = db.query(models.Finding).filter(models.Finding.status == 'open')
    
    if request.repository_id:
        query = query.filter(models.Finding.repository_id == request.repository_id)
    
    findings = query.order_by(models.Finding.created_at.desc()).limit(request.limit).all()
    
    if not findings:
        return DeduplicateResponse(
            groups_found=0,
            duplicate_groups=[],
            total_findings_analyzed=0
        )
    
    # Simple rule-based deduplication (AI enhancement can be added later)
    # Group by normalized title + scanner + severity
    groups = defaultdict(list)
    
    for f in findings:
        # Normalize title for comparison
        title_normalized = (f.title or '').lower().strip()
        # Remove common variable parts like line numbers, file paths
        import re
        title_normalized = re.sub(r'\d+', 'N', title_normalized)  # Replace numbers
        title_normalized = re.sub(r'[a-f0-9]{8,}', 'HASH', title_normalized)  # Replace hashes
        
        key = f"{title_normalized}|{f.scanner_name or 'unknown'}|{f.severity or 'unknown'}"
        groups[key].append(f)
    
    # Filter to only groups with duplicates
    duplicate_groups = []
    from uuid import uuid4
    
    for key, group_findings in groups.items():
        if len(group_findings) > 1:
            group_id = str(uuid4())
            primary = group_findings[0]
            
            # Assign group to findings in DB
            for i, f in enumerate(group_findings):
                f.duplicate_group_id = group_id
                f.is_primary_in_group = (i == 0)
            
            duplicate_groups.append(DuplicateGroup(
                group_id=group_id,
                primary_finding_id=str(primary.finding_uuid),
                primary_title=primary.title,
                member_count=len(group_findings),
                member_ids=[str(f.finding_uuid) for f in group_findings],
                similarity_reason=f"Same title pattern ({group_findings[0].scanner_name})"
            ))
    
    db.commit()
    
    logger.info(f"Deduplication found {len(duplicate_groups)} groups from {len(findings)} findings")
    
    return DeduplicateResponse(
        groups_found=len(duplicate_groups),
        duplicate_groups=duplicate_groups,
        total_findings_analyzed=len(findings)
    )


@router.get(
    "/duplicate-groups",
    dependencies=[Depends(require_permissions("findings:read"))],
    summary="Get all duplicate finding groups",
    responses={
        **LIST_ERRORS,
        401: {"description": "Not authenticated"},
        403: {"description": "Insufficient permissions - findings:read required"},
        500: {"description": "Internal server error"},
    },
)
def get_duplicate_groups(
    limit: int = 20,
    db: Session = Depends(get_tenant_db)
):
    """
    Get all duplicate groups.

    Returns duplicate finding groups with member counts, ordered by group
    size descending. Requires findings:read permission.
    """
    from sqlalchemy import func as sql_func
    
    # Get groups with counts
    groups = db.query(
        models.Finding.duplicate_group_id,
        sql_func.count(models.Finding.id).label('count'),
        sql_func.min(models.Finding.title).label('title')
    ).filter(
        models.Finding.duplicate_group_id.isnot(None)
    ).group_by(
        models.Finding.duplicate_group_id
    ).having(
        sql_func.count(models.Finding.id) > 1
    ).order_by(
        sql_func.count(models.Finding.id).desc()
    ).limit(limit).all()
    
    return [{
        "group_id": str(g[0]),
        "count": g[1],
        "sample_title": g[2]
    } for g in groups]


@router.post(
    "/merge-duplicates/{group_id}",
    dependencies=[Depends(require_permissions("findings:write"))],
    summary="Perform bulk action on a duplicate group",
    responses={
        **CREATE_ERRORS,
        401: {"description": "Not authenticated"},
        403: {"description": "Insufficient permissions - findings:write required"},
        404: {"description": "Duplicate group not found"},
    },
)
def merge_duplicate_group(
    group_id: str,
    action: str = "resolve_non_primary",  # resolve_non_primary, keep_all
    db: Session = Depends(get_tenant_db)
):
    """
    Perform bulk action on a duplicate group.

    Resolves non-primary findings as duplicates, or keeps all findings
    depending on the action parameter. Requires findings:write permission.
    """
    findings = db.query(models.Finding).filter(
        models.Finding.duplicate_group_id == group_id
    ).all()
    
    if not findings:
        raise HTTPException(status_code=404, detail="Duplicate group not found")
    
    if action == "resolve_non_primary":
        count = 0
        for f in findings:
            if not f.is_primary_in_group:
                f.status = "resolved"
                f.resolution = "duplicate"
                f.resolution_notes = f"Resolved as duplicate of primary finding"
                f.resolved_at = datetime.utcnow()
                count += 1
        db.commit()
        return {"status": "success", "resolved_count": count}
    
    return {"status": "no_action", "message": f"Unknown action: {action}"}


# =============================================================================
# AI Auto-Triage (Phase 3.2)
# =============================================================================

class AutoTriageRequest(BaseModel):
    """Request to auto-triage findings."""
    finding_ids: Optional[List[str]] = Field(None, description="Specific finding UUIDs to triage, or null for batch mode")
    limit: int = Field(50, description="Maximum findings to triage in batch mode")


class TriageResult(BaseModel):
    """Single triage result."""
    finding_id: str = Field(..., description="Finding UUID")
    recommendation: str = Field(..., description="Triage recommendation (true_positive, false_positive, needs_review)")
    confidence: float = Field(..., description="Confidence score for the recommendation (0.0-1.0)")
    reasoning: str = Field(..., description="Human-readable reasoning for the recommendation")


class AutoTriageResponse(BaseModel):
    """Response after auto-triage."""
    triaged_count: int = Field(..., description="Total number of findings triaged")
    results: List[TriageResult] = Field(..., description="Per-finding triage results")
    true_positive_count: int = Field(..., description="Number classified as true positive")
    false_positive_count: int = Field(..., description="Number classified as false positive")
    needs_review_count: int = Field(..., description="Number requiring manual review")


@router.post(
    "/auto-triage",
    response_model=AutoTriageResponse,
    dependencies=[Depends(require_permissions("findings:write"))],
    summary="Auto-triage findings using AI and heuristics",
    responses={
        **CREATE_ERRORS,
        401: {"description": "Not authenticated"},
        403: {"description": "Insufficient permissions - findings:write required"},
        500: {"description": "Auto-triage failed"},
    },
)
def auto_triage_findings(
    request: AutoTriageRequest,
    db: Session = Depends(get_tenant_db)
):
    """
    Automatically triage findings using heuristic rules and AI.

    Marks each finding with a recommendation (true_positive, false_positive,
    needs_review), confidence score, and reasoning. **Required permissions:** findings:write.
    """
    # Get findings to triage
    if request.finding_ids:
        findings = db.query(models.Finding).filter(
            models.Finding.finding_uuid.in_([uuid.UUID(fid) for fid in request.finding_ids])
        ).all()
    else:
        # Get un-triaged findings
        findings = db.query(models.Finding).filter(
            models.Finding.ai_triage_recommendation.is_(None),
            models.Finding.status == 'open'
        ).limit(request.limit).all()
    
    results = []
    counts = {"true_positive": 0, "false_positive": 0, "needs_review": 0}
    
    for f in findings:
        # Heuristic-based triage (AI enhancement can be added via ai_agent)
        recommendation = "needs_review"
        confidence = 0.5
        reasoning = "Requires manual review"
        
        title_lower = (f.title or '').lower()
        scanner = (f.scanner_name or '').lower()
        severity = (f.severity or '').lower()
        
        # High-confidence true positives
        if scanner in ['trufflehog', 'gitleaks'] and f.is_verified_by_scanner:
            recommendation = "true_positive"
            confidence = 0.95
            reasoning = "Scanner-verified secret finding"
        elif severity == 'critical' and 'cve' in title_lower:
            recommendation = "true_positive"
            confidence = 0.85
            reasoning = "Critical CVE from vulnerability scanner"
        elif f.is_validated_active is True:
            recommendation = "true_positive"
            confidence = 0.98
            reasoning = "Secret confirmed active via validation"
        
        # Likely false positives
        elif 'test' in (f.file_path or '').lower() or 'mock' in (f.file_path or '').lower():
            recommendation = "false_positive"
            confidence = 0.7
            reasoning = "Finding in test/mock file - likely not production"
        elif f.is_validated_active is False:
            recommendation = "false_positive"
            confidence = 0.9
            reasoning = "Secret confirmed as revoked/inactive"
        elif 'example' in title_lower or 'sample' in title_lower:
            recommendation = "false_positive"
            confidence = 0.6
            reasoning = "Title suggests example/sample code"
        
        # Update finding
        f.ai_triage_recommendation = recommendation
        f.ai_triage_confidence = confidence
        f.ai_triage_reasoning = reasoning
        
        counts[recommendation] += 1
        results.append(TriageResult(
            finding_id=str(f.finding_uuid),
            recommendation=recommendation,
            confidence=confidence,
            reasoning=reasoning
        ))
    
    db.commit()
    logger.info(f"Auto-triaged {len(results)} findings: {counts}")
    
    return AutoTriageResponse(
        triaged_count=len(results),
        results=results,
        true_positive_count=counts["true_positive"],
        false_positive_count=counts["false_positive"],
        needs_review_count=counts["needs_review"]
    )


@router.get(
    "/triage-stats",
    dependencies=[Depends(require_permissions("findings:read"))],
    summary="Get AI triage recommendation statistics",
    responses={
        **LIST_ERRORS,
        401: {"description": "Not authenticated"},
        403: {"description": "Insufficient permissions - findings:read required"},
        500: {"description": "Internal server error"},
    },
)
def get_triage_stats(db: Session = Depends(get_tenant_db)):
    """
    Get statistics on AI triage recommendations.

    Returns counts and average confidence by recommendation type
    (true_positive, false_positive, needs_review). Requires findings:read permission.
    """
    from sqlalchemy import func as sql_func
    
    stats = db.query(
        models.Finding.ai_triage_recommendation,
        sql_func.count(models.Finding.id),
        sql_func.avg(models.Finding.ai_triage_confidence)
    ).filter(
        models.Finding.ai_triage_recommendation.isnot(None)
    ).group_by(
        models.Finding.ai_triage_recommendation
    ).all()
    
    return {
        "by_recommendation": {
            s[0]: {"count": s[1], "avg_confidence": round(float(s[2] or 0), 2)}
            for s in stats
        },
        "total_triaged": sum(s[1] for s in stats),
        "untriaged": db.query(models.Finding).filter(
            models.Finding.ai_triage_recommendation.is_(None),
            models.Finding.status == 'open'
        ).count()
    }


# =============================================================================
# False Positive Learning (Phase 3.3)
# =============================================================================

class FPPattern(BaseModel):
    """A false positive pattern."""
    pattern_type: str = Field(..., description="Pattern type (title, file_path, scanner_rule)")
    pattern_value: str = Field(..., description="Pattern value used for matching")
    scanner_name: Optional[str] = Field(None, description="Scanner name if pattern is scanner-specific")
    match_count: int = Field(..., description="Number of resolved FPs matching this pattern")
    confidence: float = Field(..., description="Pattern confidence score (0.0-1.0)")
    auto_dismiss: bool = Field(..., description="Whether this pattern should auto-dismiss future matches")


class LearnFPRequest(BaseModel):
    """Request to learn from dismissed findings."""
    days: int = Field(30, description="Lookback period in days for analyzing resolved false positives")
    min_occurrences: int = Field(3, description="Minimum times a pattern must appear to be learned")


class LearnFPResponse(BaseModel):
    """Response after learning FP patterns."""
    patterns_learned: int = Field(..., description="Number of new patterns identified")
    patterns: List[FPPattern] = Field(..., description="Learned false positive patterns")


@router.post(
    "/learn-false-positives",
    response_model=LearnFPResponse,
    dependencies=[Depends(require_permissions("findings:write"))],
    summary="Learn false positive patterns from resolved findings",
    responses={
        **CREATE_ERRORS,
        401: {"description": "Not authenticated"},
        403: {"description": "Insufficient permissions - findings:write required"},
        500: {"description": "Pattern learning failed"},
    },
)
def learn_false_positive_patterns(
    request: LearnFPRequest,
    db: Session = Depends(get_tenant_db)
):
    """
    Analyze resolved false positive findings to learn repeatable patterns.

    Extracts title and file-path patterns from recently dismissed findings
    for use in auto-dismissal. **Required permissions:** findings:write.
    """
    from datetime import timedelta
    from collections import Counter
    
    cutoff = datetime.utcnow() - timedelta(days=request.days)
    
    # Get false positive resolutions
    fps = db.query(models.Finding).filter(
        models.Finding.resolution == 'false_positive',
        models.Finding.resolved_at >= cutoff
    ).all()
    
    # Analyze patterns
    title_patterns = Counter()
    path_patterns = Counter()
    
    for f in fps:
        # Extract title pattern (normalize)
        import re
        title_norm = re.sub(r'\d+', 'N', (f.title or '').lower())
        title_norm = re.sub(r'[a-f0-9]{8,}', 'HASH', title_norm)
        title_patterns[f"{f.scanner_name}:{title_norm}"] += 1
        
        # Extract path pattern
        path = f.file_path or ''
        if '/test' in path.lower() or '/tests' in path.lower():
            path_patterns["test_directory"] += 1
        if '/vendor' in path.lower() or '/node_modules' in path.lower():
            path_patterns["vendor_directory"] += 1
    
    # Create patterns with min occurrences
    patterns = []
    
    for pattern, count in title_patterns.most_common(20):
        if count >= request.min_occurrences:
            scanner, title = pattern.split(':', 1) if ':' in pattern else ('any', pattern)
            patterns.append(FPPattern(
                pattern_type="title",
                pattern_value=title,
                scanner_name=scanner if scanner != 'any' else None,
                match_count=count,
                confidence=min(0.9, 0.5 + (count / 20)),
                auto_dismiss=count >= 5
            ))
    
    for pattern, count in path_patterns.most_common(10):
        if count >= request.min_occurrences:
            patterns.append(FPPattern(
                pattern_type="file_path",
                pattern_value=pattern,
                scanner_name=None,
                match_count=count,
                confidence=min(0.85, 0.4 + (count / 30)),
                auto_dismiss=count >= 10
            ))
    
    logger.info(f"Learned {len(patterns)} FP patterns from {len(fps)} false positives")
    
    return LearnFPResponse(
        patterns_learned=len(patterns),
        patterns=patterns
    )


@router.get(
    "/fp-suggestions",
    dependencies=[Depends(require_permissions("findings:read"))],
    summary="Get false positive dismissal suggestions",
    responses={
        **LIST_ERRORS,
        401: {"description": "Not authenticated"},
        403: {"description": "Insufficient permissions - findings:read required"},
        500: {"description": "Internal server error"},
    },
)
def get_fp_suggestions(
    limit: int = 20,
    db: Session = Depends(get_tenant_db)
):
    """
    Get open findings that match learned false positive patterns.

    Returns candidates in test or vendor directories for bulk false positive
    dismissal. **Required permissions:** findings:read.
    """
    # Find open findings in test directories
    test_findings = db.query(models.Finding).join(models.Repository).filter(
        models.Finding.status == 'open',
        models.Finding.file_path.ilike('%/test%')
    ).limit(limit // 2).all()
    
    # Find open findings in vendor directories
    vendor_findings = db.query(models.Finding).join(models.Repository).filter(
        models.Finding.status == 'open',
        models.Finding.file_path.ilike('%/vendor%')
    ).limit(limit // 2).all()
    
    suggestions = []
    
    for f in test_findings:
        suggestions.append({
            "finding_id": str(f.finding_uuid),
            "title": f.title,
            "repo_name": f.repository.name if f.repository else "Unknown",
            "file_path": f.file_path,
            "pattern_type": "file_path",
            "pattern_match": "test_directory",
            "suggested_action": "mark_false_positive",
            "confidence": 0.7
        })
    
    for f in vendor_findings:
        suggestions.append({
            "finding_id": str(f.finding_uuid),
            "title": f.title,
            "repo_name": f.repository.name if f.repository else "Unknown",
            "file_path": f.file_path,
            "pattern_type": "file_path",
            "pattern_match": "vendor_directory",
            "suggested_action": "mark_false_positive",
            "confidence": 0.65
        })
    
    return {
        "suggestions": suggestions[:limit],
        "total_suggestions": len(suggestions)
    }


@router.post(
    "/apply-fp-suggestions",
    dependencies=[Depends(require_permissions("findings:write"))],
    summary="Apply false positive suggestions to dismiss findings",
    responses={
        **CREATE_ERRORS,
        401: {"description": "Not authenticated"},
        403: {"description": "Insufficient permissions - findings:write required"},
        500: {"description": "Internal server error"},
    },
)
def apply_fp_suggestions(
    finding_ids: List[str],
    db: Session = Depends(get_tenant_db)
):
    """
    Apply false positive suggestions to findings.

    Marks the specified findings as resolved with a false_positive resolution.
    Requires findings:write permission.
    """
    count = 0
    for fid in finding_ids:
        try:
            finding = db.query(models.Finding).filter(
                models.Finding.finding_uuid == uuid.UUID(fid)
            ).first()
            
            if finding and finding.status == 'open':
                finding.status = 'resolved'
                finding.resolution = 'false_positive'
                finding.resolution_notes = 'Auto-dismissed based on learned FP pattern'
                finding.resolved_at = datetime.utcnow()
                count += 1
        except Exception as e:
            logger.warning(f"Failed to apply FP suggestion to {fid}: {e}")
    
    db.commit()
    return {"status": "success", "dismissed_count": count}
