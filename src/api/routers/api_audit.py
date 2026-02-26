"""
API Audit Router - Endpoints for API discovery and OpenAPI management.
"""
import json
import os
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import Response
from sqlalchemy.orm import Session
from sqlalchemy import text
from typing import List, Optional
from pydantic import BaseModel, Field
from datetime import datetime
from uuid import UUID
from loguru import logger

from ..dependencies import get_tenant_db
from .. import models
from src.rbac.dependencies import require_permissions
from src.auth.models import User
from src.auth.dependencies import get_current_user

router = APIRouter(
    prefix="/projects",
    tags=["api-audit"]
)

global_router = APIRouter(
    prefix="/api-audit",
    tags=["api-audit-global"]
)


# =============================================================================
# Response Models
# =============================================================================

class PathDictionaryItem(BaseModel):
    id: UUID = Field(..., description="Unique identifier for the dictionary item")
    word: str = Field(..., description="The path segment word")
    category: Optional[str] = Field(None, description="Category classification for the word")
    is_active: bool = Field(..., description="Whether this dictionary item is currently active")
    created_at: datetime = Field(..., description="Timestamp when the item was created")

    model_config = {"from_attributes": True}

class CreatePathDictionaryItem(BaseModel):
    word: str = Field(..., description="The path segment word to add")
    category: Optional[str] = Field(None, description="Category classification for the word")

class URILibraryItem(BaseModel):
    id: UUID = Field(..., description="Unique identifier for the URI library item")
    uri: str = Field(..., description="The URI string")
    description: Optional[str] = Field(None, description="Human-readable description of the URI")
    source: str = Field(..., description="Source from which the URI was discovered")
    is_active: bool = Field(..., description="Whether this URI is currently active")
    created_at: datetime = Field(..., description="Timestamp when the item was created")

    model_config = {"from_attributes": True}

class CreateURILibraryItem(BaseModel):
    uri: str = Field(..., description="The URI string to add to the library")
    description: Optional[str] = Field(None, description="Human-readable description of the URI")
    source: str = Field("manual", description="Source of this URI entry")

class APIEndpointResponse(BaseModel):
    id: str = Field(..., description="Unique identifier for the API endpoint")
    endpoint_url: str = Field(..., description="The discovered endpoint URL or path")
    http_method: Optional[str] = Field(None, description="HTTP method (GET, POST, PUT, DELETE, etc.)")
    direction: str = Field(..., description="Endpoint direction: 'serves' (inbound) or 'outbound'")
    auth_method: Optional[str] = Field(None, description="Authentication method used (bearer, api-key, etc.)")
    file_path: Optional[str] = Field(None, description="Source file path where the endpoint was discovered")
    line_number: Optional[int] = Field(None, description="Line number in the source file")
    code_snippet: Optional[str] = Field(None, description="Code snippet where the endpoint was found")
    framework: Optional[str] = Field(None, description="Detected framework (fastapi, express, etc.)")
    confidence: Optional[str] = Field(None, description="Confidence level of the discovery")

    model_config = {"from_attributes": True}


class APIAuditSummary(BaseModel):
    total_endpoints: int = Field(..., description="Total number of discovered API endpoints")
    serves_count: int = Field(..., description="Number of inbound (serves) endpoints")
    outbound_count: int = Field(..., description="Number of outbound API call endpoints")
    auth_methods: dict = Field(..., description="Count of endpoints by auth method, e.g. {'bearer': 5, 'api-key': 3}")
    frameworks: dict = Field(..., description="Count of endpoints by framework, e.g. {'fastapi': 10, 'express': 2}")
    has_openapi_spec: bool = Field(..., description="Whether an OpenAPI specification exists for this project")


class OpenAPISpecResponse(BaseModel):
    spec_content: str = Field(..., description="The raw OpenAPI specification content")
    spec_format: str = Field(..., description="Format of the specification (yaml or json)")
    version: str = Field(..., description="OpenAPI specification version")
    endpoint_count: int = Field(..., description="Number of endpoints defined in the specification")
    generated_at: datetime = Field(..., description="Timestamp when the specification was generated")


# =============================================================================
# Endpoints
# =============================================================================

@router.get("/{project_id}/api-audit/matched-credentials", dependencies=[Depends(require_permissions("admin:manage"))],
    summary="Get AI-matched credentials",
    responses={401: {"description": "Not authenticated"}, 403: {"description": "Insufficient permissions - requires admin:manage"}, 404: {"description": "Project not found"}, 500: {"description": "AI matching service error"}})
async def get_matched_credentials(project_id: str, server_url: str = None, db: Session = Depends(get_tenant_db)):
    """
    Get AI-matched credentials with certainty scores.
    Returns credentials matched to their likely target services.
    Requires admin:manage permission.
    """
    import sys
    sys.path.insert(0, '/app/execution')
    from ai_credential_matcher import match_credentials_with_llm
    
    project = db.query(models.Repository).filter(
        models.Repository.id == project_id
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    reports_dir = os.environ.get("REPORTS_DIR", "/app/vulnerability_reports")
    
    try:
        matched = await match_credentials_with_llm(project.name, server_url, reports_dir)
        return {
            "success": True,
            "credentials": matched,
            "total": len(matched),
            "high_confidence": len([m for m in matched if m['certainty'] >= 80])
        }
    except Exception as e:
        return {"success": False, "error": str(e), "credentials": []}


# =============================================================================
# AI Correlation Endpoints
# =============================================================================

@router.get("/{project_id}/api-audit/credential-url-correlations", dependencies=[Depends(require_permissions("admin:manage"))],
    summary="Get credential-URL correlations",
    responses={401: {"description": "Not authenticated"}, 403: {"description": "Insufficient permissions - requires admin:manage"}, 404: {"description": "Project not found"}, 500: {"description": "AI correlation service error"}})
async def get_credential_url_correlations(project_id: str, db: Session = Depends(get_tenant_db)):
    """
    Get AI-powered credential-URL correlations with confidence scores.
    
    Returns credentials matched to their likely target URLs based on:
    - File proximity (credentials and URLs in same file)
    - Code context (credentials used with specific URLs)
    - Environment matching (dev credentials with dev URLs)
    - Service type matching (Azure key → Azure URL)
    
    NEW: Uses OSINT-first approach (v2) that:
    1. Pre-validates URLs to check if they require authentication
    2. Skips PUBLIC URLs that don't need credentials
    3. Only maps credentials to AUTH_REQUIRED URLs
    """
    import sys
    sys.path.insert(0, '/app/execution')
    from ai_credential_matcher import correlate_credentials_to_urls_v2
    
    project = db.query(models.Repository).filter(
        models.Repository.id == project_id
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    reports_dir = os.environ.get("REPORTS_DIR", "/app/vulnerability_reports")
    
    try:
        # Use v2 with pre-validation to filter out public URLs
        correlations = await correlate_credentials_to_urls_v2(project.name, reports_dir, pre_validate=True)
        
        # Separate by auth requirement
        auth_required = [c for c in correlations if c.get('requires_auth', True)]
        public_skipped = [c for c in correlations if c.get('requires_auth') == False]
        
        return {
            "success": True,
            "correlations": auth_required,
            "total": len(auth_required),
            "high_confidence": len([c for c in auth_required if c['confidence'] >= 70]),
            "public_urls_skipped": len(public_skipped),
            "message": f"Pre-validated {len(correlations)} URLs, skipped {len(public_skipped)} public endpoints"
        }
    except Exception as e:
        logger.bind(router="api_audit").exception(f"Correlation error: {e}")
        return {"success": False, "error": str(e), "correlations": []}


@router.get("/{project_id}/api-audit/inbound-url-correlations", dependencies=[Depends(require_permissions("admin:manage"))],
    summary="Get inbound endpoint-server correlations",
    responses={401: {"description": "Not authenticated"}, 403: {"description": "Insufficient permissions - requires admin:manage"}, 404: {"description": "Project not found"}, 500: {"description": "AI correlation service error"}})
async def get_inbound_url_correlations(project_id: str, db: Session = Depends(get_tenant_db)):
    """
    Get AI-powered inbound endpoint-server correlations with confidence scores.
    
    Returns inbound API endpoints matched to their likely server URLs based on:
    - Path patterns matching server base URLs
    - Framework detection (Express routes → Node servers)
    - Environment indicators in file paths
    """
    import sys
    sys.path.insert(0, '/app/execution')
    from ai_credential_matcher import correlate_inbound_endpoints_to_servers
    
    project = db.query(models.Repository).filter(
        models.Repository.id == project_id
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    reports_dir = os.environ.get("REPORTS_DIR", "/app/vulnerability_reports")
    
    try:
        correlations = await correlate_inbound_endpoints_to_servers(project.name, reports_dir)
        return {
            "success": True,
            "correlations": correlations,
            "total": len(correlations),
            "high_confidence": len([c for c in correlations if c['confidence'] >= 70])
        }
    except Exception as e:
        return {"success": False, "error": str(e), "correlations": []}


@router.get("/{project_id}/api-audit/outbound-url-correlations", dependencies=[Depends(require_permissions("admin:manage"))],
    summary="Get outbound endpoint-server correlations",
    responses={401: {"description": "Not authenticated"}, 403: {"description": "Insufficient permissions - requires admin:manage"}, 404: {"description": "Project not found"}, 500: {"description": "AI correlation service error"}})
async def get_outbound_url_correlations(project_id: str, db: Session = Depends(get_tenant_db)):
    """
    Get AI-powered outbound endpoint-server correlations with confidence scores.
    
    Returns outbound API calls matched to their likely target server URLs based on:
    - Direct URL extraction from code
    - Environment variable references
    - Service type matching
    - Credential associations
    """
    import sys
    sys.path.insert(0, '/app/execution')
    from ai_credential_matcher import correlate_outbound_endpoints_to_servers
    
    project = db.query(models.Repository).filter(
        models.Repository.id == project_id
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    reports_dir = os.environ.get("REPORTS_DIR", "/app/vulnerability_reports")
    
    try:
        correlations = await correlate_outbound_endpoints_to_servers(project.name, reports_dir)
        return {
            "success": True,
            "correlations": correlations,
            "total": len(correlations),
            "high_confidence": len([c for c in correlations if c['confidence'] >= 70])
        }
    except Exception as e:
        return {"success": False, "error": str(e), "correlations": []}


@router.get("/{project_id}/api-audit/server-credential-correlations", dependencies=[Depends(require_permissions("admin:manage"))],
    summary="Get server-credential correlations",
    responses={401: {"description": "Not authenticated"}, 403: {"description": "Insufficient permissions - requires admin:manage"}, 404: {"description": "Project not found"}, 500: {"description": "AI correlation service error"}})
async def get_server_credential_correlations(project_id: str, db: Session = Depends(get_tenant_db)):
    """
    Get AI-powered server-credential correlations with confidence scores.
    
    Returns API servers matched to their likely associated credentials based on:
    - Environment matching (prod servers → prod credentials)
    - Domain patterns in credential code
    - Service type matching (Azure URLs → Azure keys)
    """
    import sys
    sys.path.insert(0, '/app/execution')
    from ai_credential_matcher import correlate_servers_with_credentials
    
    project = db.query(models.Repository).filter(
        models.Repository.id == project_id
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    reports_dir = os.environ.get("REPORTS_DIR", "/app/vulnerability_reports")
    
    try:
        correlations = await correlate_servers_with_credentials(project.name, reports_dir)
        return {
            "success": True,
            "correlations": correlations,
            "total": len(correlations),
            "high_confidence": len([c for c in correlations if c['top_confidence'] >= 70])
        }
    except Exception as e:
        return {"success": False, "error": str(e), "correlations": []}


@router.get("/{project_id}/api-audit/swagger-server-credentials", dependencies=[Depends(require_permissions("admin:manage"))],
    summary="Get credentials mapped to Swagger servers",
    responses={401: {"description": "Not authenticated"}, 403: {"description": "Insufficient permissions - requires admin:manage"}, 404: {"description": "Project not found"}, 500: {"description": "Credential mapping service error"}})
async def get_swagger_server_credentials(project_id: str, db: Session = Depends(get_tenant_db)):
    """
    Get credentials mapped to discovered Swagger/OpenAPI servers.
    
    This endpoint returns the servers discovered in swagger files
    along with their matched credentials for connection/authentication testing.
    
    The server URLs match what's shown in SwaggerUI > API Discovery > servers discovered.
    The credentials can be used for testing API connectivity and authentication.
    """
    import sys
    sys.path.insert(0, '/app/execution')
    from ai_credential_matcher import map_credentials_to_swagger_servers
    
    project = db.query(models.Repository).filter(
        models.Repository.id == project_id
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    reports_dir = os.environ.get("REPORTS_DIR", "/app/vulnerability_reports")
    
    try:
        mappings = await map_credentials_to_swagger_servers(project.name, reports_dir)
        return {
            "success": True,
            "mappings": mappings,
            "total_servers": len(mappings),
            "servers_with_credentials": len([m for m in mappings if m['credential_count'] > 0])
        }
    except Exception as e:
        return {"success": False, "error": str(e), "mappings": []}


@router.get("/{project_id}/api-audit/swagger-files", dependencies=[Depends(require_permissions("admin:manage"))],
    summary="List available Swagger files",
    responses={401: {"description": "Not authenticated"}, 403: {"description": "Insufficient permissions - requires admin:manage"}, 404: {"description": "Project not found"}})
def list_swagger_files(project_id: str, db: Session = Depends(get_tenant_db)):
    """
    List all available swagger files for a project.
    Returns per-server swagger specs generated by AI discovery.
    Requires admin:manage permission.
    """
    project = db.query(models.Repository).filter(
        models.Repository.id == project_id
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    reports_dir = os.environ.get("REPORTS_DIR", "/app/vulnerability_reports")
    project_dir = os.path.join(reports_dir, project.name)
    
    swagger_files = []
    
    if os.path.exists(project_dir):
        for filename in os.listdir(project_dir):
            if filename.endswith('_swagger.yaml'):
                base_name = filename.replace('_swagger.yaml', '')
                json_file = f"{base_name}_swagger.json"
                yaml_path = os.path.join(project_dir, filename)
                json_path = os.path.join(project_dir, json_file)
                
                # Extract server info from the YAML
                server_url = ""
                path_count = 0
                try:
                    import yaml
                    with open(yaml_path, 'r') as f:
                        spec = yaml.safe_load(f)
                        servers = spec.get('servers', [])
                        if servers:
                            server_url = servers[0].get('url', '')
                        path_count = len(spec.get('paths', {}))
                except (IOError, ImportError, KeyError, AttributeError) as e:
                    logger.bind(router="api_audit").debug(f"Failed to parse OpenAPI spec {yaml_path}: {str(e)}")
                except Exception as e:
                    # Catch yaml.YAMLError without requiring yaml at module level
                    if 'yaml' in str(type(e).__module__).lower():
                        logger.bind(router="api_audit").debug(f"YAML parsing error for {yaml_path}: {str(e)}")
                    else:
                        raise
                
                swagger_files.append({
                    'name': base_name,
                    'server_url': server_url,
                    'yaml_file': filename,
                    'json_file': json_file if os.path.exists(json_path) else None,
                    'path_count': path_count
                })
    
    return {"files": swagger_files, "count": len(swagger_files)}


@router.get("/{project_id}/api-audit/swagger-file/{filename}", dependencies=[Depends(require_permissions("admin:manage"))],
    summary="Download a specific Swagger file",
    responses={401: {"description": "Not authenticated"}, 403: {"description": "Insufficient permissions - requires admin:manage"}, 404: {"description": "Project or file not found"}})
def download_swagger_file(project_id: str, filename: str, db: Session = Depends(get_tenant_db)):
    """
    Download a specific swagger file by filename.
    Returns the file content with appropriate media type headers.
    Requires admin:manage permission.
    """
    project = db.query(models.Repository).filter(
        models.Repository.id == project_id
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    reports_dir = os.environ.get("REPORTS_DIR", "/app/vulnerability_reports")
    file_path = os.path.join(reports_dir, project.name, filename)
    
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")
    
    # Determine content type
    if filename.endswith('.json'):
        media_type = "application/json"
    elif filename.endswith('.yaml') or filename.endswith('.yml'):
        media_type = "application/x-yaml"
    else:
        media_type = "application/octet-stream"
    
    with open(file_path, 'r') as f:
        content = f.read()
    
    return Response(
        content=content,
        media_type=media_type,
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


@router.get("/{project_id}/api-audit", response_model=List[APIEndpointResponse], dependencies=[Depends(require_permissions("admin:manage"))],
    summary="List discovered API endpoints",
    responses={401: {"description": "Not authenticated"}, 403: {"description": "Insufficient permissions - requires admin:manage"}, 404: {"description": "Project not found"}})
def get_project_api_endpoints(project_id: str, db: Session = Depends(get_tenant_db)):
    """
    Get all discovered API endpoints for a project.
    Returns both inbound (serves) and outbound endpoints sorted by direction and URL.
    Requires admin:manage permission.
    """
    # Verify project exists
    project = db.query(models.Repository).filter(
        models.Repository.id == project_id
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    endpoints = db.query(models.APIEndpoint).filter(
        models.APIEndpoint.repository_id == project_id
    ).order_by(
        models.APIEndpoint.direction,
        models.APIEndpoint.endpoint_url
    ).all()
    
    return [
        APIEndpointResponse(
            id=str(ep.id),
            endpoint_url=ep.endpoint_url,
            http_method=ep.http_method,
            direction=ep.direction,
            auth_method=ep.auth_method,
            file_path=ep.file_path,
            line_number=ep.line_number,
            code_snippet=ep.code_snippet,
            framework=ep.framework,
            confidence=ep.confidence
        )
        for ep in endpoints
    ]


@router.get("/{project_id}/api-audit/summary", response_model=APIAuditSummary, dependencies=[Depends(require_permissions("admin:manage"))],
    summary="Get API audit summary statistics",
    responses={401: {"description": "Not authenticated"}, 403: {"description": "Insufficient permissions - requires admin:manage"}, 404: {"description": "Project not found"}})
def get_project_api_summary(project_id: str, db: Session = Depends(get_tenant_db)):
    """
    Get summary statistics for the API audit of a project.
    Includes counts by direction, auth method, framework, and OpenAPI spec availability.
    Requires admin:manage permission.
    """
    project = db.query(models.Repository).filter(
        models.Repository.id == project_id
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    endpoints = db.query(models.APIEndpoint).filter(
        models.APIEndpoint.repository_id == project_id
    ).all()
    
    # Count by direction
    serves_count = sum(1 for ep in endpoints if ep.direction == 'serves')
    outbound_count = sum(1 for ep in endpoints if ep.direction == 'outbound')
    
    # Count by auth method
    auth_methods = {}
    for ep in endpoints:
        method = ep.auth_method or 'none'
        auth_methods[method] = auth_methods.get(method, 0) + 1
    
    # Count by framework
    frameworks = {}
    for ep in endpoints:
        fw = ep.framework or 'unknown'
        frameworks[fw] = frameworks.get(fw, 0) + 1
    
    # Check for OpenAPI spec
    has_spec = db.query(models.OpenAPISpec).filter(
        models.OpenAPISpec.repository_id == project_id
    ).first() is not None
    
    return APIAuditSummary(
        total_endpoints=len(endpoints),
        serves_count=serves_count,
        outbound_count=outbound_count,
        auth_methods=auth_methods,
        frameworks=frameworks,
        has_openapi_spec=has_spec
    )


@router.get("/{project_id}/api-audit/openapi", dependencies=[Depends(require_permissions("admin:manage"))],
    summary="Download project OpenAPI specification",
    responses={401: {"description": "Not authenticated"}, 403: {"description": "Insufficient permissions - requires admin:manage"}, 404: {"description": "Project or OpenAPI specification not found"}})
def get_project_openapi_spec(
    project_id: str,
    format: str = "yaml",
    db: Session = Depends(get_tenant_db)
):
    """
    Get the OpenAPI specification for a project in YAML or JSON format.
    The specification must be generated by an API audit scan first.
    Requires admin:manage permission.
    """
    project = db.query(models.Repository).filter(
        models.Repository.id == project_id
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    spec = db.query(models.OpenAPISpec).filter(
        models.OpenAPISpec.repository_id == project_id
    ).first()
    
    if not spec:
        raise HTTPException(status_code=404, detail="OpenAPI specification not found. Run API audit scan first.")
    
    # Return appropriate content type based on format
    if format == "json":
        return Response(
            content=spec.spec_content,
            media_type="application/json",
            headers={
                "Content-Disposition": f'attachment; filename="{project.name}_openapi.json"'
            }
        )
    else:
        return Response(
            content=spec.spec_content,
            media_type="application/x-yaml",
            headers={
                "Content-Disposition": f'attachment; filename="{project.name}_openapi.yaml"'
            }
        )


from fastapi.responses import HTMLResponse

@router.get("/{project_id}/api-audit/swagger", response_class=HTMLResponse, dependencies=[Depends(require_permissions("admin:manage"))],
    summary="Serve SwaggerUI page for project",
    responses={401: {"description": "Not authenticated"}, 403: {"description": "Insufficient permissions - requires admin:manage"}, 404: {"description": "Project not found"}})
def get_project_swagger_ui(project_id: str, spec_url: str = None, db: Session = Depends(get_tenant_db)):
    """
    Serve a local SwaggerUI page that loads the project's OpenAPI spec.
    Optionally accepts spec_url to load a specific swagger file.
    Automatically injects discovered credentials for authenticated API testing.
    """
    project = db.query(models.Repository).filter(
        models.Repository.id == project_id
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # Load discovered credentials from endpoints file
    reports_dir = os.environ.get("REPORTS_DIR", "/app/vulnerability_reports")
    endpoints_file = os.path.join(reports_dir, project.name, f"{project.name}_api_endpoints.json")
    
    credentials = []
    if os.path.exists(endpoints_file):
        try:
            with open(endpoints_file, 'r') as f:
                data = json.load(f)
            
            for ep in data.get('outbound_endpoints', []):
                secret_type = ep.get('metadata', {}).get('secret_type', '')
                if not secret_type or secret_type == 'api_url':
                    continue
                
                # Extract the actual credential value from the code
                code = ep.get('code', '')
                value = ''
                
                # Parse the credential value from different formats
                if '=' in code:
                    parts = code.split('=', 1)
                    if len(parts) > 1:
                        value = parts[1].strip().strip('"').strip("'")
                elif ':' in code:
                    parts = code.split(':', 1)
                    if len(parts) > 1:
                        value = parts[1].strip().strip('"').strip("'")
                
                if value:
                    credentials.append({
                        'type': secret_type,
                        'name': ep.get('metadata', {}).get('key_name', secret_type),
                        'value': value[:100],  # Truncate for safety
                        'environment': ep.get('metadata', {}).get('environment', 'unknown'),
                        'file': os.path.basename(ep.get('path', '')),
                        'category': _classify_credential_type(secret_type)
                    })
        except Exception as e:
            print(f"Error loading credentials: {e}")
    
    # Group credentials by category for the selector
    cred_groups = {}
    for cred in credentials:
        cat = cred['category']
        if cat not in cred_groups:
            cred_groups[cat] = []
        cred_groups[cat].append(cred)
    
    # Generate credentials JSON for JavaScript
    creds_json = json.dumps(credentials)
    
    # Generate SwaggerUI HTML page with credential selector
    html_content = f'''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>API Audit - {project.name} - SwaggerUI</title>
    <link rel="stylesheet" href="https://unpkg.com/swagger-ui-dist@5.9.0/swagger-ui.css" />
    <style>
        body {{
            margin: 0;
            padding: 0;
            background: #000 !important;
        }}
        /* Force ALL text to be white/bright */
        .swagger-ui,
        .swagger-ui * {{
            color: #fff !important;
        }}
        .swagger-ui .topbar {{
            display: none;
        }}
        .swagger-ui {{
            max-width: 1400px;
            margin: 0 auto;
        }}
        .swagger-ui,
        .swagger-ui .opblock-tag,
        .swagger-ui .opblock-tag small,
        .swagger-ui .opblock-description-wrapper p,
        .swagger-ui .opblock-external-docs-wrapper p,
        .swagger-ui .opblock-title_normal p,
        .swagger-ui table thead tr th,
        .swagger-ui table thead tr td,
        .swagger-ui .response-col_status,
        .swagger-ui .response-col_links,
        .swagger-ui .response-col_description,
        .swagger-ui .responses-inner h4,
        .swagger-ui .responses-inner h5,
        .swagger-ui .opblock .opblock-section-header h4,
        .swagger-ui .parameter__name,
        .swagger-ui .parameter__type,
        .swagger-ui .parameter__deprecated,
        .swagger-ui .parameter__in,
        .swagger-ui label,
        .swagger-ui .btn,
        .swagger-ui select,
        .swagger-ui input,
        .swagger-ui .model-title,
        .swagger-ui section.models h4,
        .swagger-ui .model,
        .swagger-ui .renderedMarkdown p {{
            color: #e0e0e0 !important;
        }}
        .swagger-ui .opblock {{
            background: #1a1a1a !important;
            border-color: #333 !important;
        }}
        .swagger-ui .opblock .opblock-summary {{
            border-color: #333 !important;
        }}
        .swagger-ui .opblock .opblock-section-header {{
            background: #222 !important;
            border-color: #333 !important;
        }}
        .swagger-ui .opblock-body pre {{
            background: #111 !important;
            color: #e0e0e0 !important;
        }}
        .swagger-ui .responses-table {{
            background: #1a1a1a !important;
        }}
        .swagger-ui .model-box {{
            background: #1a1a1a !important;
        }}
        .swagger-ui .info {{
            background: transparent !important;
        }}
        .swagger-ui .info .title,
        .swagger-ui .info .title small {{
            color: #fff !important;
        }}
        .swagger-ui .scheme-container {{
            background: #111 !important;
            box-shadow: none !important;
        }}
        .swagger-ui .filter .operation-filter-input {{
            background: #1a1a1a !important;
            color: #fff !important;
            border-color: #333 !important;
        }}
        .swagger-ui section.models {{
            background: #0a0a0a !important;
            border-color: #333 !important;
        }}
        .swagger-ui .model-container {{
            background: #111 !important;
        }}
        .swagger-ui textarea {{
            background: #1a1a1a !important;
            color: #e0e0e0 !important;
            border-color: #333 !important;
        }}
        .swagger-ui select {{
            background: #1a1a1a !important;
            border-color: #333 !important;
        }}
        .swagger-ui .microlight {{
            background: #111 !important;
        }}
        .header {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 20px 40px;
            display: flex;
            justify-content: space-between;
            align-items: center;
        }}
        .header h1 {{
            margin: 0;
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
            font-size: 1.5rem;
        }}
        .header .subtitle {{
            opacity: 0.8;
            font-size: 0.9rem;
            margin-top: 4px;
        }}
        .back-btn {{
            background: rgba(255,255,255,0.2);
            border: 1px solid rgba(255,255,255,0.3);
            color: white;
            padding: 8px 16px;
            border-radius: 6px;
            cursor: pointer;
            font-size: 14px;
            text-decoration: none;
        }}
        .back-btn:hover {{
            background: rgba(255,255,255,0.3);
        }}
        .auth-bar {{
            background: #2a2a4e;
            padding: 15px 40px;
            border-bottom: 1px solid #3a3a6e;
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
            display: flex;
            align-items: center;
            gap: 15px;
            flex-wrap: wrap;
        }}
        .auth-bar-title {{
            color: #fff;
            font-size: 14px;
            font-weight: 600;
            display: flex;
            align-items: center;
            gap: 8px;
        }}
        .auth-select {{
            background: #1a1a2e;
            border: 1px solid #4a4a7e;
            color: white;
            padding: 8px 12px;
            border-radius: 6px;
            font-size: 13px;
            min-width: 180px;
            cursor: pointer;
        }}
        .auth-input {{
            background: #1a1a2e;
            border: 1px solid #4a4a7e;
            color: white;
            padding: 8px 12px;
            border-radius: 6px;
            font-size: 13px;
            min-width: 350px;
            font-family: monospace;
        }}
        .auth-input:focus, .auth-select:focus {{
            outline: none;
            border-color: #667eea;
        }}
        .auth-status {{
            color: #10b981;
            font-size: 12px;
            margin-left: auto;
        }}
        .auth-status.inactive {{
            color: #6b7280;
        }}
        .auth-btn {{
            background: #667eea;
            border: none;
            color: white;
            padding: 8px 16px;
            border-radius: 6px;
            cursor: pointer;
            font-size: 13px;
        }}
        .auth-btn:hover {{
            background: #5a6fd6;
        }}
        .auth-btn.clear {{
            background: #4a4a7e;
        }}
    </style>
</head>
<body>
    <div class="header">
        <div>
            <h1>🔍 API Audit - {project.name}</h1>
            <div class="subtitle">Interactive API Explorer powered by SwaggerUI</div>
        </div>
        <a href="javascript:history.back()" class="back-btn">← Back to Dashboard</a>
    </div>
    
    <div class="auth-bar">
        <div class="auth-bar-title">🔐 Authentication</div>
        <select id="authType" class="auth-select">
            <option value="">No Authentication</option>
            <option value="subscription">Azure Subscription Key</option>
            <option value="apikey">API Key (X-API-Key)</option>
            <option value="bearer">Bearer Token</option>
            <option value="basic">Basic Auth (user:pass)</option>
            <option value="custom">Custom Header</option>
        </select>
        <input type="text" id="authValue" class="auth-input" placeholder="Enter credential value..." />
        <input type="text" id="customHeader" class="auth-input" placeholder="Header name (e.g. X-Custom-Auth)" style="display:none; min-width:200px;" />
        <button class="auth-btn" onclick="applyAuth()">Apply</button>
        <button class="auth-btn clear" onclick="clearAuth()">Clear</button>
        <div id="authStatus" class="auth-status inactive">No credentials applied</div>
    </div>
    
    <div id="swagger-ui"></div>
    <script src="https://unpkg.com/swagger-ui-dist@5.9.0/swagger-ui-bundle.js"></script>
    <script>
        // Active credentials
        let activeAuth = {{ type: null, value: null, header: null }};
        
        // Show/hide custom header input
        document.getElementById('authType').addEventListener('change', function() {{
            const customHeaderInput = document.getElementById('customHeader');
            const valueInput = document.getElementById('authValue');
            if (this.value === 'custom') {{
                customHeaderInput.style.display = 'block';
                valueInput.placeholder = 'Header value...';
            }} else {{
                customHeaderInput.style.display = 'none';
                if (this.value === 'subscription') {{
                    valueInput.placeholder = 'Enter Ocp-Apim-Subscription-Key...';
                }} else if (this.value === 'apikey') {{
                    valueInput.placeholder = 'Enter API Key...';
                }} else if (this.value === 'bearer') {{
                    valueInput.placeholder = 'Enter Bearer Token (without Bearer prefix)...';
                }} else if (this.value === 'basic') {{
                    valueInput.placeholder = 'Enter username:password...';
                }} else {{
                    valueInput.placeholder = 'Enter credential value...';
                }}
            }}
        }});
        
        function applyAuth() {{
            const type = document.getElementById('authType').value;
            const value = document.getElementById('authValue').value.trim();
            const customHeader = document.getElementById('customHeader').value.trim();
            
            if (!type || !value) {{
                alert('Please select an auth type and enter a value');
                return;
            }}
            
            activeAuth = {{ type: type, value: value, header: customHeader }};
            
            const status = document.getElementById('authStatus');
            status.className = 'auth-status';
            const labels = {{
                'subscription': 'Azure Subscription Key',
                'apikey': 'API Key',
                'bearer': 'Bearer Token',
                'basic': 'Basic Auth',
                'custom': customHeader || 'Custom Header'
            }};
            status.innerHTML = '✓ ' + labels[type] + ' applied';
        }}
        
        function clearAuth() {{
            activeAuth = {{ type: null, value: null, header: null }};
            document.getElementById('authValue').value = '';
            document.getElementById('customHeader').value = '';
            const status = document.getElementById('authStatus');
            status.className = 'auth-status inactive';
            status.innerHTML = 'No credentials applied';
        }}
        
        // Initialize SwaggerUI with credential injection
        window.onload = function() {{
            const specUrl = "{spec_url if spec_url else f'/projects/{project_id}/api-audit/openapi?format=json'}";
            SwaggerUIBundle({{
                url: specUrl,
                dom_id: '#swagger-ui',
                deepLinking: true,
                presets: [
                    SwaggerUIBundle.presets.apis,
                    SwaggerUIBundle.SwaggerUIStandalonePreset
                ],
                layout: "BaseLayout",
                defaultModelsExpandDepth: -1,
                docExpansion: "list",
                filter: true,
                tryItOutEnabled: true,
                
                // Inject credentials into requests
                requestInterceptor: function(request) {{
                    if (activeAuth.type && activeAuth.value) {{
                        switch (activeAuth.type) {{
                            case 'subscription':
                                request.headers['Ocp-Apim-Subscription-Key'] = activeAuth.value;
                                break;
                            case 'apikey':
                                request.headers['X-API-Key'] = activeAuth.value;
                                request.headers['api-key'] = activeAuth.value;
                                break;
                            case 'bearer':
                                request.headers['Authorization'] = 'Bearer ' + activeAuth.value;
                                break;
                            case 'basic':
                                request.headers['Authorization'] = 'Basic ' + btoa(activeAuth.value);
                                break;
                            case 'custom':
                                if (activeAuth.header) {{
                                    request.headers[activeAuth.header] = activeAuth.value;
                                }}
                                break;
                        }}
                    }}
                    return request;
                }}
            }});
        }};
    </script>
</body>
</html>'''
    
    return HTMLResponse(content=html_content)


def _classify_credential_type(secret_type: str) -> str:
    """Classify credential type for UI grouping."""
    st = secret_type.lower()
    if 'azure' in st or 'subscription' in st:
        return 'Azure API'
    elif 'mixpanel' in st:
        return 'Mixpanel'
    elif 'firebase' in st:
        return 'Firebase'
    elif 'instabug' in st:
        return 'Instabug'
    elif 'cognito' in st:
        return 'AWS Cognito'
    elif 'box' in st:
        return 'Box API'
    elif 'github' in st:
        return 'GitHub'
    elif 'bearer' in st or 'token' in st:
        return 'Bearer Token'
    elif 'api_key' in st or 'apikey' in st:
        return 'API Key'
    else:
        return 'Other'


def _match_credential_to_server(credential: dict, server_url: str) -> bool:
    """Match a credential to a server based on environment and URL patterns."""
    cred_env = credential.get('environment', '').lower()
    url_lower = server_url.lower()
    
    # Direct environment matching
    if 'prod' in cred_env and ('api.' in url_lower or '/prod' in url_lower):
        return True
    if 'stage' in cred_env and 'stage' in url_lower:
        return True
    if 'test' in cred_env and 'test' in url_lower:
        return True
    if any(x in cred_env for x in ['dev', 'qa', 'circle']) and any(x in url_lower for x in ['dev', 'qa', 'circle']):
        return True
    
    # Azure API matching
    if 'azure' in credential.get('type', '').lower() and 'azure' in url_lower:
        return True
    
    # Cognito/ECIM matching
    if 'cognito' in cred_env and 'ecim' in url_lower:
        return True
    
    # EDP matching
    if 'edp' in cred_env and ('azure-api.net' in url_lower or 'edp' in url_lower):
        return True
    
    # SleepIQ matching
    if 'sleepiq' in cred_env and 'sleepiq' in url_lower:
        return True
    
    return False


@router.get("/{project_id}/api-audit/server-testing", response_class=HTMLResponse, dependencies=[Depends(require_permissions("admin:manage"))],
    summary="Serve API discovery testing UI",
    responses={401: {"description": "Not authenticated"}, 403: {"description": "Insufficient permissions - requires admin:manage"}, 404: {"description": "Project not found"}})
def get_server_testing_ui(project_id: str, db: Session = Depends(get_tenant_db)):
    """
    Server Testing UI - Clean table layout with modal-based AI discovery.
    Displays discovered servers and allows AI-powered API path discovery per server.
    Requires admin:manage permission.
    """
    # Debug: Log database connection info
    try:
        db_url = str(db.get_bind().url)
        logger.bind(router="api_audit", endpoint="get_swagger_server_suggestions").info(f"server-testing: Using database: {db_url}, project_id: {project_id}")
    except Exception as e:
        logger.bind(router="api_audit", endpoint="get_swagger_server_suggestions").warning(f"server-testing: Could not get db url: {e}")
    
    project = db.query(models.Repository).filter(
        models.Repository.id == project_id
    ).first()
    
    if not project:
        # Debug: Check total repos in this db
        total = db.query(models.Repository).count()
        logger.bind(router="api_audit", endpoint="get_swagger_server_suggestions").warning(f"server-testing: Project {project_id} not found. Total repos in db: {total}")
        raise HTTPException(status_code=404, detail="Project not found")
    
    # Load data from files
    reports_dir = os.environ.get("REPORTS_DIR", "/app/vulnerability_reports")
    openapi_file = os.path.join(reports_dir, project.name, f"{project.name}_openapi.yaml")
    
    servers = []
    
    # Load servers from OpenAPI
    if os.path.exists(openapi_file):
        try:
            with open(openapi_file, 'r') as f:
                spec = yaml.safe_load(f)
                for s in spec.get('servers', []):
                    servers.append(s.get('url', ''))
        except:
            pass
    
    servers_json = json.dumps(servers[:30])  # Limit to 30 servers
    
    html_content = f'''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>API Discovery - {project.name}</title>
    <style>
        * {{ box-sizing: border-box; }}
        body {{
            margin: 0;
            padding: 0;
            background: #0f0f1e;
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
            color: #e0e0e0;
        }}
        .header {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 20px 40px;
            display: flex;
            justify-content: space-between;
            align-items: center;
        }}
        .header h1 {{ margin: 0; font-size: 1.4rem; }}
        .header .subtitle {{ opacity: 0.85; font-size: 0.85rem; margin-top: 4px; }}
        .back-btn {{
            background: rgba(255,255,255,0.2);
            color: white;
            padding: 8px 16px;
            border-radius: 8px;
            text-decoration: none;
            font-size: 13px;
        }}
        .container {{ padding: 30px 40px; max-width: 1400px; margin: 0 auto; }}
        
        /* Table */
        .server-table {{
            width: 100%;
            border-collapse: collapse;
            background: #1a1a2e;
            border-radius: 12px;
            overflow: hidden;
        }}
        .server-table th {{
            background: #22223a;
            padding: 16px 20px;
            text-align: left;
            font-weight: 600;
            color: #a0a0c0;
            font-size: 12px;
            text-transform: uppercase;
            letter-spacing: 0.5px;
        }}
        .server-table td {{
            padding: 14px 20px;
            border-bottom: 1px solid #2a2a4e;
        }}
        .server-table tr:last-child td {{ border-bottom: none; }}
        .server-table tr:hover {{ background: #22223a; }}
        
        .server-url {{
            font-family: monospace;
            font-size: 13px;
            color: #60a5fa;
        }}
        
        .level-select {{
            background: #0f0f1e;
            color: white;
            border: 1px solid #4a4a7e;
            padding: 8px 12px;
            border-radius: 6px;
            font-size: 13px;
            cursor: pointer;
        }}
        
        .discover-btn {{
            background: linear-gradient(135deg, #10b981 0%, #059669 100%);
            color: white;
            border: none;
            padding: 10px 20px;
            border-radius: 8px;
            cursor: pointer;
            font-size: 13px;
            font-weight: 600;
            display: inline-flex;
            align-items: center;
            gap: 8px;
        }}
        .discover-btn:hover {{ opacity: 0.9; }}
        .discover-btn:disabled {{ opacity: 0.5; cursor: not-allowed; }}
        
        /* Modal */
        .modal-overlay {{
            display: none;
            position: fixed;
            top: 0;
            left: 0;
            right: 0;
            bottom: 0;
            background: rgba(0,0,0,0.8);
            z-index: 1000;
            justify-content: center;
            align-items: center;
        }}
        .modal-overlay.active {{ display: flex; }}
        
        .modal {{
            background: #1a1a2e;
            border-radius: 16px;
            width: 90%;
            max-width: 700px;
            max-height: 80vh;
            overflow: hidden;
            box-shadow: 0 20px 60px rgba(0,0,0,0.5);
        }}
        .modal-header {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 20px 24px;
        }}
        .modal-header h2 {{ margin: 0; font-size: 1.1rem; }}
        .modal-header .server {{ opacity: 0.85; font-size: 12px; font-family: monospace; margin-top: 6px; }}
        
        .modal-body {{
            padding: 24px;
            max-height: 50vh;
            overflow-y: auto;
        }}
        
        .progress-section {{
            text-align: center;
            padding: 40px 20px;
        }}
        .spinner {{
            width: 50px;
            height: 50px;
            border: 4px solid #2a2a4e;
            border-top-color: #667eea;
            border-radius: 50%;
            animation: spin 1s linear infinite;
            margin: 0 auto 20px;
        }}
        @keyframes spin {{ to {{ transform: rotate(360deg); }} }}
        .progress-text {{ color: #a0a0c0; font-size: 14px; }}
        
        .results-table {{
            width: 100%;
            border-collapse: collapse;
        }}
        .results-table th {{
            background: #22223a;
            padding: 12px 16px;
            text-align: left;
            font-size: 11px;
            text-transform: uppercase;
            color: #a0a0c0;
        }}
        .results-table td {{
            padding: 10px 16px;
            border-bottom: 1px solid #2a2a4e;
            font-size: 13px;
        }}
        .method-badge {{
            display: inline-block;
            padding: 4px 10px;
            border-radius: 4px;
            font-size: 11px;
            font-weight: 700;
            min-width: 50px;
            text-align: center;
        }}
        .method-badge.get {{ background: #22c55e; color: #052e16; }}
        .method-badge.post {{ background: #3b82f6; color: white; }}
        .method-badge.put {{ background: #f59e0b; color: #422006; }}
        .method-badge.delete {{ background: #ef4444; color: white; }}
        .method-badge.patch {{ background: #8b5cf6; color: white; }}
        
        .confidence {{
            color: #a0a0c0;
            font-size: 11px;
        }}
        
        .modal-footer {{
            padding: 16px 24px;
            border-top: 1px solid #2a2a4e;
            display: flex;
            justify-content: flex-end;
            gap: 12px;
        }}
        .btn {{
            padding: 10px 24px;
            border-radius: 8px;
            font-size: 13px;
            font-weight: 600;
            cursor: pointer;
            border: none;
        }}
        .btn-cancel {{ background: #4a4a7e; color: white; }}
        .btn-save {{ background: linear-gradient(135deg, #10b981 0%, #059669 100%); color: white; }}
        .btn:disabled {{ opacity: 0.5; cursor: not-allowed; }}
        
        .empty-state {{
            text-align: center;
            padding: 60px 20px;
            color: #6b7280;
        }}
        .stats {{
            display: flex;
            gap: 20px;
            margin-bottom: 16px;
            font-size: 13px;
            color: #a0a0c0;
        }}
        .stat-value {{ color: white; font-weight: 600; }}
    </style>
</head>
<body>
    <div class="header">
        <div>
            <h1>🔍 API Discovery</h1>
            <div class="subtitle">{project.name} - Discover API paths per server</div>
        </div>
        <a href="javascript:history.back()" class="back-btn">← Back</a>
    </div>
    
    <div class="container">
        <div class="stats">
            <div><span class="stat-value">{len(servers)}</span> servers discovered</div>
        </div>
        
        <table class="server-table">
            <thead>
                <tr>
                    <th style="width: 55%;">Server URL</th>
                    <th style="width: 15%;">Level</th>
                    <th style="width: 30%;">Action</th>
                </tr>
            </thead>
            <tbody id="serverTableBody">
            </tbody>
        </table>
    </div>
    
    <!-- Modal -->
    <div class="modal-overlay" id="modal">
        <div class="modal">
            <div class="modal-header">
                <h2 id="modalTitle">🤖 AI Discovery</h2>
                <div class="server" id="modalServer"></div>
            </div>
            <div class="modal-body" id="modalBody">
                <!-- Dynamic content -->
            </div>
            <div class="modal-footer" id="modalFooter">
                <button class="btn btn-cancel" onclick="closeModal()">Cancel</button>
                <button class="btn btn-save" id="saveBtn" onclick="saveResults()" disabled>Save as Swagger</button>
            </div>
        </div>
    </div>
    
    <script>
        const servers = {servers_json};
        let currentServerIdx = null;
        let discoveredPaths = [];
        
        function renderTable() {{
            const tbody = document.getElementById('serverTableBody');
            if (servers.length === 0) {{
                tbody.innerHTML = `<tr><td colspan="3" class="empty-state">No servers discovered. Run an API scan first.</td></tr>`;
                return;
            }}
            
            tbody.innerHTML = servers.map((url, idx) => `
                <tr>
                    <td><span class="server-url">${{url}}</span></td>
                    <td>
                        <select class="level-select" id="level-${{idx}}">
                            <option value="light">Light</option>
                            <option value="medium" selected>Medium</option>
                            <option value="full">Full</option>
                        </select>
                    </td>
                    <td>
                        <button class="discover-btn" onclick="startDiscovery(${{idx}})">
                            🤖 AI Discover
                        </button>
                    </td>
                </tr>
            `).join('');
        }}
        
        function openModal(serverUrl) {{
            document.getElementById('modalServer').textContent = serverUrl;
            document.getElementById('modal').classList.add('active');
        }}
        
        function closeModal() {{
            document.getElementById('modal').classList.remove('active');
            currentServerIdx = null;
            discoveredPaths = [];
        }}
        
        function showProgress(message) {{
            document.getElementById('modalBody').innerHTML = `
                <div class="progress-section">
                    <div class="spinner"></div>
                    <div class="progress-text">${{message}}</div>
                </div>
            `;
            document.getElementById('saveBtn').disabled = true;
        }}
        
        function showResults(paths, stats) {{
            discoveredPaths = paths;
            
            if (paths.length === 0) {{
                document.getElementById('modalBody').innerHTML = `
                    <div class="empty-state">
                        <p>No API paths discovered.</p>
                        <p style="font-size: 12px;">Try increasing the scan level or checking if the server is reachable.</p>
                    </div>
                `;
                document.getElementById('saveBtn').disabled = true;
                return;
            }}
            
            const tableRows = paths.map(p => `
                <tr>
                    <td><span class="method-badge ${{p.method.toLowerCase()}}">${{p.method}}</span></td>
                    <td style="font-family: monospace; font-size: 12px;">${{p.path}}</td>
                    <td><span class="confidence">${{(p.confidence * 100).toFixed(0)}}%</span></td>
                    <td style="font-size: 11px; color: #6b7280;">${{p.source || 'ai'}}</td>
                </tr>
            `).join('');
            
            document.getElementById('modalBody').innerHTML = `
                <div class="stats">
                    <div><span class="stat-value">${{paths.length}}</span> paths discovered</div>
                    <div><span class="stat-value">${{stats.probe_hits || 0}}</span> probe hits</div>
                </div>
                <table class="results-table">
                    <thead>
                        <tr>
                            <th>Method</th>
                            <th>Path</th>
                            <th>Confidence</th>
                            <th>Source</th>
                        </tr>
                    </thead>
                    <tbody>${{tableRows}}</tbody>
                </table>
            `;
            
            document.getElementById('saveBtn').disabled = false;
        }}
        
        async function startDiscovery(idx) {{
            currentServerIdx = idx;
            const serverUrl = servers[idx];
            const level = document.getElementById(`level-${{idx}}`).value;
            
            openModal(serverUrl);
            showProgress(`Probing server (${{level}} mode)...`);
            
            try {{
                const response = await fetch('/projects/{project_id}/api-audit/ai-discover', {{
                    method: 'POST',
                    headers: {{ 'Content-Type': 'application/json' }},
                    body: JSON.stringify({{ server_url: serverUrl, level: level }})
                }});
                
                const result = await response.json();
                
                if (result.success) {{
                    // Get combined paths from the response
                    const paths = result.combined_paths || [];
                    showResults(paths, {{
                        probe_hits: result.probe_hits || 0,
                        ai_paths: result.ai_paths || 0
                    }});
                }} else {{
                    document.getElementById('modalBody').innerHTML = `
                        <div class="empty-state" style="color: #ef4444;">
                            <p>Discovery failed</p>
                            <p style="font-size: 12px;">${{result.error || 'Unknown error'}}</p>
                        </div>
                    `;
                }}
            }} catch (e) {{
                document.getElementById('modalBody').innerHTML = `
                    <div class="empty-state" style="color: #ef4444;">
                        <p>Request failed</p>
                        <p style="font-size: 12px;">${{e.message}}</p>
                    </div>
                `;
            }}
        }}
        
        async function saveResults() {{
            if (currentServerIdx === null || discoveredPaths.length === 0) return;
            
            const serverUrl = servers[currentServerIdx];
            document.getElementById('saveBtn').disabled = true;
            document.getElementById('saveBtn').textContent = 'Saving...';
            
            try {{
                const response = await fetch('/projects/{project_id}/api-audit/save-server-swagger', {{
                    method: 'POST',
                    headers: {{ 'Content-Type': 'application/json' }},
                    body: JSON.stringify({{ 
                        server_url: serverUrl,
                        paths: discoveredPaths
                    }})
                }});
                
                const result = await response.json();
                
                if (result.success) {{
                    alert(`✓ Swagger saved!\\n\\nFile: ${{result.filename}}\\n\\nYou can now test these paths in SwaggerUI.`);
                    closeModal();
                }} else {{
                    alert('Failed to save: ' + (result.error || 'Unknown error'));
                    document.getElementById('saveBtn').textContent = 'Save as Swagger';
                    document.getElementById('saveBtn').disabled = false;
                }}
            }} catch (e) {{
                alert('Error: ' + e.message);
                document.getElementById('saveBtn').textContent = 'Save as Swagger';
                document.getElementById('saveBtn').disabled = false;
            }}
        }}
        
        renderTable();
    </script>
</body>
</html>'''
    
    return HTMLResponse(content=html_content)


@router.post("/{project_id}/api-audit/save-server-swagger", dependencies=[Depends(require_permissions("admin:manage"))],
    summary="Save discovered paths as Swagger file",
    responses={401: {"description": "Not authenticated"}, 403: {"description": "Insufficient permissions - requires admin:manage"}, 404: {"description": "Project not found"}, 400: {"description": "Missing server_url or no paths provided"}})
async def save_server_swagger(project_id: str, request_data: dict, db: Session = Depends(get_tenant_db)):
    """
    Save discovered API paths as a Swagger/OpenAPI file for a specific server.
    Creates both YAML and JSON versions of the specification.
    Requires admin:manage permission.
    """
    project = db.query(models.Repository).filter(
        models.Repository.id == project_id
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    server_url = request_data.get('server_url', '').strip()
    paths = request_data.get('paths', [])
    
    if not server_url:
        return {"success": False, "error": "Missing server_url"}
    
    if not paths:
        return {"success": False, "error": "No paths to save"}
    
    # Create OpenAPI spec for this server
    from urllib.parse import urlparse
    parsed = urlparse(server_url)
    server_name = parsed.netloc.replace('.', '_').replace(':', '_')
    
    spec = {
        "openapi": "3.0.3",
        "info": {
            "title": f"API - {parsed.netloc}",
            "version": "1.0.0",
            "description": f"Discovered API spec for {server_url}"
        },
        "servers": [{"url": server_url, "description": "Target server"}],
        "paths": {}
    }
    
    # Add paths
    for path_info in paths:
        path = path_info.get('path', '')
        method = path_info.get('method', 'GET').lower()
        if not path or not path.startswith('/'):
            continue
        
        if path not in spec['paths']:
            spec['paths'][path] = {}
        
        spec['paths'][path][method] = {
            "summary": path_info.get('description', f"Discovered endpoint"),
            "responses": {"200": {"description": "Success"}},
            "x-confidence": path_info.get('confidence', 0.5),
            "x-source": path_info.get('source', 'ai')
        }
    
    # Save to file
    reports_dir = os.environ.get("REPORTS_DIR", "/app/vulnerability_reports")
    project_dir = os.path.join(reports_dir, project.name)
    os.makedirs(project_dir, exist_ok=True)
    
    filename = f"{server_name}_swagger.yaml"
    filepath = os.path.join(project_dir, filename)
    
    try:
        with open(filepath, 'w') as f:
            yaml.dump(spec, f, default_flow_style=False)
        
        # Also save JSON version
        json_path = os.path.join(project_dir, f"{server_name}_swagger.json")
        with open(json_path, 'w') as f:
            json.dump(spec, f, indent=2)
        
        return {
            "success": True,
            "filename": filename,
            "path_count": len(spec['paths']),
            "filepath": filepath
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

@router.post("/{project_id}/api-audit/ai-discover", dependencies=[Depends(require_permissions("admin:manage"))],
    summary="AI-powered API path discovery",
    responses={401: {"description": "Not authenticated"}, 403: {"description": "Insufficient permissions - requires admin:manage"}, 404: {"description": "Project not found"}, 400: {"description": "Missing server_url"}, 500: {"description": "AI discovery module not available or error"}})
async def ai_discover_paths(project_id: str, request_data: dict, db: Session = Depends(get_tenant_db)):
    """
    AI-powered API path discovery.
    Combines code analysis, server probing, and LLM inference to discover API endpoints.
    Requires admin:manage permission.
    """
    from pathlib import Path as PathlibPath
    
    project = db.query(models.Repository).filter(
        models.Repository.id == project_id
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    server_url = request_data.get('server_url', '').strip()
    level = request_data.get('level', 'medium')
    
    if not server_url:
        return {"success": False, "error": "Missing server_url"}
    
    if level not in ['light', 'medium', 'full']:
        level = 'medium'
    
    # Load existing data
    reports_dir = os.environ.get("REPORTS_DIR", "/app/vulnerability_reports")
    endpoints_file = os.path.join(reports_dir, project.name, f"{project.name}_api_endpoints.json")
    openapi_file = os.path.join(reports_dir, project.name, f"{project.name}_openapi.yaml")
    openapi_json = os.path.join(reports_dir, project.name, f"{project.name}_openapi.json")
    
    # Load credentials for authenticated probing
    credentials = []
    if os.path.exists(endpoints_file):
        try:
            with open(endpoints_file, 'r') as f:
                data = json.load(f)
            for ep in data.get('outbound_endpoints', []):
                secret_type = ep.get('metadata', {}).get('secret_type', '')
                if not secret_type or secret_type == 'api_url':
                    continue
                code = ep.get('code', '')
                value = ''
                if '=' in code:
                    parts = code.split('=', 1)
                    if len(parts) > 1:
                        value = parts[1].strip().strip('"').strip("'")
                if value:
                    credentials.append({
                        'type': secret_type,
                        'value': value[:100],
                        'environment': ep.get('metadata', {}).get('environment', 'unknown'),
                    })
        except:
            pass
    
    # Filter credentials for this server
    matched_creds = [c for c in credentials if _match_credential_to_server(c, server_url)]
    
    # Try to import and run AI discovery
    try:
        import sys
        sys.path.insert(0, '/app/execution')
        from ai_api_discovery import discover_api_paths
        
        result = await discover_api_paths(
            project_id=project_id,
            project_name=project.name,
            server_url=server_url,
            level=level,
            credentials=matched_creds,
            repo_path=None  # Code already scanned, use existing clues
        )
        
        # Count results
        probe_hits = len([p for p in result.get('probe_results', []) if p.get('exists')])
        ai_paths = len(result.get('ai_paths', []))
        combined_paths = result.get('combined_paths', [])
        
        # Update OpenAPI spec with discovered paths
        paths_added = 0
        base_spec = {"openapi": "3.0.3", "info": {"title": f"{project.name} API", "version": "1.0.0"}, "paths": {}, "servers": []}
        
        if os.path.exists(openapi_file):
            try:
                with open(openapi_file, 'r') as f:
                    base_spec = yaml.safe_load(f) or base_spec
            except:
                pass
        
        # Add discovered paths
        for path_info in combined_paths:
            path = path_info.get('path', '')
            method = path_info.get('method', 'GET').lower()
            if not path or not path.startswith('/'):
                continue
            
            if path not in base_spec.get('paths', {}):
                base_spec['paths'][path] = {}
            
            if method not in base_spec['paths'][path]:
                base_spec['paths'][path][method] = {
                    "summary": path_info.get('description', f"Discovered via AI ({path_info.get('source', 'unknown')})"),
                    "responses": {"200": {"description": "Success"}},
                    "x-discovered-by": path_info.get('source', 'ai'),
                    "x-confidence": path_info.get('confidence', 0.5)
                }
                paths_added += 1
        
        # Save updated spec
        if paths_added > 0:
            try:
                with open(openapi_file, 'w') as f:
                    yaml.dump(base_spec, f, default_flow_style=False)
                with open(openapi_json, 'w') as f:
                    json.dump(base_spec, f, indent=2)
            except Exception as e:
                result['errors'] = result.get('errors', []) + [f"Failed to save spec: {e}"]
        
        return {
            "success": True,
            "paths_discovered": paths_added,
            "probe_hits": probe_hits,
            "ai_paths": ai_paths,
            "total_paths": len(base_spec.get('paths', {})),
            "level": level,
            "combined_paths": combined_paths,
            "errors": result.get('errors', [])
        }
        
    except ImportError as e:
        return {"success": False, "error": f"AI discovery module not available: {e}"}
    except Exception as e:
        return {"success": False, "error": str(e)[:200]}


@router.post("/{project_id}/api-audit/spider-server", dependencies=[Depends(require_permissions("admin:manage"))],
    summary="Spider a server for OpenAPI spec",
    responses={401: {"description": "Not authenticated"}, 403: {"description": "Insufficient permissions - requires admin:manage"}, 404: {"description": "Project not found"}, 400: {"description": "Missing server_url"}, 500: {"description": "Server error during spidering"}})
async def spider_single_server(project_id: str, request_data: dict, db: Session = Depends(get_tenant_db)):
    """
    Spider a single server for OpenAPI specification.
    Probes common OpenAPI paths and merges discovered endpoints into the project spec.
    Requires admin:manage permission.
    """
    import httpx
    
    project = db.query(models.Repository).filter(
        models.Repository.id == project_id
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    server_url = request_data.get('server_url', '').strip()
    if not server_url:
        return {"success": False, "error": "Missing server_url"}
    
    # Load existing data
    reports_dir = os.environ.get("REPORTS_DIR", "/app/vulnerability_reports")
    endpoints_file = os.path.join(reports_dir, project.name, f"{project.name}_api_endpoints.json")
    openapi_file = os.path.join(reports_dir, project.name, f"{project.name}_openapi.yaml")
    openapi_json = os.path.join(reports_dir, project.name, f"{project.name}_openapi.json")
    
    # Load credentials
    credentials = []
    if os.path.exists(endpoints_file):
        try:
            with open(endpoints_file, 'r') as f:
                data = json.load(f)
            for ep in data.get('outbound_endpoints', []):
                secret_type = ep.get('metadata', {}).get('secret_type', '')
                if not secret_type or secret_type == 'api_url':
                    continue
                code = ep.get('code', '')
                value = ''
                if '=' in code:
                    parts = code.split('=', 1)
                    if len(parts) > 1:
                        value = parts[1].strip().strip('"').strip("'")
                if value:
                    credentials.append({
                        'type': secret_type,
                        'value': value[:100],
                        'environment': ep.get('metadata', {}).get('environment', 'unknown'),
                    })
        except:
            pass
    
    # Find matching credential
    matched_cred = None
    for c in credentials:
        if _match_credential_to_server(c, server_url):
            matched_cred = c
            break
    
    # Build headers
    headers = {"User-Agent": "AuditGH-Spider/1.0"}
    if matched_cred:
        cred_type = matched_cred.get('type', '').lower()
        cred_value = matched_cred.get('value', '')
        if 'azure' in cred_type:
            headers['Ocp-Apim-Subscription-Key'] = cred_value
        elif 'bearer' in cred_type or 'token' in cred_type:
            headers['Authorization'] = f'Bearer {cred_value}'
        else:
            headers['X-API-Key'] = cred_value
    
    # Probe common OpenAPI paths
    spider_paths = ['/swagger.json', '/openapi.json', '/api-docs', '/v3/api-docs', '/v2/api-docs', '/swagger.yaml', '/openapi.yaml']
    discovered_spec = None
    source_url = None
    
    async with httpx.AsyncClient(timeout=10.0, follow_redirects=True) as client:
        for path in spider_paths:
            try:
                url = server_url.rstrip('/') + path
                resp = await client.get(url, headers=headers)
                if resp.status_code == 200:
                    try:
                        spec = resp.json()
                        if 'paths' in spec or 'openapi' in spec or 'swagger' in spec:
                            discovered_spec = spec
                            source_url = url
                            break
                    except:
                        # Try YAML
                        try:
                            spec = yaml.safe_load(resp.text)
                            if isinstance(spec, dict) and ('paths' in spec or 'openapi' in spec):
                                discovered_spec = spec
                                source_url = url
                                break
                        except:
                            pass
            except:
                pass
    
    if not discovered_spec:
        return {"success": True, "paths_found": 0, "message": "No OpenAPI spec found"}
    
    # Load existing spec
    base_spec = {"openapi": "3.0.3", "info": {"title": f"{project.name} API", "version": "1.0.0"}, "paths": {}, "servers": []}
    if os.path.exists(openapi_file):
        try:
            with open(openapi_file, 'r') as f:
                base_spec = yaml.safe_load(f) or base_spec
        except:
            pass
    
    # Merge discovered paths
    paths_added = 0
    for path, methods in discovered_spec.get('paths', {}).items():
        if path not in base_spec.get('paths', {}):
            base_spec['paths'][path] = methods
            paths_added += len([m for m in methods.keys() if m.lower() in ['get', 'post', 'put', 'delete', 'patch']])
    
    # Save updated spec
    try:
        with open(openapi_file, 'w') as f:
            yaml.dump(base_spec, f, default_flow_style=False)
        with open(openapi_json, 'w') as f:
            json.dump(base_spec, f, indent=2)
    except Exception as e:
        return {"success": False, "error": str(e)}
    
    return {
        "success": True,
        "paths_found": paths_added,
        "source_url": source_url,
        "total_paths": len(base_spec.get('paths', {}))
    }


@router.post("/{project_id}/api-audit/test-auth", dependencies=[Depends(require_permissions("admin:manage"))],
    summary="Test server authorization with credentials",
    responses={401: {"description": "Not authenticated"}, 403: {"description": "Insufficient permissions - requires admin:manage"}, 400: {"description": "Missing server_url or credential_value"}, 500: {"description": "Server error or timeout during auth test"}})
async def test_server_authorization(project_id: str, request_data: dict, db: Session = Depends(get_tenant_db)):
    """
    Test authorization against a server with the provided credential.
    Makes a HEAD request to the server with appropriate auth headers.
    Requires admin:manage permission.
    """
    import httpx
    
    server_url = request_data.get('server_url', '')
    cred_type = request_data.get('credential_type', '').lower()
    cred_value = request_data.get('credential_value', '')
    
    if not server_url or not cred_value:
        return {"success": False, "error": "Missing params"}
    
    # Build headers based on credential type
    headers = {"User-Agent": "AuditGH-API-Tester/1.0"}
    
    if 'bearer' in cred_type or 'token' in cred_type:
        headers['Authorization'] = f'Bearer {cred_value}'
    elif 'azure' in cred_type or 'subscription' in cred_type:
        headers['Ocp-Apim-Subscription-Key'] = cred_value
    elif 'api_key' in cred_type:
        headers['X-API-Key'] = cred_value
    elif 'mixpanel' in cred_type:
        import base64
        headers['Authorization'] = f'Basic {base64.b64encode(f"{cred_value}:".encode()).decode()}'
    else:
        headers['Authorization'] = f'Bearer {cred_value}'
    
    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            response = await client.head(server_url, headers=headers, follow_redirects=True)
            
            # Consider 2xx, 3xx, 401, 403 as "connectivity" successes (server responded)
            if response.status_code < 500:
                return {
                    "success": response.status_code < 400,
                    "status_code": response.status_code,
                    "message": "Auth accepted" if response.status_code < 400 else "Auth rejected"
                }
            else:
                return {"success": False, "status_code": response.status_code, "error": "Server error"}
    except httpx.TimeoutException:
        return {"success": False, "error": "Timeout"}
    except httpx.RequestError as e:
        return {"success": False, "error": str(type(e).__name__)}
    except Exception as e:
        return {"success": False, "error": str(e)[:50]}


@router.get("/{project_id}/api-audit/server/{server_index}/swagger", response_class=HTMLResponse, dependencies=[Depends(require_permissions("admin:manage"))],
    summary="Serve per-server SwaggerUI with pre-filled auth",
    responses={401: {"description": "Not authenticated"}, 403: {"description": "Insufficient permissions - requires admin:manage"}, 404: {"description": "Project or server index not found"}})
def get_server_swagger_ui(project_id: str, server_index: int, db: Session = Depends(get_tenant_db)):
    """
    Per-server SwaggerUI with pre-filled auth credentials.
    Shows only the paths for this specific server with auth already configured.
    Requires admin:manage permission.
    """
    project = db.query(models.Repository).filter(
        models.Repository.id == project_id
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # Load data
    reports_dir = os.environ.get("REPORTS_DIR", "/app/vulnerability_reports")
    endpoints_file = os.path.join(reports_dir, project.name, f"{project.name}_api_endpoints.json")
    openapi_file = os.path.join(reports_dir, project.name, f"{project.name}_openapi.yaml")
    
    servers = []
    credentials = []
    
    # Load servers from OpenAPI
    if os.path.exists(openapi_file):
        try:
            with open(openapi_file, 'r') as f:
                spec = yaml.safe_load(f)
                for s in spec.get('servers', []):
                    servers.append(s.get('url', ''))
        except:
            pass
    
    if server_index >= len(servers):
        raise HTTPException(status_code=404, detail="Server not found")
    
    server_url = servers[server_index]
    
    # Load credentials
    if os.path.exists(endpoints_file):
        try:
            with open(endpoints_file, 'r') as f:
                data = json.load(f)
            
            for ep in data.get('outbound_endpoints', []):
                secret_type = ep.get('metadata', {}).get('secret_type', '')
                if not secret_type or secret_type == 'api_url':
                    continue
                
                code = ep.get('code', '')
                value = ''
                if '=' in code:
                    parts = code.split('=', 1)
                    if len(parts) > 1:
                        value = parts[1].strip().strip('"').strip("'")
                
                if value:
                    credentials.append({
                        'type': secret_type,
                        'name': ep.get('metadata', {}).get('key_name', secret_type),
                        'value': value[:100],
                        'environment': ep.get('metadata', {}).get('environment', 'unknown'),
                    })
        except:
            pass
    
    # Match credentials to this server
    matched_creds = [c for c in credentials if _match_credential_to_server(c, server_url)]
    
    # Determine auth type for this server
    auth_type = 'none'
    auth_value = ''
    auth_scheme_name = 'apiKey'
    
    if matched_creds:
        cred = matched_creds[0]  # Use first matching credential
        cred_type = cred.get('type', '').lower()
        auth_value = cred.get('value', '')
        
        if 'azure' in cred_type or 'subscription' in cred_type:
            auth_type = 'apiKey'
            auth_scheme_name = 'AzureSubscription'
        elif 'bearer' in cred_type or 'jwt' in cred_type or 'token' in cred_type:
            auth_type = 'bearer'
            auth_scheme_name = 'BearerAuth'
        elif 'basic' in cred_type or 'mixpanel' in cred_type:
            auth_type = 'basic'
            auth_scheme_name = 'BasicAuth'
        else:
            auth_type = 'apiKey'
            auth_scheme_name = 'ApiKey'
    
    # Build mini OpenAPI spec for this server
    server_spec = {
        "openapi": "3.0.3",
        "info": {
            "title": f"API Explorer - {urlparse(server_url).netloc}",
            "version": "1.0.0",
            "description": f"Discovered API server from {project.name}"
        },
        "servers": [{"url": server_url, "description": "Discovered server"}],
        "paths": {
            "/": {
                "get": {
                    "summary": "Health Check",
                    "description": "Check server connectivity and authentication",
                    "operationId": "healthCheck",
                    "responses": {"200": {"description": "Server is reachable"}}
                }
            },
            "/api": {
                "get": {
                    "summary": "API Root",
                    "description": "Check API root endpoint",
                    "operationId": "apiRoot",
                    "responses": {"200": {"description": "API is available"}}
                }
            }
        },
        "components": {
            "securitySchemes": {}
        }
    }
    
    # Add appropriate security scheme
    if auth_type == 'apiKey':
        if 'azure' in auth_scheme_name.lower():
            server_spec['components']['securitySchemes'][auth_scheme_name] = {
                "type": "apiKey",
                "in": "header",
                "name": "Ocp-Apim-Subscription-Key",
                "description": "Azure API Subscription Key"
            }
        else:
            server_spec['components']['securitySchemes'][auth_scheme_name] = {
                "type": "apiKey",
                "in": "header", 
                "name": "X-API-Key",
                "description": "API Key"
            }
    elif auth_type == 'bearer':
        server_spec['components']['securitySchemes'][auth_scheme_name] = {
            "type": "http",
            "scheme": "bearer",
            "description": "Bearer Token / JWT"
        }
    elif auth_type == 'basic':
        server_spec['components']['securitySchemes'][auth_scheme_name] = {
            "type": "http",
            "scheme": "basic",
            "description": "Basic Authentication"
        }
    
    # Apply security to all paths
    if auth_scheme_name in server_spec['components']['securitySchemes']:
        server_spec['security'] = [{auth_scheme_name: []}]
    
    spec_json = json.dumps(server_spec)
    creds_json = json.dumps(matched_creds)
    
    # Generate HTML with pre-filled auth
    html_content = f'''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>API Explorer - {urlparse(server_url).netloc}</title>
    <link rel="stylesheet" href="https://unpkg.com/swagger-ui-dist@5.9.0/swagger-ui.css" />
    <style>
        body {{ margin: 0; padding: 0; background: #1a1a2e; }}
        .header {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 16px 30px;
        }}
        .header h1 {{ margin: 0; font-size: 1.3rem; font-family: system-ui; }}
        .header .server {{ opacity: 0.85; font-size: 0.85rem; font-family: monospace; margin-top: 4px; }}
        .auth-bar {{
            background: #22223a;
            padding: 12px 30px;
            border-bottom: 1px solid #3a3a6e;
            font-family: system-ui;
        }}
        .auth-bar label {{ color: #a0a0c0; font-size: 12px; display: block; margin-bottom: 4px; }}
        .auth-bar .auth-row {{ display: flex; gap: 12px; align-items: flex-end; flex-wrap: wrap; }}
        .auth-field {{ }}
        .auth-field input, .auth-field select {{
            background: #1a1a2e;
            border: 1px solid #4a4a7e;
            color: white;
            padding: 8px 12px;
            border-radius: 6px;
            font-size: 13px;
            min-width: 200px;
        }}
        .auth-field input:focus {{ outline: none; border-color: #667eea; }}
        .auth-btn {{
            background: linear-gradient(135deg, #10b981 0%, #059669 100%);
            color: white;
            border: none;
            padding: 8px 20px;
            border-radius: 6px;
            cursor: pointer;
            font-weight: 600;
            font-size: 13px;
        }}
        .auth-status {{
            color: #10b981;
            font-size: 12px;
            padding: 8px 12px;
        }}
        .auth-status.warning {{ color: #f59e0b; }}
        .swagger-ui .topbar {{ display: none; }}
        .swagger-ui {{ max-width: 1400px; margin: 0 auto; }}
        .back-link {{
            color: rgba(255,255,255,0.7);
            text-decoration: none;
            font-size: 13px;
            margin-left: 20px;
        }}
        .back-link:hover {{ color: white; }}
    </style>
</head>
<body>
    <div class="header">
        <h1>
            🔍 API Explorer
            <a href="javascript:history.back()" class="back-link">← Back</a>
        </h1>
        <div class="server">{server_url}</div>
    </div>
    
    <div class="auth-bar">
        <div class="auth-row">
            <div class="auth-field">
                <label>Authentication Type</label>
                <select id="authType" onchange="updateAuthFields()">
                    <option value="apiKey" {"selected" if auth_type == "apiKey" else ""}>API Key (Header)</option>
                    <option value="bearer" {"selected" if auth_type == "bearer" else ""}>Bearer Token / JWT</option>
                    <option value="basic" {"selected" if auth_type == "basic" else ""}>Basic Auth</option>
                </select>
            </div>
            
            <div class="auth-field" id="apiKeyField" style="display: {"block" if auth_type == "apiKey" else "none"}">
                <label>API Key Value</label>
                <input type="text" id="apiKeyValue" value="{auth_value if auth_type == 'apiKey' else ''}" placeholder="Enter API key..." />
            </div>
            
            <div class="auth-field" id="bearerField" style="display: {"block" if auth_type == "bearer" else "none"}">
                <label>Bearer Token</label>
                <input type="text" id="bearerValue" value="{auth_value if auth_type == 'bearer' else ''}" placeholder="Enter token..." />
            </div>
            
            <div class="auth-field" id="basicUserField" style="display: {"block" if auth_type == "basic" else "none"}">
                <label>Username</label>
                <input type="text" id="basicUsername" value="{auth_value if auth_type == 'basic' else ''}" placeholder="Username..." />
            </div>
            
            <div class="auth-field" id="basicPassField" style="display: {"block" if auth_type == "basic" else "none"}">
                <label>Password</label>
                <input type="password" id="basicPassword" placeholder="Password..." />
            </div>
            
            <button class="auth-btn" onclick="applyAuth()">🔑 Apply Authorization</button>
            <span class="auth-status" id="authStatus">{"✓ Pre-filled with discovered credential" if auth_value else "⚠ No credential discovered"}</span>
        </div>
    </div>
    
    <div id="swagger-ui"></div>
    <script src="https://unpkg.com/swagger-ui-dist@5.9.0/swagger-ui-bundle.js"></script>
    <script>
        const spec = {spec_json};
        const credentials = {creds_json};
        let swaggerUI = null;
        
        function updateAuthFields() {{
            const authType = document.getElementById('authType').value;
            document.getElementById('apiKeyField').style.display = authType === 'apiKey' ? 'block' : 'none';
            document.getElementById('bearerField').style.display = authType === 'bearer' ? 'block' : 'none';
            document.getElementById('basicUserField').style.display = authType === 'basic' ? 'block' : 'none';
            document.getElementById('basicPassField').style.display = authType === 'basic' ? 'block' : 'none';
        }}
        
        function applyAuth() {{
            const status = document.getElementById('authStatus');
            status.textContent = '✓ Authorization applied';
            status.className = 'auth-status';
            
            // Re-init SwaggerUI with new auth
            initSwaggerUI();
        }}
        
        function initSwaggerUI() {{
            const authType = document.getElementById('authType').value;
            let authHeader = '';
            
            if (authType === 'apiKey') {{
                authHeader = document.getElementById('apiKeyValue').value;
            }} else if (authType === 'bearer') {{
                authHeader = 'Bearer ' + document.getElementById('bearerValue').value;
            }} else if (authType === 'basic') {{
                const user = document.getElementById('basicUsername').value;
                const pass = document.getElementById('basicPassword').value;
                authHeader = 'Basic ' + btoa(user + ':' + pass);
            }}
            
            if (document.getElementById('swagger-ui')) {{
                document.getElementById('swagger-ui').innerHTML = '';
            }}
            
            swaggerUI = SwaggerUIBundle({{
                spec: spec,
                dom_id: '#swagger-ui',
                deepLinking: true,
                presets: [
                    SwaggerUIBundle.presets.apis,
                    SwaggerUIBundle.SwaggerUIStandalonePreset
                ],
                layout: "BaseLayout",
                defaultModelsExpandDepth: -1,
                docExpansion: "list",
                tryItOutEnabled: true,
                
                requestInterceptor: function(request) {{
                    if (authHeader) {{
                        const authType = document.getElementById('authType').value;
                        if (authType === 'apiKey') {{
                            // Detect Azure vs generic API key
                            if ('{server_url}'.includes('azure')) {{
                                request.headers['Ocp-Apim-Subscription-Key'] = authHeader;
                            }} else {{
                                request.headers['X-API-Key'] = authHeader;
                            }}
                        }} else {{
                            request.headers['Authorization'] = authHeader;
                        }}
                    }}
                    return request;
                }}
            }});
        }}
        
        window.onload = initSwaggerUI;
    </script>
</body>
</html>'''
    
    return HTMLResponse(content=html_content)


@router.post("/{project_id}/api-audit/spider-openapi", dependencies=[Depends(require_permissions("admin:manage"))],
    summary="Spider all servers for OpenAPI specs",
    responses={401: {"description": "Not authenticated"}, 403: {"description": "Insufficient permissions - requires admin:manage"}, 404: {"description": "Project not found"}, 500: {"description": "Spidering error"}})
async def spider_openapi_specs(project_id: str, db: Session = Depends(get_tenant_db)):
    """
    Spider all discovered servers for OpenAPI specifications.
    Updates the project's OpenAPI spec with discovered paths and methods.
    Requires admin:manage permission.
    """
    project = db.query(models.Repository).filter(
        models.Repository.id == project_id
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # Load data
    reports_dir = os.environ.get("REPORTS_DIR", "/app/vulnerability_reports")
    endpoints_file = os.path.join(reports_dir, project.name, f"{project.name}_api_endpoints.json")
    openapi_file = os.path.join(reports_dir, project.name, f"{project.name}_openapi.yaml")
    openapi_json = os.path.join(reports_dir, project.name, f"{project.name}_openapi.json")
    
    servers = []
    credentials = []
    
    # Load servers from OpenAPI
    base_spec = {"openapi": "3.0.3", "info": {"title": f"{project.name} API", "version": "1.0.0"}, "paths": {}, "servers": []}
    if os.path.exists(openapi_file):
        try:
            with open(openapi_file, 'r') as f:
                base_spec = yaml.safe_load(f)
                for s in base_spec.get('servers', []):
                    servers.append(s.get('url', ''))
        except:
            pass
    
    # Load credentials
    if os.path.exists(endpoints_file):
        try:
            with open(endpoints_file, 'r') as f:
                data = json.load(f)
            
            for ep in data.get('outbound_endpoints', []):
                secret_type = ep.get('metadata', {}).get('secret_type', '')
                if not secret_type or secret_type == 'api_url':
                    continue
                
                code = ep.get('code', '')
                value = ''
                if '=' in code:
                    parts = code.split('=', 1)
                    if len(parts) > 1:
                        value = parts[1].strip().strip('"').strip("'")
                
                if value:
                    credentials.append({
                        'type': secret_type,
                        'value': value[:100],
                        'environment': ep.get('metadata', {}).get('environment', 'unknown'),
                    })
        except:
            pass
    
    # Build credentials map
    creds_map = {}
    for server in servers:
        matched = [c for c in credentials if _match_credential_to_server(c, server)]
        if matched:
            creds_map[server] = matched[0]
    
    # Spider servers (limit to first 5 for speed)
    discovered_specs = {}
    import httpx
    
    async with httpx.AsyncClient(timeout=10.0, follow_redirects=True) as client:
        for server_url in servers[:5]:
            creds = creds_map.get(server_url)
            headers = {"User-Agent": "AuditGH-Spider/1.0"}
            
            if creds:
                cred_type = creds.get('type', '').lower()
                cred_value = creds.get('value', '')
                if 'azure' in cred_type:
                    headers['Ocp-Apim-Subscription-Key'] = cred_value
                elif 'bearer' in cred_type or 'token' in cred_type:
                    headers['Authorization'] = f'Bearer {cred_value}'
                else:
                    headers['X-API-Key'] = cred_value
            
            # Probe common OpenAPI paths
            spider_paths = ['/swagger.json', '/openapi.json', '/api-docs', '/v3/api-docs', '/v2/api-docs']
            for path in spider_paths:
                try:
                    url = server_url.rstrip('/') + path
                    resp = await client.get(url, headers=headers)
                    if resp.status_code == 200:
                        try:
                            spec = resp.json()
                            if 'paths' in spec or 'openapi' in spec or 'swagger' in spec:
                                discovered_specs[server_url] = {
                                    'source_url': url,
                                    'spec': spec
                                }
                                break  # Found spec for this server
                        except:
                            pass
                except:
                    pass
    
    # Merge discovered specs
    merged_paths_count = 0
    for server_url, spec_data in discovered_specs.items():
        spec = spec_data.get('spec', {})
        for path, methods in spec.get('paths', {}).items():
            if path not in base_spec.get('paths', {}):
                base_spec['paths'][path] = methods
                merged_paths_count += len([m for m in methods.keys() if m.lower() in ['get', 'post', 'put', 'delete', 'patch']])
    
    # Save updated spec
    try:
        with open(openapi_file, 'w') as f:
            yaml.dump(base_spec, f, default_flow_style=False)
        with open(openapi_json, 'w') as f:
            json.dump(base_spec, f, indent=2)
    except Exception as e:
        return {"success": False, "error": str(e)}
    
    return {
        "success": True,
        "servers_spidered": len(servers[:5]),
        "specs_found": len(discovered_specs),
        "paths_discovered": merged_paths_count,
        "discovered_from": list(discovered_specs.keys())
    }


from urllib.parse import urlparse
from src.rbac.dependencies import require_permissions


@router.get("/{project_id}/api-audit/openapi/view", response_model=OpenAPISpecResponse, dependencies=[Depends(require_permissions("admin:manage"))],
    summary="View OpenAPI specification content",
    responses={401: {"description": "Not authenticated"}, 403: {"description": "Insufficient permissions - requires admin:manage"}, 404: {"description": "Project or OpenAPI specification not found"}})
def view_project_openapi_spec(project_id: str, db: Session = Depends(get_tenant_db)):
    """
    Get the OpenAPI specification content for viewing in the UI.
    Returns the raw spec content, format, version, and endpoint count.
    Requires admin:manage permission.
    """
    project = db.query(models.Repository).filter(
        models.Repository.id == project_id
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    spec = db.query(models.OpenAPISpec).filter(
        models.OpenAPISpec.repository_id == project_id
    ).first()
    
    if not spec:
        raise HTTPException(status_code=404, detail="OpenAPI specification not found")
    
    return OpenAPISpecResponse(
        spec_content=spec.spec_content,
        spec_format=spec.spec_format,
        version=spec.version,
        endpoint_count=spec.endpoint_count,
        generated_at=spec.generated_at
    )


# =============================================================================
# Full Report Endpoint - Returns complete audit data from JSON files
# =============================================================================

import json
import os
import yaml

@router.get("/{project_id}/api-audit/full-report", dependencies=[Depends(require_permissions("admin:manage"))],
    summary="Get complete API audit report",
    responses={401: {"description": "Not authenticated"}, 403: {"description": "Insufficient permissions - requires admin:manage"}, 404: {"description": "Project not found"}})
def get_project_full_audit_report(project_id: str, db: Session = Depends(get_tenant_db)):
    """
    Get the complete API audit report data including inbound/outbound endpoints,
    discovered servers, credential risk assessment, and fingerprint data.
    Requires admin:manage permission.
    """
    project = db.query(models.Repository).filter(
        models.Repository.id == project_id
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # Try to load from vulnerability_reports JSON file
    reports_dir = os.environ.get("REPORTS_DIR", "/app/vulnerability_reports")
    repo_name = project.name
    
    endpoints_file = os.path.join(reports_dir, repo_name, f"{repo_name}_api_endpoints.json")
    openapi_file = os.path.join(reports_dir, repo_name, f"{repo_name}_openapi.yaml")
    fingerprint_file = os.path.join(reports_dir, repo_name, f"{repo_name}_fingerprint.json")
    
    result = {
        "repository": repo_name,
        "timestamp": None,
        "inbound_endpoints": [],
        "outbound_endpoints": [],
        "auth_patterns": [],
        "fingerprint": {
            "language": None,
            "frameworks": [],
            "http_clients": [],
            "config_sources": []
        },
        "servers": [],
        "credentials": {
            "high": [],
            "medium": [],
            "low": []
        }
    }
    
    # Load endpoints JSON
    if os.path.exists(endpoints_file):
        try:
            with open(endpoints_file, 'r') as f:
                data = json.load(f)
                result["timestamp"] = data.get("timestamp")
                result["inbound_endpoints"] = data.get("inbound_endpoints", [])
                result["outbound_endpoints"] = data.get("outbound_endpoints", [])
                result["auth_patterns"] = data.get("auth_patterns", [])
                
                # Classify credentials by risk
                for ep in result["outbound_endpoints"]:
                    secret_type = ep.get("metadata", {}).get("secret_type", "")
                    if not secret_type or secret_type == "api_url":
                        continue
                    
                    cred = {
                        "type": secret_type,
                        "environment": ep.get("metadata", {}).get("environment", "unknown"),
                        "file": os.path.basename(ep.get("path", "")),
                        "code": ep.get("code", "")[:80],
                        "attack_vector": None
                    }
                    
                    # Classify by risk level
                    if any(x in secret_type.lower() for x in ['azure_key', 'shared_key', 'secret', 'signature']):
                        cred["attack_vector"] = "Infrastructure access, service impersonation"
                        result["credentials"]["high"].append(cred)
                    elif any(x in secret_type.lower() for x in ['mixpanel', 'firebase', 'instabug', 'api_key']):
                        if 'mixpanel' in secret_type.lower():
                            cred["attack_vector"] = "Send fake analytics, impersonate users"
                        elif 'firebase' in secret_type.lower():
                            cred["attack_vector"] = "Push notification abuse, Firebase access"
                        elif 'instabug' in secret_type.lower():
                            cred["attack_vector"] = "Submit fake bug reports"
                        else:
                            cred["attack_vector"] = "API abuse, data injection"
                        result["credentials"]["medium"].append(cred)
                    elif any(x in secret_type.lower() for x in ['client_id', 'cognito']):
                        cred["attack_vector"] = "API reconnaissance"
                        result["credentials"]["low"].append(cred)
        except Exception as e:
            print(f"Error loading endpoints file: {e}")
    else:
        # FALLBACK: Load from database if files don't exist
        endpoints = db.query(models.APIEndpoint).filter(
            models.APIEndpoint.repository_id == project_id
        ).all()
        
        for ep in endpoints:
            endpoint_data = {
                "category": "db",
                "rule_id": "database",
                "path": ep.file_path or "",
                "line": ep.line_number or 0,
                "code": ep.code_snippet or ep.endpoint_url,
                "endpoint_path": ep.endpoint_url,
                "message": f"From database: {ep.framework or 'unknown'}",
                "metadata": {
                    "category": "api-discovery",
                    "subcategory": ep.direction,
                    "framework": ep.framework
                }
            }
            if ep.direction == 'serves':
                result["inbound_endpoints"].append(endpoint_data)
            else:
                result["outbound_endpoints"].append(endpoint_data)
    
    # Load OpenAPI spec for servers
    if os.path.exists(openapi_file):
        try:
            with open(openapi_file, 'r') as f:
                spec = yaml.safe_load(f)
                result["servers"] = spec.get("servers", [])
        except Exception as e:
            print(f"Error loading OpenAPI file: {e}")
    else:
        # FALLBACK: Load from database
        spec = db.query(models.OpenAPISpec).filter(
            models.OpenAPISpec.repository_id == project_id
        ).first()
        if spec:
            try:
                spec_data = yaml.safe_load(spec.spec_content) if spec.spec_format == 'yaml' else json.loads(spec.spec_content)
                result["servers"] = spec_data.get("servers", [])
            except:
                pass
    
    # Load fingerprint
    if os.path.exists(fingerprint_file):
        try:
            with open(fingerprint_file, 'r') as f:
                fp = json.load(f)
                result["fingerprint"] = fp
        except Exception as e:
            print(f"Error loading fingerprint file: {e}")
    
    return result


# =============================================================================
# Global Config Endpoints (Path Dictionary & URI Library)
# =============================================================================

@global_router.get("/dictionary", response_model=List[PathDictionaryItem],
    summary="List API path dictionary words",
    responses={401: {"description": "Not authenticated"}})
def get_path_dictionary(db: Session = Depends(get_tenant_db)):
    """
    List all words in the API path dictionary used for endpoint discovery.
    Words are returned sorted alphabetically.
    """
    return db.query(models.APIPathDictionary).order_by(models.APIPathDictionary.word).all()

@global_router.post("/dictionary", response_model=PathDictionaryItem,
    summary="Add word to API path dictionary",
    responses={401: {"description": "Not authenticated"}, 400: {"description": "Word already exists in dictionary"}})
def add_path_dictionary_item(item: CreatePathDictionaryItem, db: Session = Depends(get_tenant_db)):
    """
    Add a new word to the API path dictionary for endpoint discovery.
    Rejects duplicates if the word already exists.
    """
    # Check if exists
    existing = db.query(models.APIPathDictionary).filter(
        models.APIPathDictionary.word == item.word
    ).first()
    
    if existing:
        raise HTTPException(status_code=400, detail="Word already exists in dictionary")
    
    new_item = models.APIPathDictionary(
        word=item.word,
        category=item.category,
        is_active=True
    )
    db.add(new_item)
    db.commit()
    db.refresh(new_item)
    return new_item

@global_router.delete("/dictionary/{word}",
    summary="Delete word from API path dictionary",
    responses={401: {"description": "Not authenticated"}, 404: {"description": "Word not found"}})
def delete_path_dictionary_item(word: str, db: Session = Depends(get_tenant_db)):
    """
    Delete a word from the API path dictionary.
    Returns 404 if the word does not exist.
    """
    item = db.query(models.APIPathDictionary).filter(
        models.APIPathDictionary.word == word
    ).first()
    
    if not item:
        raise HTTPException(status_code=404, detail="Word not found")
    
    db.delete(item)
    db.commit()
    return {"success": True, "message": "Word deleted"}

@global_router.get("/uri-library", response_model=List[URILibraryItem],
    summary="List URI library entries",
    responses={401: {"description": "Not authenticated"}})
def get_uri_library(db: Session = Depends(get_tenant_db)):
    """
    List all URIs in the library used for API audit enrichment.
    URIs are returned sorted alphabetically.
    """
    return db.query(models.APIURILibrary).order_by(models.APIURILibrary.uri).all()

@global_router.post("/uri-library", response_model=URILibraryItem,
    summary="Add URI to the library",
    responses={401: {"description": "Not authenticated"}, 400: {"description": "URI already exists in library"}})
def add_uri_library_item(item: CreateURILibraryItem, db: Session = Depends(get_tenant_db)):
    """
    Add a new URI to the library for API audit enrichment.
    Rejects duplicates if the URI already exists.
    """
    # Check if exists
    existing = db.query(models.APIURILibrary).filter(
        models.APIURILibrary.uri == item.uri
    ).first()
    
    if existing:
        raise HTTPException(status_code=400, detail="URI already exists in library")
    
    new_item = models.APIURILibrary(
        uri=item.uri,
        description=item.description,
        source=item.source,
        is_active=True
    )
    db.add(new_item)
    db.commit()
    db.refresh(new_item)
    return new_item

@global_router.delete("/uri-library/{item_id}",
    summary="Delete URI from the library",
    responses={401: {"description": "Not authenticated"}, 404: {"description": "URI not found"}})
def delete_uri_library_item(item_id: str, db: Session = Depends(get_tenant_db)):
    """
    Delete a URI from the library by its ID.
    Returns 404 if the URI does not exist.
    """
    item = db.query(models.APIURILibrary).filter(
        models.APIURILibrary.id == item_id
    ).first()
    
    if not item:
        raise HTTPException(status_code=404, detail="URI not found")
    
    db.delete(item)
    db.commit()
    return {"success": True, "message": "URI deleted"}


# =============================================================================
# Credential Testing UI - Test authentication against discovered API servers
# =============================================================================

@router.get("/{project_id}/api-audit/credential-testing", response_class=HTMLResponse, dependencies=[Depends(require_permissions("admin:manage"))],
    summary="Serve credential testing UI",
    responses={401: {"description": "Not authenticated"}, 403: {"description": "Insufficient permissions - requires admin:manage"}, 404: {"description": "Project not found"}})
async def get_credential_testing_ui(project_id: str, db: Session = Depends(get_tenant_db)):
    """
    Credential Testing UI for testing authentication against discovered API servers
    using matched credentials from the codebase.
    Requires admin:manage permission.
    """
    project = db.query(models.Repository).filter(
        models.Repository.id == project_id
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # Get swagger server credentials mapping
    reports_dir = os.environ.get("REPORTS_DIR", "/app/vulnerability_reports")
    
    try:
        from execution.ai_credential_matcher import map_credentials_to_swagger_servers
        mappings = await map_credentials_to_swagger_servers(project.name, reports_dir)
    except Exception as e:
        logger.bind(router="api_audit", endpoint="get_swagger_server_suggestions").exception(f"Failed to get credential mappings: {e}")
        mappings = []
    
    mappings_json = json.dumps(mappings[:50])  # Limit to 50 servers
    
    html_content = f'''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Credential Testing - {project.name}</title>
    <style>
        * {{ box-sizing: border-box; }}
        body {{
            margin: 0;
            padding: 0;
            background: #0f0f1e;
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
            color: #e0e0e0;
        }}
        .header {{
            background: linear-gradient(135deg, #10b981 0%, #059669 100%);
            color: white;
            padding: 20px 40px;
            display: flex;
            justify-content: space-between;
            align-items: center;
        }}
        .header h1 {{ margin: 0; font-size: 1.4rem; }}
        .header .subtitle {{ opacity: 0.85; font-size: 0.85rem; margin-top: 4px; }}
        .back-btn {{
            background: rgba(255,255,255,0.2);
            color: white;
            padding: 8px 16px;
            border-radius: 8px;
            text-decoration: none;
            font-size: 13px;
        }}
        .container {{ padding: 30px 40px; max-width: 1600px; margin: 0 auto; }}
        
        .server-card {{
            background: #1a1a2e;
            border-radius: 12px;
            margin-bottom: 20px;
            overflow: hidden;
            border: 1px solid #2a2a4e;
        }}
        .server-header {{
            background: #22223a;
            padding: 16px 20px;
            display: flex;
            justify-content: space-between;
            align-items: center;
            cursor: pointer;
        }}
        .server-header:hover {{ background: #2a2a4e; }}
        .server-url {{
            font-family: monospace;
            font-size: 14px;
            color: #60a5fa;
        }}
        .server-env {{
            background: #4a4a7e;
            color: white;
            padding: 4px 10px;
            border-radius: 4px;
            font-size: 11px;
            text-transform: uppercase;
        }}
        .server-body {{
            padding: 20px;
            display: none;
        }}
        .server-body.expanded {{ display: block; }}
        
        .cred-table {{
            width: 100%;
            border-collapse: collapse;
            font-size: 13px;
        }}
        .cred-table th {{
            text-align: left;
            padding: 10px 12px;
            background: #0f0f1e;
            color: #a0a0c0;
            font-weight: 600;
            font-size: 11px;
            text-transform: uppercase;
        }}
        .cred-table td {{
            padding: 10px 12px;
            border-bottom: 1px solid #2a2a4e;
        }}
        .cred-table tr:last-child td {{ border-bottom: none; }}
        
        .cred-type {{
            background: #3b82f6;
            color: white;
            padding: 3px 8px;
            border-radius: 4px;
            font-size: 11px;
            font-weight: 600;
        }}
        .cred-value {{
            font-family: monospace;
            font-size: 12px;
            background: #0f0f1e;
            padding: 4px 8px;
            border-radius: 4px;
            max-width: 300px;
            overflow: hidden;
            text-overflow: ellipsis;
            white-space: nowrap;
        }}
        .confidence {{
            display: flex;
            align-items: center;
            gap: 6px;
        }}
        .confidence-bar {{
            width: 60px;
            height: 6px;
            background: #2a2a4e;
            border-radius: 3px;
            overflow: hidden;
        }}
        .confidence-fill {{
            height: 100%;
            border-radius: 3px;
        }}
        .confidence-high {{ background: #22c55e; }}
        .confidence-med {{ background: #f59e0b; }}
        .confidence-low {{ background: #ef4444; }}
        
        .test-btn {{
            background: linear-gradient(135deg, #8b5cf6 0%, #7c3aed 100%);
            color: white;
            border: none;
            padding: 6px 14px;
            border-radius: 6px;
            cursor: pointer;
            font-size: 12px;
            font-weight: 600;
        }}
        .test-btn:hover {{ opacity: 0.9; }}
        .test-btn:disabled {{ opacity: 0.5; cursor: not-allowed; }}
        
        .test-all-btn {{
            background: linear-gradient(135deg, #10b981 0%, #059669 100%);
            color: white;
            border: none;
            padding: 10px 24px;
            border-radius: 8px;
            cursor: pointer;
            font-size: 14px;
            font-weight: 600;
            margin-bottom: 20px;
        }}
        
        .result-badge {{
            padding: 3px 8px;
            border-radius: 4px;
            font-size: 11px;
            font-weight: 600;
        }}
        .result-success {{ background: #22c55e; color: white; }}
        .result-fail {{ background: #ef4444; color: white; }}
        .result-pending {{ background: #6b7280; color: white; }}
        
        .stats {{
            display: flex;
            gap: 30px;
            margin-bottom: 20px;
            font-size: 14px;
        }}
        .stat {{ color: #a0a0c0; }}
        .stat-value {{ color: white; font-weight: 600; font-size: 18px; }}
        
        .empty-state {{
            text-align: center;
            padding: 60px 20px;
            color: #6b7280;
        }}
        
        .expand-icon {{
            transition: transform 0.2s;
        }}
        .expanded .expand-icon {{
            transform: rotate(180deg);
        }}
    </style>
</head>
<body>
    <div class="header">
        <div>
            <h1>🔐 Credential Testing</h1>
            <div class="subtitle">{project.name} - Test authentication against discovered API servers</div>
        </div>
        <a href="javascript:history.back()" class="back-btn">← Back</a>
    </div>
    
    <div class="container">
        <div class="stats">
            <div class="stat">
                <div class="stat-value" id="serverCount">0</div>
                <div>Servers</div>
            </div>
            <div class="stat">
                <div class="stat-value" id="credCount">0</div>
                <div>Credentials</div>
            </div>
            <div class="stat">
                <div class="stat-value" id="testedCount">0</div>
                <div>Tested</div>
            </div>
        </div>
        
        <button class="test-all-btn" onclick="testAllServers()">🚀 Test All Servers</button>
        
        <div id="serverList"></div>
    </div>
    
    <script>
        const mappings = {mappings_json};
        let testResults = {{}};
        
        function updateStats() {{
            document.getElementById('serverCount').textContent = mappings.length;
            let totalCreds = 0;
            mappings.forEach(m => totalCreds += m.credentials.length);
            document.getElementById('credCount').textContent = totalCreds;
            document.getElementById('testedCount').textContent = Object.keys(testResults).length;
        }}
        
        function toggleServer(idx) {{
            const body = document.getElementById(`server-body-${{idx}}`);
            const header = document.getElementById(`server-header-${{idx}}`);
            body.classList.toggle('expanded');
            header.classList.toggle('expanded');
        }}
        
        function getConfidenceClass(conf) {{
            if (conf >= 70) return 'confidence-high';
            if (conf >= 40) return 'confidence-med';
            return 'confidence-low';
        }}
        
        function renderServers() {{
            const container = document.getElementById('serverList');
            
            if (mappings.length === 0) {{
                container.innerHTML = `
                    <div class="empty-state">
                        <h3>No servers with credentials found</h3>
                        <p>Run an API audit scan to discover servers and match credentials.</p>
                    </div>
                `;
                updateStats();
                return;
            }}
            
            container.innerHTML = mappings.map((server, idx) => `
                <div class="server-card">
                    <div class="server-header" id="server-header-${{idx}}" onclick="toggleServer(${{idx}})">
                        <div style="display: flex; align-items: center; gap: 12px;">
                            <span class="expand-icon">▼</span>
                            <span class="server-url">${{server.server_url}}</span>
                            <span class="server-env">${{server.server_environment || 'unknown'}}</span>
                        </div>
                        <div style="display: flex; align-items: center; gap: 12px;">
                            <span style="color: #a0a0c0; font-size: 12px;">${{server.credentials.length}} credentials</span>
                            <span id="server-result-${{idx}}" class="result-badge result-pending">Not Tested</span>
                        </div>
                    </div>
                    <div class="server-body" id="server-body-${{idx}}">
                        <table class="cred-table">
                            <thead>
                                <tr>
                                    <th>Type</th>
                                    <th>Value</th>
                                    <th>Environment</th>
                                    <th>Confidence</th>
                                    <th>Action</th>
                                    <th>Result</th>
                                </tr>
                            </thead>
                            <tbody>
                                ${{server.credentials.map((cred, credIdx) => `
                                    <tr>
                                        <td><span class="cred-type">${{cred.credential_type}}</span></td>
                                        <td><span class="cred-value">${{cred.credential_value}}</span></td>
                                        <td>${{cred.environment || 'unknown'}}</td>
                                        <td>
                                            <div class="confidence">
                                                <div class="confidence-bar">
                                                    <div class="confidence-fill ${{getConfidenceClass(cred.confidence)}}" style="width: ${{cred.confidence}}%"></div>
                                                </div>
                                                <span>${{cred.confidence}}%</span>
                                            </div>
                                        </td>
                                        <td>
                                            <button class="test-btn" id="test-btn-${{idx}}-${{credIdx}}" onclick="testCredential(${{idx}}, ${{credIdx}})">
                                                Test
                                            </button>
                                        </td>
                                        <td>
                                            <span id="cred-result-${{idx}}-${{credIdx}}" class="result-badge result-pending">-</span>
                                        </td>
                                    </tr>
                                `).join('')}}
                            </tbody>
                        </table>
                    </div>
                </div>
            `).join('');
            
            updateStats();
        }}
        
        async function testCredential(serverIdx, credIdx) {{
            const server = mappings[serverIdx];
            const cred = server.credentials[credIdx];
            const btn = document.getElementById(`test-btn-${{serverIdx}}-${{credIdx}}`);
            const resultSpan = document.getElementById(`cred-result-${{serverIdx}}-${{credIdx}}`);
            
            btn.disabled = true;
            btn.textContent = 'Testing...';
            resultSpan.className = 'result-badge result-pending';
            resultSpan.textContent = 'Testing...';
            
            try {{
                const response = await fetch('/projects/{project_id}/api-audit/test-credential', {{
                    method: 'POST',
                    headers: {{ 'Content-Type': 'application/json' }},
                    body: JSON.stringify({{
                        server_url: server.server_url,
                        credential_type: cred.credential_type,
                        credential_value: cred.credential_value
                    }})
                }});
                
                const result = await response.json();
                
                if (result.success && result.reachable) {{
                    resultSpan.className = 'result-badge result-success';
                    resultSpan.textContent = result.status_code ? `${{result.status_code}} OK` : 'Reachable';
                }} else {{
                    resultSpan.className = 'result-badge result-fail';
                    resultSpan.textContent = result.error || result.status_code || 'Failed';
                }}
                
                testResults[`${{serverIdx}}-${{credIdx}}`] = result;
                updateStats();
                
            }} catch (e) {{
                resultSpan.className = 'result-badge result-fail';
                resultSpan.textContent = 'Error';
            }}
            
            btn.disabled = false;
            btn.textContent = 'Test';
        }}
        
        async function testAllServers() {{
            for (let i = 0; i < mappings.length; i++) {{
                const server = mappings[i];
                if (server.credentials.length > 0) {{
                    // Test the highest confidence credential for each server
                    await testCredential(i, 0);
                    // Small delay between tests
                    await new Promise(r => setTimeout(r, 500));
                }}
            }}
        }}
        
        renderServers();
    </script>
</body>
</html>'''
    
    return HTMLResponse(content=html_content)


@router.post("/{project_id}/api-audit/test-credential", dependencies=[Depends(require_permissions("admin:manage"))],
    summary="Test a credential against a server URL",
    responses={401: {"description": "Not authenticated"}, 403: {"description": "Insufficient permissions - requires admin:manage"}, 404: {"description": "Project not found"}, 400: {"description": "Missing server_url"}, 500: {"description": "Server error or timeout"}})
async def test_credential(project_id: str, request_data: dict, db: Session = Depends(get_tenant_db)):
    """
    Test a credential against a server URL.
    Attempts to make a request to the server using the provided credential and auth type.
    Requires admin:manage permission.
    """
    import httpx
    
    project = db.query(models.Repository).filter(
        models.Repository.id == project_id
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    server_url = request_data.get('server_url', '').strip()
    credential_type = request_data.get('credential_type', '').strip()
    credential_value = request_data.get('credential_value', '').strip()
    
    if not server_url:
        return {"success": False, "error": "Missing server_url"}
    
    # Normalize URL
    if not server_url.startswith(('http://', 'https://')):
        server_url = 'https://' + server_url
    
    # Build headers based on credential type
    headers = {
        'User-Agent': 'AuditGH-CredentialTester/1.0',
        'Accept': 'application/json'
    }
    
    cred_type_lower = credential_type.lower()
    
    if 'api' in cred_type_lower and 'key' in cred_type_lower:
        # Try common API key header names
        headers['X-API-Key'] = credential_value
        headers['Api-Key'] = credential_value
    elif 'bearer' in cred_type_lower or 'token' in cred_type_lower or 'jwt' in cred_type_lower:
        headers['Authorization'] = f'Bearer {credential_value}'
    elif 'basic' in cred_type_lower:
        import base64
        encoded = base64.b64encode(credential_value.encode()).decode()
        headers['Authorization'] = f'Basic {encoded}'
    elif 'client' in cred_type_lower and 'id' in cred_type_lower:
        headers['X-Client-Id'] = credential_value
    elif 'client' in cred_type_lower and 'secret' in cred_type_lower:
        headers['X-Client-Secret'] = credential_value
    elif 'subscription' in cred_type_lower or 'ocp-apim' in cred_type_lower:
        headers['Ocp-Apim-Subscription-Key'] = credential_value
    else:
        # Generic - try as Authorization header
        headers['Authorization'] = credential_value
    
    try:
        async with httpx.AsyncClient(timeout=10.0, verify=False) as client:
            response = await client.get(server_url, headers=headers, follow_redirects=True)
            
            return {
                "success": True,
                "reachable": True,
                "status_code": response.status_code,
                "response_size": len(response.content),
                "headers_sent": list(headers.keys())
            }
    except httpx.TimeoutException:
        return {
            "success": True,
            "reachable": False,
            "error": "Timeout"
        }
    except httpx.ConnectError as e:
        return {
            "success": True,
            "reachable": False,
            "error": f"Connection failed: {str(e)[:50]}"
        }
    except Exception as e:
        return {
            "success": False,
            "reachable": False,
            "error": str(e)[:100]
        }


# =============================================================================
# AI Credential-URL Testing Agent Endpoints
# =============================================================================

class CredentialUrlTestRequest(BaseModel):
    """Request model for credential-URL testing"""
    target_url: str = Field(..., description="Target URL to test the credential against")
    credential_type: str = Field(..., description="Type of credential (bearer, api_key, azure_subscription, etc.)")
    credential_value: str = Field(..., description="Credential value to use for authentication")
    credential_environment: Optional[str] = Field("", description="Environment where the credential was found")
    confidence_score: Optional[int] = Field(0, description="AI confidence score for the credential-URL match")
    test_mode: Optional[str] = Field("cautious", description="Test mode: none, cautious, or insane")


class CredentialUrlTestAllRequest(BaseModel):
    """Request model for testing all credential-URL pairs"""
    test_mode: Optional[str] = Field("cautious", description="Test mode: none, cautious, or insane")


def _get_current_organization_id(db: Session) -> Optional[str]:
    """
    Get the current organization ID from the global context.
    Returns None if no organization is selected.
    """
    from ..database import get_current_org_id
    return get_current_org_id()


@router.get("/{project_id}/api-audit/credential-url-test-results", dependencies=[Depends(require_permissions("admin:manage"))],
    summary="Get credential-URL test results",
    responses={401: {"description": "Not authenticated"}, 403: {"description": "Insufficient permissions - requires admin:manage"}})
async def get_credential_url_test_results(
    project_id: str,
    auth_status: Optional[str] = None,
    threat_level: Optional[str] = None,
    limit: int = 100,
    db: Session = Depends(get_tenant_db)
):
    """
    Get all credential-URL test results for a project.
    Multi-tenant: Results are filtered by current organization.
    Supports filtering by auth_status and threat_level.
    Requires admin:manage permission.
    """
    org_id = _get_current_organization_id(db)
    
    query = db.query(models.CredentialUrlTestResult).filter(
        models.CredentialUrlTestResult.repository_id == project_id
    )
    
    if org_id:
        query = query.filter(models.CredentialUrlTestResult.organization_id == org_id)
    
    if auth_status:
        query = query.filter(models.CredentialUrlTestResult.auth_status == auth_status)
    
    if threat_level:
        query = query.filter(models.CredentialUrlTestResult.threat_level == threat_level)
    
    results = query.order_by(models.CredentialUrlTestResult.tested_at.desc()).limit(limit).all()
    
    return {
        "success": True,
        "results": [
            {
                "id": str(r.id),
                "target_url": r.target_url,
                "credential_type": r.credential_type,
                "credential_environment": r.credential_environment,
                "confidence_score": r.confidence_score,
                "auth_status": r.auth_status,
                "auth_status_code": r.auth_status_code,
                "auth_response_time_ms": r.auth_response_time_ms,
                "discovered_paths_count": r.discovered_paths_count,
                "hidden_paths_found": r.hidden_paths_found,
                "threat_level": r.threat_level,
                "ai_overview": r.ai_overview,
                "ai_risk_assessment": r.ai_risk_assessment,
                "ai_recommendations": r.ai_recommendations,
                "tested_at": r.tested_at.isoformat() if r.tested_at else None,
                "test_duration_seconds": r.test_duration_seconds
            }
            for r in results
        ],
        "total": len(results)
    }


@router.get("/{project_id}/api-audit/credential-url-test-results/{result_id}", dependencies=[Depends(require_permissions("admin:manage"))],
    summary="Get detailed credential-URL test result",
    responses={401: {"description": "Not authenticated"}, 403: {"description": "Insufficient permissions - requires admin:manage"}, 404: {"description": "Test result not found"}})
async def get_credential_url_test_result_detail(
    project_id: str,
    result_id: str,
    db: Session = Depends(get_tenant_db)
):
    """
    Get detailed test result for a specific credential-URL test.
    Includes full discovered paths, sample data, OSINT findings, and unmasked credentials.
    Requires admin:manage permission.
    """
    org_id = _get_current_organization_id(db)
    
    query = db.query(models.CredentialUrlTestResult).filter(
        models.CredentialUrlTestResult.id == result_id,
        models.CredentialUrlTestResult.repository_id == project_id
    )
    
    if org_id:
        query = query.filter(models.CredentialUrlTestResult.organization_id == org_id)
    
    result = query.first()
    
    if not result:
        raise HTTPException(status_code=404, detail="Test result not found")
    
    return {
        "success": True,
        "result": {
            "id": str(result.id),
            "target_url": result.target_url,
            "credential_type": result.credential_type,
            "credential_value": result.credential_value,  # Full unmasked value for security analyst validation
            "credential_environment": result.credential_environment,
            "confidence_score": result.confidence_score,
            "auth_status": result.auth_status,
            "auth_status_code": result.auth_status_code,
            "auth_response_time_ms": result.auth_response_time_ms,
            "auth_error_message": result.auth_error_message,
            "auth_headers_used": result.auth_headers_used,
            "auth_request_headers": result.auth_request_headers,  # Full request headers with actual values
            "discovered_paths": result.discovered_paths,
            "discovered_paths_count": result.discovered_paths_count,
            "hidden_paths_found": result.hidden_paths_found,
            "sample_data_retrieved": result.sample_data_retrieved,
            "data_sensitivity_indicators": result.data_sensitivity_indicators,
            "osint_findings": result.osint_findings,
            "github_repos_found": result.github_repos_found,
            "documentation_links_found": result.documentation_links_found,
            "ai_overview": result.ai_overview,
            "ai_risk_assessment": result.ai_risk_assessment,
            "ai_recommendations": result.ai_recommendations,
            "threat_level": result.threat_level,
            "test_mode": result.test_mode,
            "tested_at": result.tested_at.isoformat() if result.tested_at else None,
            "test_duration_seconds": result.test_duration_seconds,
            "llm_provider": result.llm_provider,
            "llm_model": result.llm_model
        }
    }


@router.delete("/{project_id}/api-audit/credential-url-test-results/{result_id}",
    summary="Delete a credential-URL test result",
    responses={401: {"description": "Not authenticated"}, 404: {"description": "Test result not found"}})
async def delete_credential_url_test_result(
    project_id: str,
    result_id: str,
    db: Session = Depends(get_tenant_db)
):
    """
    Delete a specific credential-URL test result by ID.
    Scoped to the current organization in multi-tenant mode.
    """
    org_id = _get_current_organization_id(db)
    
    query = db.query(models.CredentialUrlTestResult).filter(
        models.CredentialUrlTestResult.id == result_id,
        models.CredentialUrlTestResult.repository_id == project_id
    )
    
    if org_id:
        query = query.filter(models.CredentialUrlTestResult.organization_id == org_id)
    
    result = query.first()
    
    if not result:
        raise HTTPException(status_code=404, detail="Test result not found")
    
    db.delete(result)
    db.commit()
    
    return {"success": True, "message": "Test result deleted"}


@router.get("/{project_id}/api-audit/credential-url-test-status", dependencies=[Depends(require_permissions("admin:manage"))],
    summary="Get credential-URL test initialization status",
    responses={401: {"description": "Not authenticated"}, 403: {"description": "Insufficient permissions - requires admin:manage"}})
async def get_credential_url_test_status(
    project_id: str,
    db: Session = Depends(get_tenant_db)
):
    """
    Get the auto-test initialization status for a project.
    Returns whether initial testing has been completed.
    Requires admin:manage permission.
    """
    org_id = _get_current_organization_id(db)
    
    status = db.query(models.CredentialUrlTestStatus).filter(
        models.CredentialUrlTestStatus.repository_id == project_id
    )
    if org_id:
        status = status.filter(models.CredentialUrlTestStatus.organization_id == org_id)
    status = status.first()
    
    if not status:
        return {
            "initial_test_completed": False,
            "initial_test_at": None,
            "total_correlations_tested": 0,
            "total_correlations_found": 0
        }
    
    return {
        "initial_test_completed": status.initial_test_completed,
        "initial_test_at": status.initial_test_at.isoformat() if status.initial_test_at else None,
        "total_correlations_tested": status.total_correlations_tested,
        "total_correlations_found": status.total_correlations_found,
        "last_test_at": status.last_test_at.isoformat() if status.last_test_at else None
    }


@router.post("/{project_id}/api-audit/credential-url-test-status/mark-complete", dependencies=[Depends(require_permissions("admin:manage"))],
    summary="Mark initial credential-URL test as complete",
    responses={401: {"description": "Not authenticated"}, 403: {"description": "Insufficient permissions - requires admin:manage"}})
async def mark_initial_test_complete(
    project_id: str,
    total_tested: int = 0,
    total_found: int = 0,
    db: Session = Depends(get_tenant_db)
):
    """
    Mark the initial auto-test as complete for a project.
    Prevents re-testing on every page load by recording completion status.
    Requires admin:manage permission.
    """
    org_id = _get_current_organization_id(db)
    
    # Check if status record exists
    status = db.query(models.CredentialUrlTestStatus).filter(
        models.CredentialUrlTestStatus.repository_id == project_id
    )
    if org_id:
        status = status.filter(models.CredentialUrlTestStatus.organization_id == org_id)
    status = status.first()
    
    if status:
        # Update existing
        status.initial_test_completed = True
        status.initial_test_at = datetime.utcnow()
        status.total_correlations_tested = total_tested
        status.total_correlations_found = total_found
        status.last_test_at = datetime.utcnow()
    else:
        # Create new
        status = models.CredentialUrlTestStatus(
            organization_id=org_id,
            repository_id=project_id,
            initial_test_completed=True,
            initial_test_at=datetime.utcnow(),
            total_correlations_tested=total_tested,
            total_correlations_found=total_found,
            last_test_at=datetime.utcnow()
        )
        db.add(status)
    
    db.commit()
    
    return {
        "success": True,
        "initial_test_completed": True,
        "initial_test_at": status.initial_test_at.isoformat() if status.initial_test_at else None
    }


@router.delete("/{project_id}/api-audit/credential-url-reset",
    summary="Reset credential-URL test data",
    responses={401: {"description": "Not authenticated"}, 404: {"description": "Project not found"}})
async def reset_credential_url_mappings(
    project_id: str,
    db: Session = Depends(get_tenant_db)
):
    """
    Reset all credential-URL test results and status for a project.
    Deletes all test results and marks initial test as not completed,
    allowing re-running with updated matching logic.
    """
    org_id = _get_current_organization_id(db)
    
    # Verify project exists
    project = db.query(models.Repository).filter(
        models.Repository.id == project_id
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # Delete all test results for this project (scoped to org)
    results_query = db.query(models.CredentialUrlTestResult).filter(
        models.CredentialUrlTestResult.repository_id == project_id
    )
    if org_id:
        results_query = results_query.filter(
            models.CredentialUrlTestResult.organization_id == org_id
        )
    deleted_results = results_query.delete(synchronize_session=False)
    
    # Reset the test status
    status_query = db.query(models.CredentialUrlTestStatus).filter(
        models.CredentialUrlTestStatus.repository_id == project_id
    )
    if org_id:
        status_query = status_query.filter(
            models.CredentialUrlTestStatus.organization_id == org_id
        )
    status = status_query.first()
    
    if status:
        status.initial_test_completed = False
        status.initial_test_at = None
        status.total_correlations_tested = 0
        status.total_correlations_found = 0
        status.last_test_at = None
    
    db.commit()
    
    return {
        "success": True,
        "message": f"Reset credential-URL mappings for project {project.name}",
        "deleted_results": deleted_results,
        "status_reset": status is not None,
        "next_steps": [
            "Refresh the API Audit page",
            "The system will re-fetch correlations using updated matching logic",
            "Auto-test will run against the correct service-specific endpoints"
        ]
    }


@router.post("/{project_id}/api-audit/credential-url-test", dependencies=[Depends(require_permissions("admin:manage"))],
    summary="Test a single credential-URL pair with AI agent",
    responses={401: {"description": "Not authenticated"}, 403: {"description": "Insufficient permissions - requires admin:manage"}, 404: {"description": "Project not found"}, 500: {"description": "AI agent error or test failure"}})
async def test_credential_url_endpoint(
    project_id: str,
    request_data: CredentialUrlTestRequest,
    db: Session = Depends(get_tenant_db)
):
    """
    Test a single credential-URL pair using the AI Agent.
    The agent authenticates, discovers paths, analyzes data, gathers OSINT,
    and generates AI-powered risk assessment and recommendations.
    Requires admin:manage permission.
    """
    from execution.ai_credential_url_agent import test_credential_url
    
    project = db.query(models.Repository).filter(
        models.Repository.id == project_id
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # Get current organization for multi-tenant scoping
    org_id = _get_current_organization_id(db)
    
    # Validate test mode
    if request_data.test_mode not in ('none', 'cautious', 'insane'):
        raise HTTPException(status_code=400, detail="Invalid test_mode. Must be: none, cautious, or insane")
    
    try:
        # Run the AI agent
        github_token = os.environ.get("GITHUB_TOKEN", "")
        result = await test_credential_url(
            target_url=request_data.target_url,
            credential_type=request_data.credential_type,
            credential_value=request_data.credential_value,
            test_mode=request_data.test_mode,
            credential_environment=request_data.credential_environment or "",
            confidence_score=request_data.confidence_score or 0,
            github_token=github_token
        )
        
        # Check if result already exists for this URL (scoped to org)
        existing_query = db.query(models.CredentialUrlTestResult).filter(
            models.CredentialUrlTestResult.repository_id == project_id,
            models.CredentialUrlTestResult.target_url == request_data.target_url
        )
        if org_id:
            existing_query = existing_query.filter(
                models.CredentialUrlTestResult.organization_id == org_id
            )
        existing = existing_query.first()
        
        if existing:
            # Update existing record
            for key, value in result.items():
                if hasattr(existing, key) and key not in ('id', 'api_id', 'organization_id', 'repository_id', 'created_at'):
                    setattr(existing, key, value)
            existing.credential_value = request_data.credential_value  # Store actual value
            db.commit()
            db.refresh(existing)
            result_id = str(existing.id)
        else:
            # Create new record with organization scope
            new_result = models.CredentialUrlTestResult(
                organization_id=org_id,
                repository_id=project_id,
                target_url=result.get('target_url'),
                credential_type=result.get('credential_type'),
                credential_value=request_data.credential_value,  # Store actual value
                credential_environment=result.get('credential_environment'),
                confidence_score=result.get('confidence_score'),
                auth_status=result.get('auth_status'),
                auth_status_code=result.get('auth_status_code'),
                auth_response_time_ms=result.get('auth_response_time_ms'),
                auth_error_message=result.get('auth_error_message'),
                auth_headers_used=result.get('auth_headers_used', []),
                auth_request_method=result.get('auth_request_method', 'GET'),
                auth_request_url=result.get('auth_request_url'),
                auth_request_headers=result.get('auth_request_headers', {}),
                auth_request_body=result.get('auth_request_body', ''),
                auth_response_headers=result.get('auth_response_headers', {}),
                auth_response_body=result.get('auth_response_body', ''),
                auth_response_body_truncated=result.get('auth_response_body_truncated', False),
                detected_service=result.get('detected_service'),
                service_detection_score=result.get('service_detection_score', 0),
                discovered_paths=result.get('discovered_paths', []),
                discovered_paths_count=result.get('discovered_paths_count', 0),
                hidden_paths_found=result.get('hidden_paths_found', 0),
                sample_data_retrieved=result.get('sample_data_retrieved', []),
                data_sensitivity_indicators=result.get('data_sensitivity_indicators', []),
                osint_findings=result.get('osint_findings', []),
                github_repos_found=result.get('github_repos_found', 0),
                documentation_links_found=result.get('documentation_links_found', 0),
                ai_overview=result.get('ai_overview'),
                ai_risk_assessment=result.get('ai_risk_assessment'),
                ai_recommendations=result.get('ai_recommendations', []),
                threat_level=result.get('threat_level'),
                test_mode=result.get('test_mode'),
                tested_at=datetime.fromisoformat(result.get('tested_at')) if result.get('tested_at') else datetime.utcnow(),
                test_duration_seconds=result.get('test_duration_seconds'),
                llm_provider=result.get('llm_provider'),
                llm_model=result.get('llm_model'),
                raw_llm_responses=result.get('raw_llm_responses', [])
            )
            db.add(new_result)
            db.commit()
            db.refresh(new_result)
            result_id = str(new_result.id)
        
        return {
            "success": True,
            "result_id": result_id,
            "auth_status": result.get('auth_status'),
            "discovered_paths_count": result.get('discovered_paths_count'),
            "threat_level": result.get('threat_level'),
            "test_duration_seconds": result.get('test_duration_seconds')
        }
        
    except Exception as e:
        logger.bind(router="api_audit", endpoint="test_credential_url").exception(f"Credential-URL test failed: {e}")
        return {
            "success": False,
            "error": str(e)
        }


@router.post("/{project_id}/api-audit/credential-url-test-all", dependencies=[Depends(require_permissions("admin:manage"))],
    summary="Test all credential-URL pairs with AI agent",
    responses={401: {"description": "Not authenticated"}, 403: {"description": "Insufficient permissions - requires admin:manage"}, 404: {"description": "Project not found"}, 500: {"description": "AI agent or correlation error"}})
async def test_all_credential_urls(
    project_id: str,
    request_data: CredentialUrlTestAllRequest,
    db: Session = Depends(get_tenant_db)
):
    """
    Test all credential-URL pairs for a project using the AI Agent.
    Uses an OSINT-first approach to pre-validate URLs and skip public endpoints.
    Results are stored in the database and scoped to the current organization.
    Requires admin:manage permission.
    """
    from execution.ai_credential_matcher import correlate_credentials_to_urls_v2
    from execution.ai_credential_url_agent import test_credential_url
    
    project = db.query(models.Repository).filter(
        models.Repository.id == project_id
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # Get current organization for multi-tenant scoping
    org_id = _get_current_organization_id(db)
    
    # Get all credential-URL correlations using v2 with pre-validation
    reports_dir = os.environ.get("REPORTS_DIR", "/app/vulnerability_reports")
    
    try:
        logger.bind(router="api_audit", endpoint="test_all_credential_urls").info(f"[test_all] Getting correlations with URL pre-validation...")
        correlations = await correlate_credentials_to_urls_v2(project.name, reports_dir, pre_validate=True)

        # Filter to only AUTH_REQUIRED URLs
        auth_required_correlations = [c for c in correlations if c.get('requires_auth', True)]
        public_skipped = len(correlations) - len(auth_required_correlations)

        if public_skipped > 0:
            logger.bind(router="api_audit", endpoint="test_all_credential_urls").info(f"[test_all] Skipped {public_skipped} public URLs that don't require auth")

        correlations = auth_required_correlations
    except Exception as e:
        logger.bind(router="api_audit", endpoint="test_all_credential_urls").exception(f"Failed to get correlations: {e}")
        return {"success": False, "error": f"Failed to get correlations: {str(e)}"}
    
    if not correlations:
        return {"success": True, "message": "No credential-URL correlations found", "tested": 0}
    
    # Test each correlation
    results = []
    github_token = os.environ.get("GITHUB_TOKEN", "")
    
    for corr in correlations[:50]:  # Limit to 50 to prevent timeout
        try:
            result = await test_credential_url(
                target_url=corr.get('url', ''),
                credential_type=corr.get('credential', {}).get('type', ''),
                credential_value=corr.get('credential', {}).get('value', ''),
                test_mode=request_data.test_mode,
                credential_environment=corr.get('credential', {}).get('environment', ''),
                confidence_score=corr.get('confidence', 0),
                github_token=github_token
            )
            
            # Store result (scoped to org)
            existing_query = db.query(models.CredentialUrlTestResult).filter(
                models.CredentialUrlTestResult.repository_id == project_id,
                models.CredentialUrlTestResult.target_url == corr.get('url', '')
            )
            if org_id:
                existing_query = existing_query.filter(
                    models.CredentialUrlTestResult.organization_id == org_id
                )
            existing = existing_query.first()
            
            if existing:
                for key, value in result.items():
                    if hasattr(existing, key) and key not in ('id', 'api_id', 'organization_id', 'repository_id', 'created_at'):
                        setattr(existing, key, value)
                existing.credential_value = corr.get('credential', {}).get('value', '')
            else:
                new_result = models.CredentialUrlTestResult(
                    organization_id=org_id,
                    repository_id=project_id,
                    target_url=result.get('target_url'),
                    credential_type=result.get('credential_type'),
                    credential_value=corr.get('credential', {}).get('value', ''),
                    credential_environment=result.get('credential_environment'),
                    confidence_score=result.get('confidence_score'),
                    auth_status=result.get('auth_status'),
                    auth_status_code=result.get('auth_status_code'),
                    auth_response_time_ms=result.get('auth_response_time_ms'),
                    auth_error_message=result.get('auth_error_message'),
                    auth_headers_used=result.get('auth_headers_used', []),
                    discovered_paths=result.get('discovered_paths', []),
                    discovered_paths_count=result.get('discovered_paths_count', 0),
                    hidden_paths_found=result.get('hidden_paths_found', 0),
                    sample_data_retrieved=result.get('sample_data_retrieved', []),
                    data_sensitivity_indicators=result.get('data_sensitivity_indicators', []),
                    osint_findings=result.get('osint_findings', []),
                    github_repos_found=result.get('github_repos_found', 0),
                    documentation_links_found=result.get('documentation_links_found', 0),
                    ai_overview=result.get('ai_overview'),
                    ai_risk_assessment=result.get('ai_risk_assessment'),
                    ai_recommendations=result.get('ai_recommendations', []),
                    threat_level=result.get('threat_level'),
                    test_mode=result.get('test_mode'),
                    tested_at=datetime.fromisoformat(result.get('tested_at')) if result.get('tested_at') else datetime.utcnow(),
                    test_duration_seconds=result.get('test_duration_seconds'),
                    llm_provider=result.get('llm_provider'),
                    llm_model=result.get('llm_model'),
                    raw_llm_responses=result.get('raw_llm_responses', [])
                )
                db.add(new_result)
            
            db.commit()
            
            results.append({
                "url": corr.get('url', ''),
                "auth_status": result.get('auth_status'),
                "threat_level": result.get('threat_level')
            })
            
        except Exception as e:
            logger.bind(router="api_audit", endpoint="test_all_credential_urls").exception(f"Failed to test {corr.get('url', '')}: {e}")
            results.append({
                "url": corr.get('url', ''),
                "auth_status": "failed",
                "error": str(e)
            })
    
    return {
        "success": True,
        "tested": len(results),
        "results": results
    }


@router.get("/{project_id}/api-audit/credential-url-results", dependencies=[Depends(require_permissions("admin:manage"))],
    summary="Get all credential-URL results",
    responses={401: {"description": "Not authenticated"}, 403: {"description": "Insufficient permissions - requires admin:manage"}, 404: {"description": "Project not found"}})
async def get_credential_url_results(
    project_id: str,
    db: Session = Depends(get_tenant_db)
):
    """
    Get all credential-URL test results for a project.
    Results are scoped to the current organization in multi-tenant mode.
    Requires admin:manage permission.
    """
    project = db.query(models.Repository).filter(
        models.Repository.id == project_id
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # Get current organization for multi-tenant scoping
    org_id = _get_current_organization_id(db)
    
    results_query = db.query(models.CredentialUrlTestResult).filter(
        models.CredentialUrlTestResult.repository_id == project_id
    )
    if org_id:
        results_query = results_query.filter(
            models.CredentialUrlTestResult.organization_id == org_id
        )
    results = results_query.order_by(models.CredentialUrlTestResult.tested_at.desc()).all()
    
    return {
        "success": True,
        "count": len(results),
        "results": [
            {
                "id": str(r.id),
                "target_url": r.target_url,
                "credential_type": r.credential_type,
                "auth_status": r.auth_status,
                "auth_status_code": r.auth_status_code,
                "discovered_paths_count": r.discovered_paths_count,
                "threat_level": r.threat_level,
                "tested_at": r.tested_at.isoformat() if r.tested_at else None
            }
            for r in results
        ]
    }


@router.get("/{project_id}/api-audit/credential-url-results/{result_id}", dependencies=[Depends(require_permissions("admin:manage"))],
    summary="Get detailed credential-URL test result",
    responses={401: {"description": "Not authenticated"}, 403: {"description": "Insufficient permissions - requires admin:manage"}, 404: {"description": "Result not found"}})
async def get_credential_url_result_detail(
    project_id: str,
    result_id: str,
    db: Session = Depends(get_tenant_db)
):
    """
    Get detailed results for a specific credential-URL test including full
    unmasked credential values, discovered paths, and AI analysis.
    Requires admin:manage permission.
    """
    # Get current organization for multi-tenant scoping
    org_id = _get_current_organization_id(db)
    
    result_query = db.query(models.CredentialUrlTestResult).filter(
        models.CredentialUrlTestResult.id == result_id,
        models.CredentialUrlTestResult.repository_id == project_id
    )
    if org_id:
        result_query = result_query.filter(
            models.CredentialUrlTestResult.organization_id == org_id
        )
    result = result_query.first()
    
    if not result:
        raise HTTPException(status_code=404, detail="Result not found")
    
    return {
        "success": True,
        "result": {
            "id": str(result.id),
            "target_url": result.target_url,
            "credential_type": result.credential_type,
            "credential_value": result.credential_value,  # Full unmasked value for security analyst validation
            "credential_environment": result.credential_environment,
            "confidence_score": result.confidence_score,
            "auth_status": result.auth_status,
            "auth_status_code": result.auth_status_code,
            "auth_response_time_ms": result.auth_response_time_ms,
            "auth_error_message": result.auth_error_message,
            "auth_headers_used": result.auth_headers_used,
            "auth_request_method": result.auth_request_method,
            "auth_request_url": result.auth_request_url,
            "auth_request_headers": result.auth_request_headers,  # Full request headers with actual values
            "auth_request_body": result.auth_request_body,
            "auth_response_headers": result.auth_response_headers,
            "auth_response_body": result.auth_response_body,
            "detected_service": result.detected_service,
            "service_detection_score": result.service_detection_score,
            "discovered_paths": result.discovered_paths,
            "discovered_paths_count": result.discovered_paths_count,
            "hidden_paths_found": result.hidden_paths_found,
            "sample_data_retrieved": result.sample_data_retrieved,
            "data_sensitivity_indicators": result.data_sensitivity_indicators,
            "osint_findings": result.osint_findings,
            "github_repos_found": result.github_repos_found,
            "documentation_links_found": result.documentation_links_found,
            "ai_overview": result.ai_overview,
            "ai_risk_assessment": result.ai_risk_assessment,
            "ai_recommendations": result.ai_recommendations,
            "threat_level": result.threat_level,
            "test_mode": result.test_mode,
            "tested_at": result.tested_at.isoformat() if result.tested_at else None,
            "test_duration_seconds": result.test_duration_seconds,
            "llm_provider": result.llm_provider,
            "llm_model": result.llm_model
        }
    }


@router.get("/{project_id}/api-audit/credential-url-results/{result_id}/download", dependencies=[Depends(require_permissions("admin:manage"))],
    summary="Download credential-URL test result",
    responses={401: {"description": "Not authenticated"}, 403: {"description": "Insufficient permissions - requires admin:manage"}, 404: {"description": "Result not found"}, 400: {"description": "Unsupported format"}})
async def download_credential_url_result(
    project_id: str,
    result_id: str,
    format: str = "json",
    filename: Optional[str] = None,
    db: Session = Depends(get_tenant_db)
):
    """
    Download credential-URL test results in various formats.
    Supported formats: json, csv, markdown, pdf, docx.
    Requires admin:manage permission.
    """
    from fastapi.responses import Response
    
    result = db.query(models.CredentialUrlTestResult).filter(
        models.CredentialUrlTestResult.id == result_id,
        models.CredentialUrlTestResult.repository_id == project_id
    ).first()
    
    if not result:
        raise HTTPException(status_code=404, detail="Result not found")
    
    # Get project name for filename
    project = db.query(models.Repository).filter(
        models.Repository.id == project_id
    ).first()
    project_name = project.name if project else "unknown"
    
    # Default filename
    if not filename:
        filename = f"credential_url_report_{project_name}_{result_id[:8]}"
    
    # Build report data
    report_data = {
        "title": f"Credential-URL Security Test Report",
        "project": project_name,
        "target_url": result.target_url,
        "credential_type": result.credential_type,
        "tested_at": result.tested_at.isoformat() if result.tested_at else "N/A",
        "test_mode": result.test_mode,
        "test_duration": f"{result.test_duration_seconds}s" if result.test_duration_seconds else "N/A",
        "authentication": {
            "status": result.auth_status,
            "status_code": result.auth_status_code,
            "response_time_ms": result.auth_response_time_ms,
            "error": result.auth_error_message,
            "headers_used": result.auth_headers_used,
            "detected_service": result.detected_service,
            "service_detection_score": result.service_detection_score
        },
        "raw_request": {
            "method": result.auth_request_method or "GET",
            "url": result.auth_request_url or result.target_url,
            "headers": result.auth_request_headers or {},
            "body": result.auth_request_body or ""
        },
        "raw_response": {
            "status_code": result.auth_status_code,
            "headers": result.auth_response_headers or {},
            "body": result.auth_response_body or "",
            "body_truncated": result.auth_response_body_truncated or False
        },
        "path_discovery": {
            "total_discovered": result.discovered_paths_count,
            "hidden_paths": result.hidden_paths_found,
            "paths": result.discovered_paths
        },
        "data_analysis": {
            "samples": result.sample_data_retrieved,
            "sensitivity_indicators": result.data_sensitivity_indicators
        },
        "osint": {
            "total_findings": len(result.osint_findings or []),
            "github_repos": result.github_repos_found,
            "documentation_links": result.documentation_links_found,
            "findings": result.osint_findings
        },
        "ai_analysis": {
            "overview": result.ai_overview,
            "risk_assessment": result.ai_risk_assessment,
            "threat_level": result.threat_level,
            "recommendations": result.ai_recommendations
        },
        "metadata": {
            "llm_provider": result.llm_provider,
            "llm_model": result.llm_model
        }
    }
    
    if format.lower() == "json":
        content = json.dumps(report_data, indent=2, default=str)
        return Response(
            content=content,
            media_type="application/json",
            headers={"Content-Disposition": f"attachment; filename={filename}.json"}
        )
    
    elif format.lower() == "csv":
        # Flatten for CSV
        import csv
        import io
        
        output = io.StringIO()
        writer = csv.writer(output)
        
        # Write headers and data for discovered paths
        writer.writerow(["Section", "Field", "Value"])
        writer.writerow(["Target", "URL", result.target_url])
        writer.writerow(["Target", "Credential Type", result.credential_type])
        writer.writerow(["Authentication", "Status", result.auth_status])
        writer.writerow(["Authentication", "HTTP Code", result.auth_status_code])
        writer.writerow(["Discovery", "Paths Found", result.discovered_paths_count])
        writer.writerow(["Discovery", "Hidden Paths", result.hidden_paths_found])
        writer.writerow(["OSINT", "GitHub Repos", result.github_repos_found])
        writer.writerow(["Analysis", "Threat Level", result.threat_level])
        
        # Add discovered paths
        writer.writerow([])
        writer.writerow(["Discovered Paths"])
        writer.writerow(["Method", "Path", "Status Code", "Success"])
        for path in (result.discovered_paths or []):
            writer.writerow([
                path.get('method', ''),
                path.get('path', ''),
                path.get('status_code', ''),
                path.get('success', '')
            ])
        
        content = output.getvalue()
        return Response(
            content=content,
            media_type="text/csv",
            headers={"Content-Disposition": f"attachment; filename={filename}.csv"}
        )
    
    elif format.lower() == "markdown":
        md = f"""# Credential-URL Security Test Report

## Target Information
- **URL**: {result.target_url}
- **Credential Type**: {result.credential_type}
- **Environment**: {result.credential_environment or 'N/A'}
- **Test Mode**: {result.test_mode}
- **Tested At**: {result.tested_at.isoformat() if result.tested_at else 'N/A'}
- **Duration**: {result.test_duration_seconds}s

## Overview
{result.ai_overview or 'No overview available.'}

## Authentication Status
| Field | Value |
|-------|-------|
| Status | **{result.auth_status.upper()}** |
| HTTP Code | {result.auth_status_code or 'N/A'} |
| Response Time | {result.auth_response_time_ms}ms |
| Detected Service | {result.detected_service or 'Unknown'} ({result.service_detection_score or 0}% confidence) |
| Error | {result.auth_error_message or 'None'} |

## Raw HTTP Request

```http
{result.auth_request_method or 'GET'} {result.auth_request_url or result.target_url} HTTP/1.1
"""
        # Add request headers
        req_headers = result.auth_request_headers or {}
        for hdr_key, hdr_val in req_headers.items():
            md += f"{hdr_key}: {hdr_val}\n"
        
        if result.auth_request_body:
            md += f"\n{result.auth_request_body}\n"
        
        md += f"""```

## Raw HTTP Response

```http
HTTP/1.1 {result.auth_status_code or 'N/A'} {result.auth_status.upper() if result.auth_status == 'yes' else 'FAILED'}
"""
        # Add response headers
        resp_headers = result.auth_response_headers or {}
        for hdr_key, hdr_val in resp_headers.items():
            md += f"{hdr_key}: {hdr_val}\n"
        
        if result.auth_response_body:
            body_preview = result.auth_response_body[:2000] if len(result.auth_response_body or '') > 2000 else result.auth_response_body
            md += f"\n{body_preview}"
            if result.auth_response_body_truncated or len(result.auth_response_body or '') > 2000:
                md += "\n... [truncated]"
        
        md += f"""
```

## Discovered Paths ({result.discovered_paths_count} total, {result.hidden_paths_found} hidden)

| Method | Path | Status | Success |
|--------|------|--------|---------|
"""
        for path in (result.discovered_paths or [])[:50]:
            md += f"| {path.get('method', '')} | {path.get('path', '')} | {path.get('status_code', '')} | {'✅' if path.get('success') else '❌'} |\n"
        
        md += f"""
## OSINT Findings ({len(result.osint_findings or [])} sources)

| URL | Type | Relevance |
|-----|------|-----------|
"""
        for finding in (result.osint_findings or [])[:20]:
            md += f"| {finding.get('url', '')[:60]} | {finding.get('type', '')} | {finding.get('relevance', '')}% |\n"
        
        md += f"""
## Risk Assessment
**Threat Level**: {result.threat_level.upper() if result.threat_level else 'N/A'}

{result.ai_risk_assessment or 'No risk assessment available.'}

## Recommendations
"""
        for rec in (result.ai_recommendations or []):
            md += f"- {rec}\n"
        
        md += f"""
---
*Generated by AuditGH AI Credential-URL Testing Agent*
*LLM: {result.llm_provider or 'Rule-based'} / {result.llm_model or 'N/A'}*
"""
        
        return Response(
            content=md,
            media_type="text/markdown",
            headers={"Content-Disposition": f"attachment; filename={filename}.md"}
        )
    
    elif format.lower() == "pdf":
        # Generate PDF using reportlab
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib import colors
            import io
            
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Title
            title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=18, spaceAfter=20)
            story.append(Paragraph("Credential-URL Security Test Report", title_style))
            story.append(Spacer(1, 12))
            
            # Target info
            story.append(Paragraph(f"<b>Target URL:</b> {result.target_url}", styles['Normal']))
            story.append(Paragraph(f"<b>Credential Type:</b> {result.credential_type}", styles['Normal']))
            story.append(Paragraph(f"<b>Test Mode:</b> {result.test_mode}", styles['Normal']))
            story.append(Paragraph(f"<b>Tested At:</b> {result.tested_at}", styles['Normal']))
            story.append(Spacer(1, 20))
            
            # Overview
            story.append(Paragraph("Overview", styles['Heading2']))
            story.append(Paragraph(result.ai_overview or "No overview available.", styles['Normal']))
            story.append(Spacer(1, 12))
            
            # Auth status
            story.append(Paragraph("Authentication Status", styles['Heading2']))
            auth_data = [
                ["Status", result.auth_status.upper()],
                ["HTTP Code", str(result.auth_status_code or "N/A")],
                ["Response Time", f"{result.auth_response_time_ms}ms"],
            ]
            auth_table = Table(auth_data, colWidths=[150, 300])
            auth_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ]))
            story.append(auth_table)
            story.append(Spacer(1, 12))
            
            # Threat level
            story.append(Paragraph(f"<b>Threat Level:</b> {result.threat_level.upper() if result.threat_level else 'N/A'}", styles['Heading2']))
            story.append(Spacer(1, 12))
            
            # Recommendations
            story.append(Paragraph("Recommendations", styles['Heading2']))
            for rec in (result.ai_recommendations or ["No recommendations available."]):
                story.append(Paragraph(f"• {rec}", styles['Normal']))
            
            doc.build(story)
            
            buffer.seek(0)
            return Response(
                content=buffer.getvalue(),
                media_type="application/pdf",
                headers={"Content-Disposition": f"attachment; filename={filename}.pdf"}
            )
        except ImportError:
            raise HTTPException(status_code=500, detail="PDF generation requires reportlab library")
    
    elif format.lower() == "docx":
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import io
            
            doc = Document()
            
            # Title
            title = doc.add_heading("Credential-URL Security Test Report", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Target info
            doc.add_heading("Target Information", level=1)
            doc.add_paragraph(f"URL: {result.target_url}")
            doc.add_paragraph(f"Credential Type: {result.credential_type}")
            doc.add_paragraph(f"Test Mode: {result.test_mode}")
            doc.add_paragraph(f"Tested At: {result.tested_at}")
            
            # Overview
            doc.add_heading("Overview", level=1)
            doc.add_paragraph(result.ai_overview or "No overview available.")
            
            # Auth status
            doc.add_heading("Authentication Status", level=1)
            table = doc.add_table(rows=3, cols=2)
            table.style = 'Table Grid'
            table.cell(0, 0).text = "Status"
            table.cell(0, 1).text = result.auth_status.upper()
            table.cell(1, 0).text = "HTTP Code"
            table.cell(1, 1).text = str(result.auth_status_code or "N/A")
            table.cell(2, 0).text = "Response Time"
            table.cell(2, 1).text = f"{result.auth_response_time_ms}ms"
            
            # Threat level
            doc.add_heading("Risk Assessment", level=1)
            doc.add_paragraph(f"Threat Level: {result.threat_level.upper() if result.threat_level else 'N/A'}")
            doc.add_paragraph(result.ai_risk_assessment or "No risk assessment available.")
            
            # Recommendations
            doc.add_heading("Recommendations", level=1)
            for rec in (result.ai_recommendations or ["No recommendations available."]):
                doc.add_paragraph(rec, style='List Bullet')
            
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            return Response(
                content=buffer.getvalue(),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename={filename}.docx"}
            )
        except ImportError:
            raise HTTPException(status_code=500, detail="DOCX generation requires python-docx library")
    
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported format: {format}. Use: json, csv, markdown, pdf, docx")


